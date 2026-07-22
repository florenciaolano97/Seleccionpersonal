from __future__ import annotations

from pathlib import Path
import base64
import html
import smtplib
import ssl
import hashlib
import io
import os
import tempfile
import json
import re
import secrets
import sqlite3
import time
import traceback
import unicodedata
import uuid
from datetime import date, datetime, time as datetime_time, timedelta
from email.message import EmailMessage
from urllib.parse import urlencode

import pandas as pd
import streamlit as st
import streamlit.components.v1 as components


# =========================================================
# CONFIGURACIÓN
# =========================================================

APP_NAME = "ALBA v2 | Plataforma de Selección"
APP_VERSION = "6.0.0"
BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data"
DB_PATH = DATA_DIR / "alba_v2.db"
DATA_DIR.mkdir(exist_ok=True)

PERMISSIONS = {
    "manage_company": "Administrar datos y logo de la empresa",
    "manage_users": "Crear usuarios y administrar accesos",
    "manage_jobs": "Crear y administrar búsquedas",
    "view_candidates": "Ver candidatos y postulaciones",
    "manage_candidates": "Modificar etapas y decisiones",
    "manage_cv_pool": "Cargar CV y ejecutar preselección",
    "view_interviews": "Ver entrevistas y respuestas",
    "manage_interviews": "Crear, asignar y evaluar entrevistas",
    "view_audit": "Consultar auditoría",
}

ROLE_DEFAULTS = {
    "ADMIN": list(PERMISSIONS.keys()),
    "RECRUITER": [
        "manage_jobs",
        "view_candidates",
        "manage_candidates",
        "manage_cv_pool",
        "view_interviews",
        "manage_interviews",
    ],
    "VIEWER": [
        "view_candidates",
        "view_interviews",
    ],
}


APPLICATION_STATUSES = [
    "RECIBIDA",
    "EN REVISIÓN",
    "ENTREVISTA",
    "FINALISTA",
    "SELECCIONADA",
    "RECHAZADA",
]


INTERVIEW_STATUSES = [
    "INVITADA",
    "EN CURSO",
    "COMPLETADA",
    "EVALUADA",
    "VENCIDA",
    "CANCELADA",
]

INTERVIEW_RECOMMENDATIONS = [
    "AVANZA",
    "REVISAR",
    "NO AVANZA",
]

MAX_AUDIO_SIZE_BYTES = 10 * 1024 * 1024

INTERVIEW_NOTICE = (
    "La asistencia automática analiza únicamente respuestas escritas o "
    "transcripciones. No evalúa voz, acento, apariencia, gestos ni emociones. "
    "La decisión final debe ser revisada y justificada por una persona."
)

AVATAR_NOTICE = (
    "ALBA presenta las preguntas mediante un avatar parlante. "
    "El avatar no observa, puntúa ni analiza al candidato."
)

DID_API_BASE = "https://api.d-id.com"
DID_POLL_ATTEMPTS = 24
DID_POLL_SECONDS = 2

INVITATION_DEFAULT_DAYS = 7
INVITATION_EMAIL_NOTICE = (
    "ALBA utiliza primero el correo detectado en el CV. "
    "RR. HH. puede revisarlo y corregirlo antes del envío."
)


RECORDINGS_DIR = DATA_DIR / "recordings"
RECORDINGS_DIR.mkdir(exist_ok=True)

COMPETENCY_SCALE = {
    1: {
        "level": "Muy por debajo de lo esperado",
        "description": (
            "No demuestra la competencia o presenta conductas contrarias "
            "a lo requerido para el puesto."
        ),
    },
    2: {
        "level": "Bajo",
        "description": (
            "Demuestra la competencia de manera limitada, inconsistente "
            "o con necesidad importante de acompañamiento."
        ),
    },
    3: {
        "level": "Cumple lo esperado",
        "description": (
            "Demuestra la competencia en el nivel necesario para desempeñar "
            "el puesto de manera adecuada."
        ),
    },
    4: {
        "level": "Sobre lo esperado",
        "description": (
            "Demuestra la competencia con solidez, autonomía y ejemplos "
            "que superan los requerimientos habituales."
        ),
    },
    5: {
        "level": "Excelente",
        "description": (
            "Demuestra dominio sobresaliente, genera impacto y puede servir "
            "como referente para otras personas."
        ),
    },
}

COMPETENCY_GUIDE_DEFAULTS = {
    "Comunicación": {
        "definition": (
            "Capacidad para expresar ideas con claridad, escuchar, adaptar "
            "el mensaje y confirmar la comprensión."
        ),
        "keywords": [
            "claridad", "escucha", "audiencia", "mensaje", "comprensión",
            "feedback", "explicar", "comunicación",
        ],
        "questions": [
            "Contanos una situación en la que tuviste que explicar información compleja.",
            "¿Cómo actuaste cuando una comunicación generó una confusión?",
        ],
    },
    "Trabajo en Equipo": {
        "definition": (
            "Capacidad para colaborar, integrar miradas distintas y contribuir "
            "al logro colectivo."
        ),
        "keywords": [
            "equipo", "colaboración", "acuerdo", "apoyo", "conflicto",
            "objetivo común", "escucha",
        ],
        "questions": [
            "Relatá una experiencia con un equipo de ideas diferentes.",
            "¿Qué rol asumís normalmente en los trabajos grupales?",
        ],
    },
    "Orientación a Resultados": {
        "definition": (
            "Capacidad para definir objetivos, priorizar, hacer seguimiento "
            "y obtener resultados medibles."
        ),
        "keywords": [
            "objetivo", "resultado", "indicador", "impacto", "plazo",
            "seguimiento", "meta", "mejora",
        ],
        "questions": [
            "Contanos un objetivo desafiante que hayas alcanzado.",
            "¿Cómo mediste el impacto de tu trabajo?",
        ],
    },
    "Adaptabilidad": {
        "definition": (
            "Capacidad para ajustar conductas, prioridades y aprendizajes "
            "frente a cambios o situaciones nuevas."
        ),
        "keywords": [
            "cambio", "adaptación", "aprendizaje", "flexibilidad",
            "prioridad", "nuevo", "ajuste",
        ],
        "questions": [
            "Describí un cambio importante y cómo te adaptaste.",
            "¿Qué hacés cuando cambia una prioridad de forma inesperada?",
        ],
    },
    "Resolución de Problemas": {
        "definition": (
            "Capacidad para analizar causas, generar alternativas, tomar "
            "decisiones y verificar resultados."
        ),
        "keywords": [
            "problema", "causa", "análisis", "alternativas", "decisión",
            "solución", "resultado", "riesgo",
        ],
        "questions": [
            "Relatá un problema complejo que hayas resuelto.",
            "¿Cómo evaluaste las alternativas antes de decidir?",
        ],
    },
    "Organización": {
        "definition": (
            "Capacidad para planificar, priorizar y cumplir compromisos "
            "administrando tiempo y recursos."
        ),
        "keywords": [
            "planificación", "prioridad", "agenda", "plazo",
            "seguimiento", "organización", "recursos",
        ],
        "questions": [
            "¿Cómo organizaste varias tareas con fechas similares?",
            "¿Qué herramientas utilizás para hacer seguimiento?",
        ],
    },
    "Iniciativa": {
        "definition": (
            "Capacidad para anticiparse, proponer mejoras y actuar de forma "
            "responsable sin esperar instrucciones permanentes."
        ),
        "keywords": [
            "iniciativa", "propuesta", "anticipación", "mejora",
            "acción", "autonomía", "implementación",
        ],
        "questions": [
            "Contanos una mejora que propusiste sin que te la solicitaran.",
            "¿Qué oportunidad detectaste y cómo actuaste?",
        ],
    },
}

INTERVIEW_RECORDING_NOTICE = (
    "La grabación desde el navegador registra la cámara y el micrófono del "
    "dispositivo actual. Para guardar una videollamada externa completa, "
    "subí el archivo exportado por la plataforma utilizada."
)

DEFAULT_INTERVIEW_TEMPLATES = {
    "Entrevista inicial general": {
        "category": "GENERAL",
        "description": (
            "Entrevista estructurada para conocer motivación, experiencia, "
            "fortalezas, aprendizaje y expectativas."
        ),
        "questions": [
            {
                "competency": "Motivación",
                "question": (
                    "¿Qué te interesó de esta oportunidad y por qué considerás "
                    "que tu experiencia puede aportar al puesto?"
                ),
                "indicators": (
                    "motivación, conocimiento del puesto, aporte, experiencia, "
                    "interés, objetivos"
                ),
            },
            {
                "competency": "Experiencia relevante",
                "question": (
                    "Contanos una experiencia o proyecto relacionado con las "
                    "responsabilidades de esta posición. ¿Cuál fue tu aporte?"
                ),
                "indicators": (
                    "situación, tarea, acción, resultado, aporte, aprendizaje, "
                    "responsabilidad"
                ),
            },
            {
                "competency": "Resolución de problemas",
                "question": (
                    "Describí un problema laboral o académico complejo que "
                    "hayas tenido que resolver. ¿Cómo lo abordaste y qué ocurrió?"
                ),
                "indicators": (
                    "problema, análisis, alternativas, decisión, acción, "
                    "resultado, aprendizaje"
                ),
            },
            {
                "competency": "Trabajo en equipo",
                "question": (
                    "Contanos una situación en la que tuviste que trabajar con "
                    "personas con ideas diferentes. ¿Qué hiciste?"
                ),
                "indicators": (
                    "escucha, colaboración, comunicación, acuerdo, conflicto, "
                    "resultado, equipo"
                ),
            },
            {
                "competency": "Autoconocimiento",
                "question": (
                    "¿Cuáles son tus principales fortalezas para este rol y qué "
                    "aspecto estás trabajando actualmente para mejorar?"
                ),
                "indicators": (
                    "fortaleza, ejemplo, mejora, aprendizaje, desarrollo, "
                    "autoconocimiento"
                ),
            },
        ],
    },
    "Entrevista por competencias STAR": {
        "category": "COMPETENCIAS",
        "description": (
            "Banco de preguntas conductuales con estructura Situación, Tarea, "
            "Acción y Resultado."
        ),
        "questions": [
            {
                "competency": "Adaptabilidad",
                "question": (
                    "Relatá una situación en la que cambió una prioridad, "
                    "herramienta o forma de trabajo. ¿Cómo te adaptaste?"
                ),
                "indicators": (
                    "situación, cambio, acción, adaptación, resultado, "
                    "aprendizaje"
                ),
            },
            {
                "competency": "Iniciativa",
                "question": (
                    "Contanos una ocasión en la que propusiste o implementaste "
                    "una mejora sin que te la solicitaran."
                ),
                "indicators": (
                    "iniciativa, propuesta, acción, mejora, impacto, resultado"
                ),
            },
            {
                "competency": "Organización",
                "question": (
                    "Describí un momento en el que tuviste varias tareas o "
                    "fechas límite al mismo tiempo. ¿Cómo te organizaste?"
                ),
                "indicators": (
                    "priorización, planificación, organización, seguimiento, "
                    "plazo, resultado"
                ),
            },
            {
                "competency": "Comunicación",
                "question": (
                    "Contanos una situación en la que tuviste que explicar "
                    "información compleja o difícil a otra persona."
                ),
                "indicators": (
                    "audiencia, claridad, escucha, adaptación, comprensión, "
                    "resultado"
                ),
            },
            {
                "competency": "Orientación a resultados",
                "question": (
                    "Relatá un objetivo desafiante que hayas alcanzado. "
                    "¿Qué acciones realizaste y cómo mediste el resultado?"
                ),
                "indicators": (
                    "objetivo, indicador, acción, seguimiento, resultado, "
                    "impacto"
                ),
            },
        ],
    },
    "Jóvenes profesionales y pasantías": {
        "category": "JUNIOR",
        "description": (
            "Entrevista para perfiles con experiencia inicial, centrada en "
            "aprendizaje, potencial, colaboración e iniciativa."
        ),
        "questions": [
            {
                "competency": "Motivación y aprendizaje",
                "question": (
                    "¿Qué esperás aprender en esta experiencia y cómo se "
                    "relaciona con tus objetivos profesionales?"
                ),
                "indicators": (
                    "aprendizaje, objetivos, motivación, carrera, desarrollo"
                ),
            },
            {
                "competency": "Aplicación de conocimientos",
                "question": (
                    "Contanos un trabajo práctico, proyecto o experiencia "
                    "académica de la que te sientas orgulloso/a."
                ),
                "indicators": (
                    "proyecto, responsabilidad, herramientas, aporte, "
                    "resultado, aprendizaje"
                ),
            },
            {
                "competency": "Trabajo en equipo",
                "question": (
                    "Describí una experiencia de trabajo grupal. ¿Qué rol "
                    "asumiste y cómo contribuiste?"
                ),
                "indicators": (
                    "rol, colaboración, comunicación, aporte, conflicto, "
                    "resultado"
                ),
            },
            {
                "competency": "Organización",
                "question": (
                    "¿Cómo organizás estudio, trabajo y otras responsabilidades "
                    "cuando coinciden varias entregas?"
                ),
                "indicators": (
                    "agenda, prioridad, planificación, plazo, seguimiento"
                ),
            },
            {
                "competency": "Iniciativa",
                "question": (
                    "Contanos una situación en la que buscaste aprender algo "
                    "nuevo o resolver una dificultad por tu cuenta."
                ),
                "indicators": (
                    "iniciativa, búsqueda, aprendizaje, acción, resultado"
                ),
            },
        ],
    },
    "Liderazgo": {
        "category": "LIDERAZGO",
        "description": (
            "Entrevista para posiciones de supervisión, jefatura y liderazgo."
        ),
        "questions": [
            {
                "competency": "Liderazgo de equipos",
                "question": (
                    "Contanos una situación en la que tuviste que conducir a "
                    "un equipo hacia un objetivo exigente."
                ),
                "indicators": (
                    "objetivo, equipo, planificación, comunicación, "
                    "seguimiento, resultado"
                ),
            },
            {
                "competency": "Gestión de conflictos",
                "question": (
                    "Describí un conflicto entre integrantes del equipo y cómo "
                    "interveniste."
                ),
                "indicators": (
                    "escucha, conflicto, mediación, acuerdo, seguimiento, "
                    "resultado"
                ),
            },
            {
                "competency": "Toma de decisiones",
                "question": (
                    "Relatá una decisión difícil que hayas tenido que tomar "
                    "con información limitada."
                ),
                "indicators": (
                    "información, alternativas, riesgo, decisión, resultado, "
                    "aprendizaje"
                ),
            },
            {
                "competency": "Desarrollo de personas",
                "question": (
                    "¿Cómo identificaste y acompañaste una necesidad de "
                    "desarrollo en una persona de tu equipo?"
                ),
                "indicators": (
                    "observación, feedback, plan, acompañamiento, mejora, "
                    "resultado"
                ),
            },
            {
                "competency": "Mejora continua",
                "question": (
                    "Contanos una mejora de proceso que hayas liderado. "
                    "¿Qué impacto generó?"
                ),
                "indicators": (
                    "problema, análisis, mejora, implementación, indicador, "
                    "impacto"
                ),
            },
        ],
    },
}

MAX_CV_SIZE_BYTES = 5 * 1024 * 1024
MAX_CV_BATCH = 50

SCORING_NOTICE = (
    "El puntaje es una ayuda de preselección y no reemplaza la decisión "
    "humana. No utiliza nombre, correo, DNI, edad, teléfono, país, "
    "provincia ni localidad."
)

AGE_FILTER_NOTICE = (
    "El filtro de edad es exclusivamente administrativo. No modifica "
    "el puntaje, la recomendación ni el orden del ranking. Debe utilizarse "
    "solo cuando exista una razón objetiva y documentada."
)

SUGGESTION_MIN_SCORE = 35.0
SUGGESTION_TOP_LIMIT = 20

ANALYSIS_PRESETS = {
    "General": [],
    "Tecnología / Sistemas": [
        "sql",
        "git",
        "api",
        "metodologias agiles",
        "resolucion de problemas",
    ],
    "Ingeniería / Producción": [
        "mejora continua",
        "calidad",
        "procesos",
        "seguridad",
        "lean",
    ],
    "Compras / Logística": [
        "excel",
        "proveedores",
        "negociacion",
        "sap",
        "ingles",
    ],
    "Recursos Humanos": [
        "seleccion",
        "entrevistas",
        "capacitacion",
        "legislacion laboral",
        "excel",
    ],
    "Comercial / Ventas": [
        "ventas",
        "clientes",
        "negociacion",
        "crm",
        "objetivos",
    ],
    "Administración / Finanzas": [
        "excel",
        "analisis",
        "presupuestos",
        "sap",
        "contabilidad",
    ],
}


# =========================================================
# UTILIDADES
# =========================================================

def now_iso() -> str:
    return datetime.now().isoformat(timespec="seconds")


def normalize_email(email: str) -> str:
    return (email or "").strip().lower()


def validate_email(email: str) -> bool:
    email = normalize_email(email)
    return bool(re.fullmatch(r"[^@\s]+@[^@\s]+\.[^@\s]+", email))


def validate_password(password: str) -> str | None:
    if len(password) < 8:
        return "La contraseña debe tener al menos 8 caracteres."
    if not re.search(r"[A-Za-z]", password):
        return "La contraseña debe incluir al menos una letra."
    if not re.search(r"\d", password):
        return "La contraseña debe incluir al menos un número."
    return None


def hash_password(password: str, salt: str | None = None) -> tuple[str, str]:
    salt = salt or secrets.token_hex(16)
    digest = hashlib.pbkdf2_hmac(
        "sha256",
        password.encode("utf-8"),
        salt.encode("utf-8"),
        180_000,
    )
    return digest.hex(), salt


def verify_password(password: str, expected_hash: str, salt: str) -> bool:
    current_hash, _ = hash_password(password, salt)
    return secrets.compare_digest(current_hash, expected_hash)


def validate_logo(uploaded_file) -> tuple[bytes, str]:
    if uploaded_file is None:
        raise ValueError("Tenés que subir el logo de la empresa.")

    raw = uploaded_file.getvalue()
    if not raw:
        raise ValueError("El archivo del logo está vacío.")

    try:
        from PIL import Image
    except ImportError as exc:
        raise RuntimeError(
            "No está instalada la dependencia Pillow. "
            "Revisá requirements.txt y reiniciá la aplicación."
        ) from exc

    try:
        image = Image.open(io.BytesIO(raw))
        image.verify()
    except Exception as exc:
        raise ValueError("El archivo subido no es una imagen válida.") from exc

    mime_type = uploaded_file.type or "image/png"
    if mime_type not in {
        "image/png",
        "image/jpeg",
        "image/webp",
    }:
        raise ValueError("El logo debe ser PNG, JPG, JPEG o WEBP.")

    return raw, mime_type


# =========================================================
# BASE DE DATOS
# =========================================================

def get_connection() -> sqlite3.Connection:
    conn = sqlite3.connect(
        DB_PATH,
        check_same_thread=False,
        timeout=30,
    )
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    conn.execute("PRAGMA busy_timeout = 30000")
    return conn


def table_columns(table_name: str) -> set[str]:
    conn = get_connection()
    try:
        rows = conn.execute(
            f"PRAGMA table_info({table_name})"
        ).fetchall()
        return {row["name"] for row in rows}
    finally:
        conn.close()


def add_column_if_missing(
    table_name: str,
    column_name: str,
    definition: str,
) -> None:
    if column_name in table_columns(table_name):
        return

    conn = get_connection()
    try:
        conn.execute(
            f"ALTER TABLE {table_name} "
            f"ADD COLUMN {column_name} {definition}"
        )
        conn.commit()
    finally:
        conn.close()


def init_db() -> None:
    conn = get_connection()
    try:
        cur = conn.cursor()

        cur.execute("""
            CREATE TABLE IF NOT EXISTS companies (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL,
                industry TEXT,
                country TEXT,
                logo_blob BLOB,
                logo_mime TEXT,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                email TEXT NOT NULL UNIQUE,
                password_hash TEXT NOT NULL,
                password_salt TEXT NOT NULL,
                full_name TEXT NOT NULL,
                account_type TEXT NOT NULL,
                company_id INTEGER,
                role TEXT,
                permissions_json TEXT DEFAULT '[]',
                active INTEGER DEFAULT 1,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(company_id) REFERENCES companies(id)
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS candidates (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL UNIQUE,
                phone TEXT,
                dni TEXT,
                city TEXT,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(user_id) REFERENCES users(id)
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS jobs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                company_id INTEGER NOT NULL,
                title TEXT NOT NULL,
                area TEXT,
                seniority TEXT,
                description TEXT,
                status TEXT DEFAULT 'ABIERTA',
                created_by INTEGER,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(company_id) REFERENCES companies(id),
                FOREIGN KEY(created_by) REFERENCES users(id)
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS applications (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                candidate_id INTEGER NOT NULL,
                job_id INTEGER NOT NULL,
                status TEXT DEFAULT 'RECIBIDA',
                created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                UNIQUE(candidate_id, job_id),
                FOREIGN KEY(candidate_id) REFERENCES candidates(id),
                FOREIGN KEY(job_id) REFERENCES jobs(id)
            )
        """)


        cur.execute("""
            CREATE TABLE IF NOT EXISTS candidate_documents (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                candidate_id INTEGER NOT NULL,
                filename TEXT NOT NULL,
                mime_type TEXT,
                content_blob BLOB NOT NULL,
                extracted_text TEXT,
                parsed_json TEXT,
                uploaded_by INTEGER,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(candidate_id) REFERENCES candidates(id),
                FOREIGN KEY(uploaded_by) REFERENCES users(id)
            )
        """)


        cur.execute("""
            CREATE TABLE IF NOT EXISTS company_candidates (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                company_id INTEGER NOT NULL,
                candidate_id INTEGER NOT NULL,
                source TEXT,
                added_by INTEGER,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                UNIQUE(company_id, candidate_id),
                FOREIGN KEY(company_id) REFERENCES companies(id),
                FOREIGN KEY(candidate_id) REFERENCES candidates(id),
                FOREIGN KEY(added_by) REFERENCES users(id)
            )
        """)


        cur.execute("""
            CREATE TABLE IF NOT EXISTS candidate_job_suggestions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                company_id INTEGER NOT NULL,
                candidate_id INTEGER NOT NULL,
                job_id INTEGER NOT NULL,
                score_total REAL,
                recommendation TEXT,
                report_json TEXT,
                status TEXT DEFAULT 'ACTIVA',
                generated_at TEXT DEFAULT CURRENT_TIMESTAMP,
                UNIQUE(company_id, candidate_id, job_id),
                FOREIGN KEY(company_id) REFERENCES companies(id),
                FOREIGN KEY(candidate_id) REFERENCES candidates(id),
                FOREIGN KEY(job_id) REFERENCES jobs(id)
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS job_scoring_config (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                job_id INTEGER NOT NULL UNIQUE,
                preset_name TEXT,
                criteria_json TEXT NOT NULL,
                thresholds_json TEXT NOT NULL,
                updated_by INTEGER,
                updated_at TEXT DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(job_id) REFERENCES jobs(id),
                FOREIGN KEY(updated_by) REFERENCES users(id)
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS avatar_settings (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                company_id INTEGER NOT NULL UNIQUE,
                enabled INTEGER DEFAULT 0,
                display_name TEXT DEFAULT 'ALBA',
                source_url TEXT,
                voice_provider TEXT DEFAULT 'microsoft',
                voice_id TEXT,
                welcome_text TEXT,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                updated_at TEXT,
                FOREIGN KEY(company_id) REFERENCES companies(id)
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS avatar_question_videos (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                company_id INTEGER NOT NULL,
                question_id INTEGER NOT NULL,
                content_hash TEXT NOT NULL,
                talk_id TEXT,
                status TEXT DEFAULT 'PENDIENTE',
                result_url TEXT,
                error_text TEXT,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                updated_at TEXT,
                UNIQUE(company_id, question_id, content_hash),
                FOREIGN KEY(company_id) REFERENCES companies(id),
                FOREIGN KEY(question_id) REFERENCES interview_questions(id)
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS interview_templates (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                company_id INTEGER NOT NULL,
                name TEXT NOT NULL,
                description TEXT,
                category TEXT,
                active INTEGER DEFAULT 1,
                created_by INTEGER,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                updated_at TEXT,
                FOREIGN KEY(company_id) REFERENCES companies(id),
                FOREIGN KEY(created_by) REFERENCES users(id)
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS interview_questions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                template_id INTEGER NOT NULL,
                position INTEGER DEFAULT 1,
                competency TEXT,
                question_text TEXT NOT NULL,
                indicators_text TEXT,
                max_score REAL DEFAULT 5,
                required INTEGER DEFAULT 1,
                response_mode TEXT DEFAULT 'TEXTO_O_AUDIO',
                active INTEGER DEFAULT 1,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                updated_at TEXT,
                FOREIGN KEY(template_id) REFERENCES interview_templates(id)
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS interviews (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                application_id INTEGER NOT NULL,
                template_id INTEGER NOT NULL,
                status TEXT DEFAULT 'INVITADA',
                due_date TEXT,
                intro_message TEXT,
                assigned_by INTEGER,
                assigned_at TEXT DEFAULT CURRENT_TIMESTAMP,
                started_at TEXT,
                completed_at TEXT,
                auto_score REAL,
                auto_recommendation TEXT,
                auto_report_json TEXT,
                human_score REAL,
                final_recommendation TEXT,
                final_reason TEXT,
                evaluated_by INTEGER,
                evaluated_at TEXT,
                updated_at TEXT,
                UNIQUE(application_id, template_id),
                FOREIGN KEY(application_id) REFERENCES applications(id),
                FOREIGN KEY(template_id) REFERENCES interview_templates(id),
                FOREIGN KEY(assigned_by) REFERENCES users(id),
                FOREIGN KEY(evaluated_by) REFERENCES users(id)
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS interview_invitations (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                interview_id INTEGER NOT NULL UNIQUE,
                candidate_user_id INTEGER NOT NULL,
                recipient_email TEXT NOT NULL,
                token_hash TEXT NOT NULL UNIQUE,
                expires_at TEXT NOT NULL,
                used_at TEXT,
                sent_at TEXT,
                send_status TEXT DEFAULT 'PENDIENTE',
                send_error TEXT,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                updated_at TEXT,
                FOREIGN KEY(interview_id) REFERENCES interviews(id),
                FOREIGN KEY(candidate_user_id) REFERENCES users(id)
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS interview_responses (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                interview_id INTEGER NOT NULL,
                question_id INTEGER NOT NULL,
                answer_text TEXT,
                audio_blob BLOB,
                audio_mime TEXT,
                audio_filename TEXT,
                submitted_at TEXT,
                updated_at TEXT,
                UNIQUE(interview_id, question_id),
                FOREIGN KEY(interview_id) REFERENCES interviews(id),
                FOREIGN KEY(question_id) REFERENCES interview_questions(id)
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS interview_evaluations (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                interview_id INTEGER NOT NULL,
                question_id INTEGER NOT NULL,
                evaluator_type TEXT NOT NULL,
                score REAL,
                evidence_text TEXT,
                strengths_text TEXT,
                gaps_text TEXT,
                notes TEXT,
                evaluated_by INTEGER,
                evaluated_at TEXT,
                UNIQUE(interview_id, question_id, evaluator_type),
                FOREIGN KEY(interview_id) REFERENCES interviews(id),
                FOREIGN KEY(question_id) REFERENCES interview_questions(id),
                FOREIGN KEY(evaluated_by) REFERENCES users(id)
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS competency_guides (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                company_id INTEGER NOT NULL,
                name TEXT NOT NULL,
                definition TEXT,
                keywords_json TEXT DEFAULT '[]',
                questions_json TEXT DEFAULT '[]',
                critical_default INTEGER DEFAULT 0,
                active INTEGER DEFAULT 1,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                updated_at TEXT,
                UNIQUE(company_id, name),
                FOREIGN KEY(company_id) REFERENCES companies(id)
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS interview_sessions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                interview_id INTEGER NOT NULL UNIQUE,
                recording_blob BLOB,
                recording_mime TEXT,
                recording_filename TEXT,
                transcript_text TEXT,
                interviewer_notes TEXT,
                analysis_json TEXT,
                created_by INTEGER,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                updated_at TEXT,
                FOREIGN KEY(interview_id) REFERENCES interviews(id),
                FOREIGN KEY(created_by) REFERENCES users(id)
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS interview_competency_scores (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                interview_id INTEGER NOT NULL,
                competency_name TEXT NOT NULL,
                score INTEGER NOT NULL,
                critical INTEGER DEFAULT 0,
                evidence_text TEXT,
                notes TEXT,
                evaluated_by INTEGER,
                evaluated_at TEXT,
                UNIQUE(interview_id, competency_name),
                FOREIGN KEY(interview_id) REFERENCES interviews(id),
                FOREIGN KEY(evaluated_by) REFERENCES users(id)
            )
        """)

        cur.execute("""
            CREATE TABLE IF NOT EXISTS audit_log (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                company_id INTEGER,
                user_id INTEGER,
                event_type TEXT NOT NULL,
                entity_type TEXT,
                entity_id INTEGER,
                details TEXT,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(company_id) REFERENCES companies(id),
                FOREIGN KEY(user_id) REFERENCES users(id)
            )
        """)

        conn.commit()
    finally:
        conn.close()

    add_column_if_missing("companies", "logo_blob", "BLOB")
    add_column_if_missing("companies", "logo_mime", "TEXT")
    add_column_if_missing("users", "role", "TEXT")
    add_column_if_missing(
        "users",
        "permissions_json",
        "TEXT DEFAULT '[]'",
    )
    add_column_if_missing(
        "users",
        "active",
        "INTEGER DEFAULT 1",
    )

    for column_name, definition in {
        "location": "TEXT",
        "work_mode": "TEXT",
        "contract_type": "TEXT",
        "responsibilities": "TEXT",
        "must_have": "TEXT",
        "desirable": "TEXT",
        "competencies": "TEXT",
        "source_filename": "TEXT",
        "updated_at": "TEXT",
    }.items():
        add_column_if_missing(
            "jobs",
            column_name,
            definition,
        )


    # Compatibilidad con bases creadas por versiones anteriores.
    for column_name, definition in {
        "company_id": "INTEGER",
        "user_id": "INTEGER",
        "event_type": "TEXT DEFAULT ''",
        "entity_type": "TEXT",
        "entity_id": "INTEGER",
        "details": "TEXT",
        "created_at": "TEXT",
    }.items():
        add_column_if_missing(
            "audit_log",
            column_name,
            definition,
        )

    for column_name, definition in {
        "user_id": "INTEGER",
        "phone": "TEXT",
        "dni": "TEXT",
        "city": "TEXT",
    }.items():
        add_column_if_missing(
            "candidates",
            column_name,
            definition,
        )


    for column_name, definition in {
        "source": "TEXT DEFAULT 'PORTAL'",
        "headline": "TEXT",
        "education_summary": "TEXT",
        "experience_summary": "TEXT",
        "skills_text": "TEXT",
        "languages_text": "TEXT",
        "birth_date": "TEXT",
        "country": "TEXT",
        "province": "TEXT",
        "city": "TEXT",
        "locality": "TEXT",
        "tags_text": "TEXT",
        "updated_at": "TEXT",
    }.items():
        add_column_if_missing(
            "candidates",
            column_name,
            definition,
        )

    for column_name, definition in {
        "score_total": "REAL",
        "score_breakdown_json": "TEXT",
        "screening_recommendation": "TEXT",
        "screening_summary": "TEXT",
        "screened_at": "TEXT",
        "manual_recommendation": "TEXT",
        "manual_reason": "TEXT",
        "recruiter_notes": "TEXT",
        "assigned_recruiter_id": "INTEGER",
        "starred": "INTEGER DEFAULT 0",
        "updated_at": "TEXT",
    }.items():
        add_column_if_missing(
            "applications",
            column_name,
            definition,
        )


    for column_name, definition in {
        "competency_result_json": "TEXT",
        "traffic_light": "TEXT",
    }.items():
        add_column_if_missing(
            "interviews",
            column_name,
            definition,
        )


def fetch_all(
    query: str,
    params: tuple = (),
) -> list[dict]:
    conn = get_connection()
    try:
        rows = conn.execute(query, params).fetchall()
        return [dict(row) for row in rows]
    finally:
        conn.close()


def fetch_one(
    query: str,
    params: tuple = (),
) -> dict | None:
    conn = get_connection()
    try:
        row = conn.execute(query, params).fetchone()
        return dict(row) if row else None
    finally:
        conn.close()


def execute(
    query: str,
    params: tuple = (),
    retries: int = 3,
) -> int:
    last_error = None

    for attempt in range(retries):
        conn = get_connection()
        try:
            cur = conn.execute(query, params)
            conn.commit()
            return cur.lastrowid
        except sqlite3.OperationalError as exc:
            conn.rollback()
            last_error = exc

            if (
                "locked" in str(exc).lower()
                and attempt < retries - 1
            ):
                time.sleep(0.35 * (attempt + 1))
                continue

            raise
        finally:
            conn.close()

    if last_error:
        raise last_error

    raise RuntimeError("No se pudo ejecutar la operación.")


def log_event(
    company_id: int | None,
    user_id: int | None,
    event_type: str,
    entity_type: str = "",
    entity_id: int | None = None,
    details: dict | str | None = None,
) -> bool:
    """
    La auditoría nunca debe bloquear la operación principal.
    """
    if isinstance(details, dict):
        details = json.dumps(details, ensure_ascii=False)

    try:
        execute(
            """
            INSERT INTO audit_log(
                company_id,
                user_id,
                event_type,
                entity_type,
                entity_id,
                details,
                created_at
            )
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (
                company_id,
                user_id,
                event_type,
                entity_type,
                entity_id,
                details or "",
                now_iso(),
            ),
        )
        return True

    except Exception as exc:
        try:
            st.session_state["audit_warning"] = (
                "La acción principal se guardó, pero no pudo "
                f"registrarse en auditoría: {exc}"
            )
        except Exception:
            pass
        return False


# =========================================================
# AUTENTICACIÓN Y SESIÓN
# =========================================================

def create_user(
    email: str,
    password: str,
    full_name: str,
    account_type: str,
    company_id: int | None = None,
    role: str | None = None,
    permissions: list[str] | None = None,
) -> int:
    email = normalize_email(email)

    if not validate_email(email):
        raise ValueError("Ingresá un correo electrónico válido.")

    password_error = validate_password(password)
    if password_error:
        raise ValueError(password_error)

    if not full_name.strip():
        raise ValueError("Ingresá nombre y apellido.")

    if fetch_one("SELECT id FROM users WHERE email = ?", (email,)):
        raise ValueError("Ya existe un usuario registrado con ese correo.")

    password_hash, salt = hash_password(password)

    return execute(
        """
        INSERT INTO users(
            email, password_hash, password_salt, full_name,
            account_type, company_id, role,
            permissions_json, active, created_at
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, 1, ?)
        """,
        (
            email,
            password_hash,
            salt,
            full_name.strip(),
            account_type,
            company_id,
            role,
            json.dumps(permissions or [], ensure_ascii=False),
            now_iso(),
        ),
    )


def register_company_admin(
    company_name: str,
    industry: str,
    country: str,
    logo_file,
    full_name: str,
    email: str,
    password: str,
) -> int:
    if not company_name.strip():
        raise ValueError("Ingresá el nombre de la empresa.")

    logo_blob, logo_mime = validate_logo(logo_file)

    company_id = execute(
        """
        INSERT INTO companies(
            name, industry, country,
            logo_blob, logo_mime, created_at
        )
        VALUES (?, ?, ?, ?, ?, ?)
        """,
        (
            company_name.strip(),
            industry.strip(),
            country.strip(),
            logo_blob,
            logo_mime,
            now_iso(),
        ),
    )

    try:
        user_id = create_user(
            email=email,
            password=password,
            full_name=full_name,
            account_type="COMPANY",
            company_id=company_id,
            role="ADMIN",
            permissions=ROLE_DEFAULTS["ADMIN"],
        )
    except Exception:
        execute("DELETE FROM companies WHERE id = ?", (company_id,))
        raise

    log_event(
        company_id,
        user_id,
        "REGISTER_COMPANY",
        "company",
        company_id,
        {"company_name": company_name},
    )

    return company_id


def register_candidate(
    full_name: str,
    email: str,
    password: str,
    phone: str = "",
    dni: str = "",
    city: str = "",
) -> int:
    user_id = create_user(
        email=email,
        password=password,
        full_name=full_name,
        account_type="CANDIDATE",
    )

    candidate_id = execute(
        """
        INSERT INTO candidates(
            user_id, phone, dni, city, created_at
        )
        VALUES (?, ?, ?, ?, ?)
        """,
        (
            user_id,
            phone.strip(),
            dni.strip(),
            city.strip(),
            now_iso(),
        ),
    )

    log_event(
        None,
        user_id,
        "REGISTER_CANDIDATE",
        "candidate",
        candidate_id,
        {"email": normalize_email(email)},
    )

    return candidate_id


def authenticate(email: str, password: str) -> dict | None:
    user = fetch_one(
        "SELECT * FROM users WHERE email = ? AND active = 1",
        (normalize_email(email),),
    )

    if not user:
        return None

    if not verify_password(
        password,
        user["password_hash"],
        user["password_salt"],
    ):
        return None

    try:
        user["permissions"] = json.loads(
            user.get("permissions_json") or "[]"
        )
    except json.JSONDecodeError:
        user["permissions"] = []

    return user


def current_user() -> dict | None:
    return st.session_state.get("auth_user")


def refresh_user(user_id: int) -> dict | None:
    user = fetch_one(
        "SELECT * FROM users WHERE id = ? AND active = 1",
        (user_id,),
    )

    if not user:
        return None

    try:
        user["permissions"] = json.loads(
            user.get("permissions_json") or "[]"
        )
    except json.JSONDecodeError:
        user["permissions"] = []

    return user


def has_permission(permission: str) -> bool:
    user = current_user()

    if not user or user.get("account_type") != "COMPANY":
        return False

    if user.get("role") == "ADMIN":
        return True

    return permission in user.get("permissions", [])


def logout() -> None:
    st.session_state.pop("auth_user", None)
    st.rerun()


# =========================================================
# UI DE ACCESO
# =========================================================

def render_login() -> None:
    st.subheader("Iniciar sesión")

    with st.form("login_form"):
        email = st.text_input("Correo electrónico")
        password = st.text_input("Contraseña", type="password")
        submitted = st.form_submit_button(
            "Ingresar",
            type="primary",
            use_container_width=True,
        )

        if submitted:
            user = authenticate(email, password)

            if not user:
                st.error("Correo o contraseña incorrectos.")
            else:
                st.session_state["auth_user"] = user

                log_event(
                    user.get("company_id"),
                    user["id"],
                    "LOGIN",
                    "user",
                    user["id"],
                    {"email": user["email"]},
                )

                st.rerun()


def render_registration() -> None:
    st.subheader("Crear cuenta")

    account_type = st.radio(
        "Tipo de cuenta",
        ["Empresa", "Candidato"],
        horizontal=True,
    )

    if account_type == "Empresa":
        with st.form("register_company"):
            st.markdown("#### Datos de la empresa")

            company_name = st.text_input("Nombre de la empresa")
            c1, c2 = st.columns(2)
            industry = c1.text_input("Industria")
            country = c2.text_input("País", value="Argentina")

            logo_file = st.file_uploader(
                "Logo de la empresa",
                type=["png", "jpg", "jpeg", "webp"],
            )

            st.markdown("#### Usuario administrador")

            full_name = st.text_input("Nombre y apellido")
            email = st.text_input("Correo electrónico")
            password = st.text_input("Contraseña", type="password")
            repeat = st.text_input("Repetir contraseña", type="password")

            submitted = st.form_submit_button(
                "Crear empresa",
                type="primary",
                use_container_width=True,
            )

            if submitted:
                try:
                    if password != repeat:
                        raise ValueError("Las contraseñas no coinciden.")

                    register_company_admin(
                        company_name,
                        industry,
                        country,
                        logo_file,
                        full_name,
                        email,
                        password,
                    )

                    st.success(
                        "Empresa y usuario administrador creados. "
                        "Ya podés iniciar sesión."
                    )
                except Exception as exc:
                    st.error(str(exc))

    else:
        with st.form("register_candidate"):
            full_name = st.text_input("Nombre y apellido")
            email = st.text_input("Correo electrónico")
            password = st.text_input("Contraseña", type="password")
            repeat = st.text_input("Repetir contraseña", type="password")

            c1, c2 = st.columns(2)
            phone = c1.text_input("Teléfono")
            dni = c2.text_input("DNI")

            city = st.text_input("Ciudad")

            consent = st.checkbox(
                "Acepto el tratamiento de mis datos para gestionar "
                "postulaciones. Los datos administrativos no se usarán "
                "para puntuar ni rankear."
            )

            submitted = st.form_submit_button(
                "Crear cuenta",
                type="primary",
                use_container_width=True,
            )

            if submitted:
                try:
                    if password != repeat:
                        raise ValueError("Las contraseñas no coinciden.")
                    if not consent:
                        raise ValueError(
                            "Tenés que aceptar el consentimiento."
                        )

                    register_candidate(
                        full_name,
                        email,
                        password,
                        phone,
                        dni,
                        city,
                    )

                    st.success("Cuenta creada. Ya podés iniciar sesión.")
                except Exception as exc:
                    st.error(str(exc))


# =========================================================
# EMPRESA: ENCABEZADO Y CONFIGURACIÓN
# =========================================================

def get_company(company_id: int) -> dict | None:
    return fetch_one(
        "SELECT * FROM companies WHERE id = ?",
        (company_id,),
    )


def render_company_header(company: dict) -> None:
    c1, c2 = st.columns([1, 5])

    with c1:
        if company.get("logo_blob"):
            st.image(company["logo_blob"], width=120)

    with c2:
        st.title(company["name"])
        st.caption(
            f"{company.get('industry') or 'Industria no informada'} · "
            f"{company.get('country') or 'País no informado'}"
        )


def render_company_settings(user: dict, company: dict) -> None:
    st.subheader("Configuración de empresa")

    if not has_permission("manage_company"):
        st.warning("No tenés permiso para modificar la empresa.")
        return

    with st.form("company_settings_form"):
        name = st.text_input(
            "Nombre de la empresa",
            value=company["name"],
        )
        industry = st.text_input(
            "Industria",
            value=company.get("industry") or "",
        )
        country = st.text_input(
            "País",
            value=company.get("country") or "",
        )

        logo_file = st.file_uploader(
            "Reemplazar logo",
            type=["png", "jpg", "jpeg", "webp"],
        )

        submitted = st.form_submit_button(
            "Guardar cambios",
            type="primary",
        )

        if submitted:
            if not name.strip():
                st.error("Ingresá el nombre de la empresa.")
                return

            if logo_file is None:
                execute(
                    """
                    UPDATE companies
                    SET name = ?, industry = ?, country = ?
                    WHERE id = ?
                    """,
                    (
                        name.strip(),
                        industry.strip(),
                        country.strip(),
                        company["id"],
                    ),
                )
            else:
                logo_blob, logo_mime = validate_logo(logo_file)

                execute(
                    """
                    UPDATE companies
                    SET name = ?, industry = ?, country = ?,
                        logo_blob = ?, logo_mime = ?
                    WHERE id = ?
                    """,
                    (
                        name.strip(),
                        industry.strip(),
                        country.strip(),
                        logo_blob,
                        logo_mime,
                        company["id"],
                    ),
                )

            log_event(
                company["id"],
                user["id"],
                "UPDATE_COMPANY",
                "company",
                company["id"],
                {"name": name},
            )

            st.success("Empresa actualizada.")
            st.rerun()


# =========================================================
# EMPRESA: USUARIOS Y PERMISOS
# =========================================================

def count_active_admins(company_id: int) -> int:
    result = fetch_one(
        """
        SELECT COUNT(*) AS total
        FROM users
        WHERE company_id = ?
          AND role = 'ADMIN'
          AND active = 1
        """,
        (company_id,),
    )
    return int(result["total"])


def render_user_management(user: dict) -> None:
    st.subheader("Usuarios, roles y permisos")

    if not has_permission("manage_users"):
        st.warning("No tenés permiso para administrar usuarios.")
        return

    with st.expander("Crear usuario interno", expanded=True):
        with st.form("create_internal_user"):
            c1, c2 = st.columns(2)
            full_name = c1.text_input("Nombre y apellido")
            email = c2.text_input("Correo electrónico")

            c3, c4 = st.columns(2)
            role = c3.selectbox(
                "Rol base",
                ["ADMIN", "RECRUITER", "VIEWER"],
            )
            password = c4.text_input(
                "Contraseña provisoria",
                type="password",
            )

            st.markdown("**Permisos personalizados**")

            selected_permissions: list[str] = []
            permission_columns = st.columns(2)
            defaults = ROLE_DEFAULTS[role]

            for index, (key, label) in enumerate(PERMISSIONS.items()):
                checked = permission_columns[index % 2].checkbox(
                    label,
                    value=key in defaults,
                    key=f"new_permission_{key}",
                )
                if checked:
                    selected_permissions.append(key)

            submitted = st.form_submit_button(
                "Crear usuario",
                type="primary",
            )

            if submitted:
                try:
                    new_user_id = create_user(
                        email=email,
                        password=password,
                        full_name=full_name,
                        account_type="COMPANY",
                        company_id=user["company_id"],
                        role=role,
                        permissions=selected_permissions,
                    )

                    log_event(
                        user["company_id"],
                        user["id"],
                        "CREATE_USER",
                        "user",
                        new_user_id,
                        {
                            "email": normalize_email(email),
                            "role": role,
                        },
                    )

                    st.success("Usuario creado correctamente.")
                    st.rerun()
                except Exception as exc:
                    st.error(str(exc))

    users = fetch_all(
        """
        SELECT
            id, full_name, email, role,
            permissions_json, active, created_at
        FROM users
        WHERE company_id = ?
        ORDER BY created_at DESC
        """,
        (user["company_id"],),
    )

    if not users:
        st.info("No hay usuarios internos.")
        return

    display_rows = []

    for row in users:
        try:
            permissions = json.loads(
                row.get("permissions_json") or "[]"
            )
        except json.JSONDecodeError:
            permissions = []

        display_rows.append(
            {
                "ID": row["id"],
                "Nombre": row["full_name"],
                "Email": row["email"],
                "Rol": row["role"],
                "Activo": "Sí" if row["active"] else "No",
                "Permisos": ", ".join(
                    PERMISSIONS.get(item, item)
                    for item in permissions
                ),
            }
        )

    st.dataframe(
        pd.DataFrame(display_rows),
        use_container_width=True,
        hide_index=True,
    )

    st.markdown("#### Editar accesos")

    options = {
        f"{row['full_name']} — {row['email']}": row
        for row in users
    }

    selected_label = st.selectbox(
        "Usuario",
        list(options.keys()),
    )
    selected_user = options[selected_label]

    try:
        current_permissions = json.loads(
            selected_user.get("permissions_json") or "[]"
        )
    except json.JSONDecodeError:
        current_permissions = []

    with st.form("edit_internal_user"):
        role_options = ["ADMIN", "RECRUITER", "VIEWER"]
        role = st.selectbox(
            "Rol",
            role_options,
            index=role_options.index(
                selected_user.get("role") or "VIEWER"
            ),
        )

        active = st.checkbox(
            "Usuario activo",
            value=bool(selected_user["active"]),
        )

        updated_permissions: list[str] = []
        permission_columns = st.columns(2)

        for index, (key, label) in enumerate(PERMISSIONS.items()):
            checked = permission_columns[index % 2].checkbox(
                label,
                value=key in current_permissions,
                key=f"edit_permission_{selected_user['id']}_{key}",
            )
            if checked:
                updated_permissions.append(key)

        reset_password = st.text_input(
            "Nueva contraseña (opcional)",
            type="password",
            help="Dejá el campo vacío para conservar la contraseña actual.",
        )

        submitted = st.form_submit_button(
            "Guardar accesos",
            type="primary",
        )

        if submitted:
            if selected_user["id"] == user["id"] and not active:
                st.error("No podés desactivar tu propia cuenta.")
                return

            if (
                selected_user.get("role") == "ADMIN"
                and role != "ADMIN"
                and bool(selected_user["active"])
                and count_active_admins(user["company_id"]) <= 1
            ):
                st.error(
                    "La empresa debe conservar al menos un administrador activo."
                )
                return

            if (
                selected_user.get("role") == "ADMIN"
                and not active
                and count_active_admins(user["company_id"]) <= 1
            ):
                st.error(
                    "La empresa debe conservar al menos un administrador activo."
                )
                return

            execute(
                """
                UPDATE users
                SET role = ?, permissions_json = ?, active = ?
                WHERE id = ? AND company_id = ?
                """,
                (
                    role,
                    json.dumps(
                        updated_permissions,
                        ensure_ascii=False,
                    ),
                    int(active),
                    selected_user["id"],
                    user["company_id"],
                ),
            )

            if reset_password.strip():
                password_error = validate_password(reset_password)
                if password_error:
                    st.error(password_error)
                    return

                password_hash, salt = hash_password(reset_password)

                execute(
                    """
                    UPDATE users
                    SET password_hash = ?, password_salt = ?
                    WHERE id = ? AND company_id = ?
                    """,
                    (
                        password_hash,
                        salt,
                        selected_user["id"],
                        user["company_id"],
                    ),
                )

            log_event(
                user["company_id"],
                user["id"],
                "UPDATE_USER_ACCESS",
                "user",
                selected_user["id"],
                {
                    "role": role,
                    "active": active,
                    "permissions": updated_permissions,
                    "password_reset": bool(reset_password.strip()),
                },
            )

            st.success("Accesos actualizados.")
            st.rerun()



# =========================================================
# IMPORTACIÓN Y AUTOCOMPLETADO DE BÚSQUEDAS
# =========================================================

JOB_SECTION_ALIASES = {
    "title": [
        "puesto", "título del puesto", "titulo del puesto",
        "posición", "posicion", "cargo", "vacante",
    ],
    "area": ["área", "area", "departamento", "sector"],
    "seniority": ["seniority", "nivel", "jerarquía", "jerarquia"],
    "location": [
        "ubicación", "ubicacion", "localidad",
        "lugar de trabajo", "sede",
    ],
    "work_mode": ["modalidad", "modalidad de trabajo"],
    "contract_type": [
        "tipo de contratación", "tipo de contratacion",
        "contrato", "jornada",
    ],
    "description": [
        "descripción", "descripcion", "objetivo del puesto",
        "misión", "mision", "resumen",
    ],
    "responsibilities": [
        "responsabilidades", "funciones", "tareas",
        "principales tareas",
    ],
    "must_have": [
        "requisitos excluyentes", "requisitos obligatorios",
        "excluyentes", "must have",
    ],
    "desirable": [
        "requisitos deseables", "deseables",
        "se valorará", "se valorara", "nice to have",
    ],
    "competencies": [
        "competencias", "habilidades", "skills",
        "competencias requeridas",
    ],
}


def clean_import_line(line: str) -> str:
    return re.sub(r"\s+", " ", str(line)).strip(" \t-•|:")


def extract_pdf_text(raw: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError(
            "No está instalada la dependencia pypdf. "
            "Revisá requirements.txt y reiniciá la aplicación."
        ) from exc

    reader = PdfReader(io.BytesIO(raw))
    pages = [(page.extract_text() or "") for page in reader.pages]
    text = "\n".join(pages).strip()

    if not text:
        raise ValueError(
            "El PDF no contiene texto extraíble. "
            "Si es un escaneo, convertílo a PDF con texto antes de subirlo."
        )

    return text


def extract_docx_text(raw: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "No está instalada la dependencia python-docx. "
            "Revisá requirements.txt y reiniciá la aplicación."
        ) from exc

    document = Document(io.BytesIO(raw))
    blocks = [paragraph.text for paragraph in document.paragraphs]

    for table in document.tables:
        for row in table.rows:
            blocks.append(" | ".join(cell.text for cell in row.cells))

    return "\n".join(blocks).strip()


def extract_excel_text(raw: bytes) -> str:
    workbook = pd.ExcelFile(io.BytesIO(raw))
    blocks: list[str] = []

    for sheet_name in workbook.sheet_names[:10]:
        frame = pd.read_excel(
            workbook,
            sheet_name=sheet_name,
            header=None,
        ).dropna(how="all")

        blocks.append(f"HOJA: {sheet_name}")

        for row in frame.values.tolist():
            values = [
                clean_import_line(value)
                for value in row
                if pd.notna(value) and clean_import_line(value)
            ]
            if values:
                blocks.append(" | ".join(values))

    return "\n".join(blocks).strip()


def extract_job_file_text(uploaded_file) -> str:
    extension = Path(uploaded_file.name).suffix.lower()
    raw = uploaded_file.getvalue()

    if extension == ".pdf":
        return extract_pdf_text(raw)
    if extension == ".docx":
        return extract_docx_text(raw)
    if extension in {".xlsx", ".xls"}:
        return extract_excel_text(raw)

    raise ValueError("Usá un archivo PDF, DOCX, XLSX o XLS.")


def parse_job_text(text: str) -> dict[str, str]:
    result = {field: "" for field in JOB_SECTION_ALIASES}
    section_values = {field: [] for field in JOB_SECTION_ALIASES}
    current_section: str | None = None

    lines = [
        clean_import_line(line)
        for line in text.splitlines()
        if clean_import_line(line)
    ]

    for line in lines:
        detected_field = None
        inline_value = ""

        for field, aliases in JOB_SECTION_ALIASES.items():
            for alias in aliases:
                match = re.match(
                    rf"^{re.escape(alias)}\s*[:\-]\s*(.*)$",
                    line,
                    flags=re.IGNORECASE,
                )
                if match:
                    detected_field = field
                    inline_value = match.group(1).strip()
                    break

                if line.lower() == alias.lower():
                    detected_field = field
                    break

            if detected_field:
                break

        if detected_field:
            current_section = detected_field
            if inline_value:
                section_values[detected_field].append(inline_value)
            continue

        if current_section:
            section_values[current_section].append(line)

    for field, values in section_values.items():
        result[field] = "\n".join(values).strip()

    full_lower = text.lower()

    if not result["title"] and lines:
        result["title"] = next(
            (
                line for line in lines[:15]
                if 2 <= len(line.split()) <= 12 and len(line) <= 120
            ),
            "",
        )

    if not result["seniority"]:
        seniority_map = [
            ("Pasantía", ["pasantía", "pasantia", "internship"]),
            ("Junior", [" junior", "jr.", " jr "]),
            ("Semi Senior", ["semi senior", "semisenior", "ssr"]),
            ("Senior", [" senior", "sr.", " sr "]),
            ("Liderazgo", ["supervisor", "jefatura", "líder", "lider"]),
            ("Dirección", ["dirección", "direccion", "director", "gerencia"]),
        ]
        padded = f" {full_lower} "
        for label, keywords in seniority_map:
            if any(keyword in padded for keyword in keywords):
                result["seniority"] = label
                break

    if not result["work_mode"]:
        for label, keywords in {
            "Presencial": ["presencial"],
            "Híbrido": ["híbrido", "hibrido"],
            "Remoto": ["remoto", "home office", "teletrabajo"],
        }.items():
            if any(keyword in full_lower for keyword in keywords):
                result["work_mode"] = label
                break

    if not result["description"]:
        result["description"] = text[:3000].strip()

    return result


def save_new_job(
    user: dict,
    values: dict,
    source_filename: str = "",
) -> int:
    return execute(
        """
        INSERT INTO jobs(
            company_id, title, area, seniority,
            location, work_mode, contract_type,
            description, responsibilities, must_have,
            desirable, competencies, status,
            source_filename, created_by, created_at, updated_at
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            user["company_id"],
            values["title"].strip(),
            values["area"].strip(),
            values["seniority"],
            values["location"].strip(),
            values["work_mode"],
            values["contract_type"].strip(),
            values["description"].strip(),
            values["responsibilities"].strip(),
            values["must_have"].strip(),
            values["desirable"].strip(),
            values["competencies"].strip(),
            values.get("status", "ABIERTA"),
            source_filename,
            user["id"],
            now_iso(),
            now_iso(),
        ),
    )



def update_job_record(
    user: dict,
    job_id: int,
    values: dict,
) -> None:
    """
    Confirma primero los cambios de la búsqueda.
    La auditoría se intenta después y no puede revertirlos.
    """
    conn = get_connection()

    try:
        cursor = conn.execute(
            """
            UPDATE jobs
            SET title = ?,
                area = ?,
                seniority = ?,
                location = ?,
                work_mode = ?,
                contract_type = ?,
                description = ?,
                responsibilities = ?,
                must_have = ?,
                desirable = ?,
                competencies = ?,
                status = ?,
                updated_at = ?
            WHERE id = ?
              AND company_id = ?
            """,
            (
                values["title"].strip(),
                values["area"].strip(),
                values["seniority"],
                values["location"].strip(),
                values["work_mode"],
                values["contract_type"].strip(),
                values["description"].strip(),
                values["responsibilities"].strip(),
                values["must_have"].strip(),
                values["desirable"].strip(),
                values["competencies"].strip(),
                values["status"],
                now_iso(),
                job_id,
                user["company_id"],
            ),
        )

        if cursor.rowcount != 1:
            conn.rollback()
            raise ValueError(
                "La búsqueda no existe o no pertenece "
                "a la empresa del usuario."
            )

        conn.commit()

    except Exception:
        conn.rollback()
        raise

    finally:
        conn.close()

    log_event(
        user["company_id"],
        user["id"],
        "UPDATE_JOB",
        "job",
        job_id,
        {
            "title": values["title"],
            "status": values["status"],
        },
    )


def render_job_fields(
    prefix: str,
    initial: dict | None = None,
    include_status: bool = False,
) -> dict:
    initial = initial or {}

    seniority_options = [
        "Pasantía", "Junior", "Semi Senior",
        "Senior", "Liderazgo", "Dirección",
        "No especificado",
    ]
    work_mode_options = [
        "Presencial", "Híbrido", "Remoto", "No especificado"
    ]

    initial_seniority = initial.get("seniority") or "No especificado"
    if initial_seniority not in seniority_options:
        initial_seniority = "No especificado"

    initial_mode = initial.get("work_mode") or "No especificado"
    if initial_mode not in work_mode_options:
        initial_mode = "No especificado"

    title = st.text_input(
        "Puesto",
        value=initial.get("title") or "",
        key=f"{prefix}_title",
    )

    c1, c2 = st.columns(2)
    area = c1.text_input(
        "Área",
        value=initial.get("area") or "",
        key=f"{prefix}_area",
    )
    seniority = c2.selectbox(
        "Seniority",
        seniority_options,
        index=seniority_options.index(initial_seniority),
        key=f"{prefix}_seniority",
    )

    c3, c4 = st.columns(2)
    location = c3.text_input(
        "Ubicación",
        value=initial.get("location") or "",
        key=f"{prefix}_location",
    )
    work_mode = c4.selectbox(
        "Modalidad",
        work_mode_options,
        index=work_mode_options.index(initial_mode),
        key=f"{prefix}_work_mode",
    )

    contract_type = st.text_input(
        "Tipo de contratación / jornada",
        value=initial.get("contract_type") or "",
        key=f"{prefix}_contract_type",
    )

    description = st.text_area(
        "Descripción del puesto",
        value=initial.get("description") or "",
        height=160,
        key=f"{prefix}_description",
    )
    responsibilities = st.text_area(
        "Responsabilidades y tareas",
        value=initial.get("responsibilities") or "",
        height=130,
        key=f"{prefix}_responsibilities",
    )
    must_have = st.text_area(
        "Requisitos excluyentes",
        value=initial.get("must_have") or "",
        height=120,
        key=f"{prefix}_must_have",
    )
    desirable = st.text_area(
        "Requisitos deseables",
        value=initial.get("desirable") or "",
        height=120,
        key=f"{prefix}_desirable",
    )
    competencies = st.text_area(
        "Competencias",
        value=initial.get("competencies") or "",
        height=110,
        key=f"{prefix}_competencies",
    )

    status = initial.get("status") or "ABIERTA"
    if include_status:
        status_options = ["ABIERTA", "PAUSADA", "CERRADA"]
        if status not in status_options:
            status = "ABIERTA"
        status = st.selectbox(
            "Estado",
            status_options,
            index=status_options.index(status),
            key=f"{prefix}_status",
        )

    return {
        "title": title,
        "area": area,
        "seniority": seniority,
        "location": location,
        "work_mode": work_mode,
        "contract_type": contract_type,
        "description": description,
        "responsibilities": responsibilities,
        "must_have": must_have,
        "desirable": desirable,
        "competencies": competencies,
        "status": status,
    }



# =========================================================
# EMPRESA: BÚSQUEDAS Y POSTULACIONES
# =========================================================

def render_jobs(user: dict) -> None:
    st.subheader("Búsquedas laborales")

    create_tab, manage_tab = st.tabs(
        ["Crear o importar", "Administrar búsquedas"]
    )

    with create_tab:
        if not has_permission("manage_jobs"):
            st.warning("No tenés permiso para crear búsquedas.")
        else:
            mode = st.radio(
                "Método de carga",
                ["Importar archivo", "Crear manualmente"],
                horizontal=True,
                key="job_creation_mode",
            )

            if mode == "Importar archivo":
                uploaded_file = st.file_uploader(
                    "Subí el requerimiento en Excel, PDF o Word",
                    type=["xlsx", "xls", "pdf", "docx"],
                    key="job_import_file",
                )

                st.caption(
                    "El sistema extrae texto del archivo y propone valores. "
                    "Siempre podés revisar y modificar los campos antes de guardar."
                )

                if uploaded_file is not None and st.button(
                    "Analizar y autocompletar",
                    type="primary",
                    key="analyze_job_file",
                ):
                    try:
                        extracted_text = extract_job_file_text(uploaded_file)
                        parsed = parse_job_text(extracted_text)
                        st.session_state["job_import_parsed"] = parsed
                        st.session_state["job_import_filename"] = uploaded_file.name
                        st.session_state["job_import_preview"] = extracted_text[:5000]
                        st.success(
                            "Archivo analizado. Revisá y editá la información."
                        )
                        st.rerun()
                    except Exception as exc:
                        st.error(f"No se pudo analizar el archivo: {exc}")

                parsed = st.session_state.get("job_import_parsed")
                if parsed:
                    with st.expander("Texto detectado en el archivo"):
                        st.text(
                            st.session_state.get("job_import_preview", "")
                        )

                    with st.form("imported_job_form"):
                        values = render_job_fields(
                            "imported_job",
                            initial=parsed,
                        )
                        submitted = st.form_submit_button(
                            "Guardar búsqueda",
                            type="primary",
                            use_container_width=True,
                        )

                        if submitted:
                            if not values["title"].strip():
                                st.error("Ingresá el nombre del puesto.")
                            else:
                                try:
                                    job_id = save_new_job(
                                        user,
                                        values,
                                        st.session_state.get(
                                            "job_import_filename",
                                            "",
                                        ),
                                    )
                                    log_event(
                                        user["company_id"],
                                        user["id"],
                                        "CREATE_JOB_FROM_FILE",
                                        "job",
                                        job_id,
                                        {
                                            "title": values["title"],
                                            "source_filename": st.session_state.get(
                                                "job_import_filename",
                                                "",
                                            ),
                                        },
                                    )
                                    for key in [
                                        "job_import_parsed",
                                        "job_import_filename",
                                        "job_import_preview",
                                    ]:
                                        st.session_state.pop(key, None)
                                    st.success(
                                        "Búsqueda importada y guardada."
                                    )
                                    st.rerun()
                                except Exception as exc:
                                    st.error(
                                        f"No se pudo guardar la búsqueda: {exc}"
                                    )

            else:
                with st.form("manual_job_form"):
                    values = render_job_fields("manual_job")
                    submitted = st.form_submit_button(
                        "Guardar búsqueda",
                        type="primary",
                        use_container_width=True,
                    )

                    if submitted:
                        if not values["title"].strip():
                            st.error("Ingresá el nombre del puesto.")
                        else:
                            try:
                                job_id = save_new_job(user, values)
                                log_event(
                                    user["company_id"],
                                    user["id"],
                                    "CREATE_JOB",
                                    "job",
                                    job_id,
                                    {"title": values["title"]},
                                )
                                st.success(
                                    "Búsqueda creada correctamente."
                                )
                                st.rerun()
                            except Exception as exc:
                                st.error(
                                    f"No se pudo guardar la búsqueda: {exc}"
                                )

    with manage_tab:
        jobs = fetch_all(
            """
            SELECT *
            FROM jobs
            WHERE company_id = ?
            ORDER BY created_at DESC
            """,
            (user["company_id"],),
        )

        if not jobs:
            st.info("Todavía no hay búsquedas.")
            return

        summary_rows = [
            {
                "ID": job["id"],
                "Puesto": job["title"],
                "Área": job.get("area") or "",
                "Seniority": job.get("seniority") or "",
                "Modalidad": job.get("work_mode") or "",
                "Estado": job.get("status") or "ABIERTA",
                "Archivo de origen": job.get("source_filename") or "",
                "Creada": job.get("created_at") or "",
            }
            for job in jobs
        ]

        st.dataframe(
            pd.DataFrame(summary_rows),
            use_container_width=True,
            hide_index=True,
        )

        if not has_permission("manage_jobs"):
            st.info(
                "Tenés acceso de consulta, pero no permiso para editar."
            )
            return

        job_options = {
            f"#{job['id']} — {job['title']}": job
            for job in jobs
        }
        selected_label = st.selectbox(
            "Seleccionar búsqueda para editar",
            list(job_options.keys()),
            key="selected_job_edit",
        )
        selected_job = job_options[selected_label]

        with st.form(f"edit_job_form_{selected_job['id']}"):
            edited_values = render_job_fields(
                f"edit_job_{selected_job['id']}",
                initial=selected_job,
                include_status=True,
            )

            submitted = st.form_submit_button(
                "Guardar cambios",
                type="primary",
                use_container_width=True,
            )

            if submitted:
                if not edited_values["title"].strip():
                    st.error("Ingresá el nombre del puesto.")
                else:
                    try:
                        update_job_record(
                            user=user,
                            job_id=selected_job["id"],
                            values=edited_values,
                        )

                        st.session_state[
                            "job_update_success"
                        ] = (
                            "Búsqueda actualizada correctamente."
                        )

                    except sqlite3.Error as exc:
                        st.error(
                            "No se pudo actualizar la búsqueda "
                            "por un error de base de datos."
                        )
                        st.code(str(exc))

                    except Exception as exc:
                        st.error(
                            "No se pudo actualizar la búsqueda."
                        )
                        st.exception(exc)

        if st.session_state.pop(
            "job_update_success",
            None,
        ):
            st.success(
                "Búsqueda actualizada correctamente. "
                "Los cambios quedaron confirmados en la base."
            )
            st.caption(
                f"Corrección activa: versión {APP_VERSION}"
            )




# =========================================================
# CANDIDATOS, CV Y PRESELECCIÓN
# =========================================================

CV_SECTION_ALIASES = {
    "experience_summary": [
        "experiencia",
        "experiencia laboral",
        "experiencia profesional",
        "antecedentes laborales",
        "employment",
        "work experience",
    ],
    "education_summary": [
        "educación",
        "educacion",
        "formación",
        "formacion",
        "estudios",
        "academic background",
        "education",
    ],
    "skills_text": [
        "habilidades",
        "competencias",
        "herramientas",
        "tecnologías",
        "tecnologias",
        "skills",
        "technical skills",
    ],
    "languages_text": [
        "idiomas",
        "languages",
    ],
}

SCORING_STOPWORDS = {
    "para", "como", "con", "los", "las", "una", "uno", "del", "por",
    "que", "se", "de", "en", "y", "o", "al", "el", "la", "un",
    "requisito", "requisitos", "experiencia", "conocimiento",
    "conocimientos", "manejo", "nivel", "años", "anos", "deseable",
    "excluyente", "excluyentes", "competencia", "competencias",
}


def normalize_match_text(value: str) -> str:
    normalized = unicodedata.normalize(
        "NFKD",
        value or "",
    )
    without_accents = "".join(
        character
        for character in normalized
        if not unicodedata.combining(character)
    )
    without_accents = without_accents.lower()
    without_accents = re.sub(
        r"[^a-z0-9+#.\s-]",
        " ",
        without_accents,
    )
    return re.sub(r"\s+", " ", without_accents).strip()


def validate_cv_file(uploaded_file) -> tuple[bytes, str]:
    if uploaded_file is None:
        raise ValueError("Seleccioná un CV.")

    raw = uploaded_file.getvalue()
    if not raw:
        raise ValueError("El CV está vacío.")

    if len(raw) > MAX_CV_SIZE_BYTES:
        raise ValueError(
            "El CV supera el máximo permitido de 5 MB."
        )

    extension = Path(uploaded_file.name).suffix.lower()
    if extension not in {".pdf", ".docx"}:
        raise ValueError("El CV debe estar en PDF o DOCX.")

    mime_type = (
        uploaded_file.type
        or (
            "application/pdf"
            if extension == ".pdf"
            else "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        )
    )
    return raw, mime_type


def extract_cv_text(uploaded_file) -> tuple[bytes, str, str]:
    raw, mime_type = validate_cv_file(uploaded_file)
    extension = Path(uploaded_file.name).suffix.lower()

    if extension == ".pdf":
        text = extract_pdf_text(raw)
    else:
        text = extract_docx_text(raw)

    if len(text.strip()) < 30:
        raise ValueError(
            "No se pudo extraer suficiente texto del CV."
        )

    return raw, mime_type, text.strip()


def extract_named_cv_sections(text: str) -> dict[str, str]:
    results = {
        field: ""
        for field in CV_SECTION_ALIASES
    }
    collected = {
        field: []
        for field in CV_SECTION_ALIASES
    }
    current_field = None

    lines = [
        clean_import_line(line)
        for line in text.splitlines()
        if clean_import_line(line)
    ]

    for line in lines:
        normalized_line = normalize_match_text(line)
        detected_field = None
        inline_value = ""

        for field, aliases in CV_SECTION_ALIASES.items():
            for alias in aliases:
                normalized_alias = normalize_match_text(alias)

                if normalized_line == normalized_alias:
                    detected_field = field
                    break

                prefix_pattern = (
                    rf"^{re.escape(normalized_alias)}\s*[:\-]\s*(.*)$"
                )
                match = re.match(
                    prefix_pattern,
                    normalized_line,
                )
                if match:
                    detected_field = field
                    if ":" in line:
                        inline_value = line.split(":", 1)[1].strip()
                    elif "-" in line:
                        inline_value = line.split("-", 1)[1].strip()
                    break

            if detected_field:
                break

        if detected_field:
            current_field = detected_field
            if inline_value:
                collected[detected_field].append(inline_value)
            continue

        if current_field:
            collected[current_field].append(line)

    for field, values in collected.items():
        results[field] = "\n".join(values[:30]).strip()

    return results


def extract_candidate_name(
    text: str,
    fallback_filename: str,
) -> str:
    lines = [
        clean_import_line(line)
        for line in text.splitlines()
        if clean_import_line(line)
    ]

    for line in lines[:12]:
        if "@" in line:
            continue
        if re.search(r"\d{6,}", line):
            continue
        if normalize_match_text(line) in {
            "curriculum vitae",
            "curriculum",
            "cv",
            "resume",
        }:
            continue

        words = line.split()
        if 2 <= len(words) <= 6 and len(line) <= 80:
            return line.title()

    fallback = Path(fallback_filename).stem
    fallback = re.sub(r"[_\-]+", " ", fallback)
    fallback = re.sub(
        r"\b(cv|curriculum|vitae|resume)\b",
        " ",
        fallback,
        flags=re.I,
    )
    fallback = re.sub(r"\s+", " ", fallback).strip()
    return fallback.title() or "Candidato sin identificar"


def parse_cv_text(
    text: str,
    fallback_filename: str,
) -> dict[str, str]:
    email_match = re.search(
        r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}",
        text,
    )
    phone_match = re.search(
        r"(?:\+?\d[\d\s().-]{7,}\d)",
        text,
    )

    sections = extract_named_cv_sections(text)

    lines = [
        clean_import_line(line)
        for line in text.splitlines()
        if clean_import_line(line)
    ]

    headline_lines = []
    for line in lines[:18]:
        if "@" in line:
            continue
        if phone_match and phone_match.group(0) in line:
            continue
        if line.lower() == extract_candidate_name(
            text,
            fallback_filename,
        ).lower():
            continue
        headline_lines.append(line)
        if len(headline_lines) >= 3:
            break

    return {
        "full_name": extract_candidate_name(
            text,
            fallback_filename,
        ),
        "email": (
            normalize_email(email_match.group(0))
            if email_match
            else ""
        ),
        "phone": (
            re.sub(r"\s+", " ", phone_match.group(0)).strip()
            if phone_match
            else ""
        ),
        "headline": " · ".join(headline_lines)[:500],
        "education_summary": sections["education_summary"][:5000],
        "experience_summary": sections["experience_summary"][:7000],
        "skills_text": sections["skills_text"][:4000],
        "languages_text": sections["languages_text"][:2000],
    }


def extract_scoring_terms(value: str) -> list[str]:
    if not value:
        return []

    segments = re.split(
        r"[\n;•|]+",
        value,
    )
    terms: list[str] = []

    for segment in segments:
        for part in segment.split(","):
            cleaned = normalize_match_text(part)
            cleaned = re.sub(
                r"^(requisitos?|excluyentes?|deseables?|"
                r"competencias?|habilidades?)\s*[:\-]?\s*",
                "",
                cleaned,
            ).strip()

            if not cleaned:
                continue

            words = [
                word
                for word in cleaned.split()
                if len(word) >= 3
                and word not in SCORING_STOPWORDS
            ]

            if not words:
                continue

            if len(words) <= 7:
                terms.append(" ".join(words))
            else:
                terms.extend(words)

    unique_terms = []
    seen = set()

    for term in terms:
        if term not in seen:
            unique_terms.append(term)
            seen.add(term)

    return unique_terms[:50]




def normalized_location_value(value: str | None) -> str:
    return normalize_match_text(value or "")


def build_location_label(candidate: dict) -> str:
    parts = [
        candidate.get("country"),
        candidate.get("province"),
        candidate.get("city"),
        candidate.get("locality"),
    ]
    clean_parts = [
        str(part).strip()
        for part in parts
        if part and str(part).strip()
    ]
    return " / ".join(clean_parts) if clean_parts else "Sin ubicación"


def recommendation_explanation(report: dict) -> dict:
    """
    Convierte el informe técnico en una explicación de negocio
    más clara y exhaustiva para RR. HH.
    """
    total = report.get("total")
    recommendation = report.get("recommendation") or "REVISIÓN MANUAL"
    config = report.get("config") or {}
    thresholds = config.get("thresholds") or {}

    advance_threshold = float(thresholds.get("advance", 75))
    review_threshold = float(thresholds.get("review", 50))
    minimum_required = float(
        thresholds.get("minimum_required_coverage", 50)
    )

    dimensions = report.get("dimensions") or []
    active_dimensions = [
        item
        for item in dimensions
        if item.get("criteria_count")
    ]

    mandatory_breaches = []
    fulfilled_mandatory = []
    strongest_dimensions = []
    weakest_dimensions = []

    for dimension in active_dimensions:
        coverage = round(float(dimension.get("coverage") or 0) * 100, 1)
        item = {
            "label": dimension.get("label") or "Criterio",
            "coverage": coverage,
            "matched": dimension.get("matched") or [],
            "missing": dimension.get("missing") or [],
            "required": bool(dimension.get("required")),
            "points": dimension.get("points"),
            "weight": dimension.get("weight"),
        }

        if item["required"]:
            if coverage < minimum_required:
                mandatory_breaches.append(item)
            else:
                fulfilled_mandatory.append(item)

        if coverage >= 70:
            strongest_dimensions.append(item)
        elif coverage < 50:
            weakest_dimensions.append(item)

    if total is None:
        headline = (
            "No se pudo emitir una recomendación automática porque faltan "
            "criterios configurados o no existe un CV analizable."
        )
    elif recommendation == "AVANZA":
        headline = (
            f"El candidato alcanza {float(total):.1f}% de coincidencia, "
            f"supera el umbral de avance de {advance_threshold:.1f}% "
            "y no presenta incumplimientos críticos en los criterios obligatorios."
        )
    elif recommendation == "REVISAR":
        headline = (
            f"El candidato alcanza {float(total):.1f}% de coincidencia. "
            f"Supera el umbral de revisión de {review_threshold:.1f}%, "
            f"pero no llega al umbral de avance de {advance_threshold:.1f}%."
        )
    elif mandatory_breaches:
        breached_labels = ", ".join(
            item["label"]
            for item in mandatory_breaches
        )
        headline = (
            f"El candidato obtiene {float(total):.1f}% de coincidencia, "
            "pero queda clasificado como baja coincidencia porque no alcanza "
            f"la cobertura mínima obligatoria de {minimum_required:.1f}% "
            f"en: {breached_labels}."
        )
    else:
        headline = (
            f"El candidato obtiene {float(total):.1f}% de coincidencia, "
            f"por debajo del umbral de revisión de {review_threshold:.1f}%."
        )

    decision_factors = []

    for item in strongest_dimensions:
        matched = ", ".join(item["matched"][:10]) or "sin coincidencias listadas"
        decision_factors.append(
            f"Fortaleza en {item['label']}: {item['coverage']:.1f}% "
            f"de cobertura. Evidencias: {matched}."
        )

    for item in weakest_dimensions:
        missing = ", ".join(item["missing"][:10]) or "sin brechas listadas"
        decision_factors.append(
            f"Brecha en {item['label']}: {item['coverage']:.1f}% "
            f"de cobertura. No se encontró: {missing}."
        )

    if mandatory_breaches:
        for item in mandatory_breaches:
            decision_factors.append(
                f"Criterio obligatorio incumplido: {item['label']} "
                f"({item['coverage']:.1f}% de cobertura; mínimo requerido "
                f"{minimum_required:.1f}%)."
            )

    next_steps = []
    if recommendation == "AVANZA":
        next_steps = [
            "Validar en entrevista las competencias que no aparecen claramente en el CV.",
            "Confirmar disponibilidad, condiciones y motivación.",
            "Verificar referencias y documentación antes de la decisión final.",
        ]
    elif recommendation == "REVISAR":
        next_steps = [
            "Realizar una entrevista de preselección enfocada en las brechas detectadas.",
            "Solicitar ejemplos concretos de experiencia en los requisitos parcialmente cubiertos.",
            "Comparar el perfil con candidatos de puntaje similar antes de decidir.",
        ]
    else:
        next_steps = [
            "Revisar manualmente si el CV usa sinónimos o información no detectada.",
            "Confirmar si los criterios obligatorios fueron correctamente configurados.",
            "Mantener en el banco para búsquedas con mejor correspondencia, salvo descarte humano justificado.",
        ]

    return {
        "headline": headline,
        "decision_factors": decision_factors,
        "mandatory_breaches": mandatory_breaches,
        "fulfilled_mandatory": fulfilled_mandatory,
        "strongest_dimensions": strongest_dimensions,
        "weakest_dimensions": weakest_dimensions,
        "next_steps": next_steps,
        "thresholds": {
            "advance": advance_threshold,
            "review": review_threshold,
            "minimum_required": minimum_required,
        },
    }


def render_detailed_recommendation(report: dict) -> None:
    explanation = recommendation_explanation(report)

    st.markdown("### Conclusión de ALBA")
    st.info(explanation["headline"])

    total = report.get("total")
    thresholds = explanation["thresholds"]

    c1, c2, c3 = st.columns(3)
    c1.metric(
        "Puntaje obtenido",
        "Pendiente" if total is None else f"{float(total):.1f}%",
    )
    c2.metric(
        "Umbral para revisar",
        f"{thresholds['review']:.1f}%",
    )
    c3.metric(
        "Umbral para avanzar",
        f"{thresholds['advance']:.1f}%",
    )

    st.markdown("#### Factores determinantes")
    if explanation["decision_factors"]:
        for factor in explanation["decision_factors"]:
            st.write(f"• {factor}")
    else:
        st.write("No hay factores suficientes para emitir una explicación detallada.")

    if explanation["mandatory_breaches"]:
        st.error(
            "Existen requisitos obligatorios por debajo de la cobertura mínima."
        )
        for breach in explanation["mandatory_breaches"]:
            st.write(
                f"• **{breach['label']}**: "
                f"{breach['coverage']:.1f}% de cobertura."
            )

    if explanation["fulfilled_mandatory"]:
        with st.expander("Requisitos obligatorios cumplidos"):
            for item in explanation["fulfilled_mandatory"]:
                st.write(
                    f"• **{item['label']}**: "
                    f"{item['coverage']:.1f}% de cobertura."
                )

    st.markdown("#### Próximos pasos sugeridos")
    for item in explanation["next_steps"]:
        st.write(f"• {item}")

    st.caption(
        "Esta explicación se basa en el contenido profesional del CV y "
        "en los criterios configurados para la búsqueda. No utiliza edad, "
        "nombre, correo, teléfono ni ubicación para calcular la recomendación."
    )


def generate_candidate_job_suggestions(
    company_id: int,
    job_id: int | None = None,
) -> dict:
    jobs_query = """
        SELECT *
        FROM jobs
        WHERE company_id = ?
          AND status = 'ABIERTA'
    """
    job_params: tuple = (company_id,)

    if job_id is not None:
        jobs_query += " AND id = ?"
        job_params = (company_id, job_id)

    jobs = fetch_all(jobs_query, job_params)

    candidates = fetch_all(
        """
        SELECT
            candidates.*,
            users.full_name,
            users.email
        FROM company_candidates
        JOIN candidates
            ON candidates.id = company_candidates.candidate_id
        JOIN users
            ON users.id = candidates.user_id
        WHERE company_candidates.company_id = ?
        ORDER BY company_candidates.created_at DESC
        """,
        (company_id,),
    )

    generated = 0
    skipped_without_cv = 0

    for job in jobs:
        for candidate in candidates:
            existing_application = fetch_one(
                """
                SELECT id
                FROM applications
                WHERE candidate_id = ?
                  AND job_id = ?
                """,
                (candidate["id"], job["id"]),
            )
            if existing_application:
                continue

            latest_cv = get_latest_cv(candidate["id"])
            if not latest_cv or not latest_cv.get("extracted_text"):
                skipped_without_cv += 1
                continue

            report = calculate_match_score(
                job,
                latest_cv["extracted_text"],
            )
            score = report.get("total")

            status = (
                "ACTIVA"
                if score is not None
                and float(score) >= SUGGESTION_MIN_SCORE
                else "BAJA"
            )

            existing = fetch_one(
                """
                SELECT id
                FROM candidate_job_suggestions
                WHERE company_id = ?
                  AND candidate_id = ?
                  AND job_id = ?
                """,
                (company_id, candidate["id"], job["id"]),
            )

            values = (
                score,
                report.get("recommendation"),
                json.dumps(report, ensure_ascii=False),
                status,
                now_iso(),
                company_id,
                candidate["id"],
                job["id"],
            )

            if existing:
                execute(
                    """
                    UPDATE candidate_job_suggestions
                    SET score_total = ?,
                        recommendation = ?,
                        report_json = ?,
                        status = ?,
                        generated_at = ?
                    WHERE company_id = ?
                      AND candidate_id = ?
                      AND job_id = ?
                    """,
                    values,
                )
            else:
                execute(
                    """
                    INSERT INTO candidate_job_suggestions(
                        score_total,
                        recommendation,
                        report_json,
                        status,
                        generated_at,
                        company_id,
                        candidate_id,
                        job_id
                    )
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    values,
                )

            generated += 1

    return {
        "generated": generated,
        "jobs": len(jobs),
        "candidates": len(candidates),
        "skipped_without_cv": skipped_without_cv,
    }


def get_job_suggestions(
    company_id: int,
    job_id: int,
) -> list[dict]:
    return fetch_all(
        """
        SELECT
            candidate_job_suggestions.*,
            candidates.country,
            candidates.province,
            candidates.city,
            candidates.locality,
            candidates.headline,
            candidates.tags_text,
            users.full_name,
            users.email
        FROM candidate_job_suggestions
        JOIN candidates
            ON candidates.id = candidate_job_suggestions.candidate_id
        JOIN users
            ON users.id = candidates.user_id
        WHERE candidate_job_suggestions.company_id = ?
          AND candidate_job_suggestions.job_id = ?
          AND candidate_job_suggestions.status = 'ACTIVA'
        ORDER BY
            candidate_job_suggestions.score_total DESC,
            candidate_job_suggestions.generated_at DESC
        LIMIT ?
        """,
        (
            company_id,
            job_id,
            SUGGESTION_TOP_LIMIT,
        ),
    )


def calculate_age(birth_date_value: str | None) -> int | None:
    if not birth_date_value:
        return None

    try:
        born = date.fromisoformat(str(birth_date_value)[:10])
    except (TypeError, ValueError):
        return None

    today = date.today()
    return (
        today.year
        - born.year
        - ((today.month, today.day) < (born.month, born.day))
    )


def detect_analysis_preset(job: dict) -> str:
    text = normalize_match_text(
        " ".join(
            [
                job.get("title") or "",
                job.get("area") or "",
                job.get("description") or "",
            ]
        )
    )

    rules = [
        (
            "Tecnología / Sistemas",
            ["sistemas", "software", "it", "datos", "programador", "developer"],
        ),
        (
            "Ingeniería / Producción",
            ["ingenier", "produccion", "planta", "calidad", "procesos"],
        ),
        (
            "Compras / Logística",
            ["compras", "logistica", "abastecimiento", "supply"],
        ),
        (
            "Recursos Humanos",
            ["recursos humanos", "rrhh", "seleccion", "talento"],
        ),
        (
            "Comercial / Ventas",
            ["comercial", "ventas", "vendedor", "cuentas"],
        ),
        (
            "Administración / Finanzas",
            ["administracion", "finanzas", "contabilidad", "control de gestion"],
        ),
    ]

    for preset_name, keywords in rules:
        if any(keyword in text for keyword in keywords):
            return preset_name

    return "General"


def default_scoring_config(
    job: dict,
    preset_name: str | None = None,
) -> dict:
    resolved_preset = (
        detect_analysis_preset(job)
        if not preset_name or preset_name == "Automático"
        else preset_name
    )

    groups = [
        {
            "key": "must_have",
            "label": "Requisitos excluyentes",
            "terms": extract_scoring_terms(job.get("must_have") or ""),
            "weight": 50.0,
            "required": True,
        },
        {
            "key": "desirable",
            "label": "Requisitos deseables",
            "terms": extract_scoring_terms(job.get("desirable") or ""),
            "weight": 20.0,
            "required": False,
        },
        {
            "key": "competencies",
            "label": "Competencias",
            "terms": extract_scoring_terms(job.get("competencies") or ""),
            "weight": 15.0,
            "required": False,
        },
        {
            "key": "context",
            "label": "Contexto del puesto",
            "terms": extract_scoring_terms(
                " | ".join(
                    [
                        job.get("title") or "",
                        job.get("area") or "",
                        job.get("seniority") or "",
                    ]
                )
            ),
            "weight": 10.0,
            "required": False,
        },
        {
            "key": "preset",
            "label": f"Criterios sugeridos · {resolved_preset}",
            "terms": list(ANALYSIS_PRESETS.get(resolved_preset, [])),
            "weight": 5.0,
            "required": False,
        },
        {
            "key": "custom",
            "label": "Criterios personalizados",
            "terms": [],
            "weight": 0.0,
            "required": False,
        },
    ]

    return {
        "preset_name": resolved_preset,
        "groups": groups,
        "thresholds": {
            "advance": 75.0,
            "review": 50.0,
            "minimum_required_coverage": 50.0,
        },
    }


def get_job_scoring_config(job: dict) -> dict:
    saved = fetch_one(
        """
        SELECT preset_name, criteria_json, thresholds_json
        FROM job_scoring_config
        WHERE job_id = ?
        """,
        (job["id"],),
    )

    if not saved:
        return default_scoring_config(job)

    try:
        groups = json.loads(saved.get("criteria_json") or "[]")
        thresholds = json.loads(saved.get("thresholds_json") or "{}")
    except json.JSONDecodeError:
        return default_scoring_config(job)

    if not isinstance(groups, list) or not groups:
        return default_scoring_config(job)

    return {
        "preset_name": saved.get("preset_name") or "General",
        "groups": groups,
        "thresholds": thresholds,
    }


def save_job_scoring_config(
    user: dict,
    job: dict,
    preset_name: str,
    groups: list[dict],
    thresholds: dict,
) -> None:
    if not groups:
        raise ValueError("Cargá al menos un grupo de criterios.")

    active_weight = sum(
        float(group.get("weight") or 0)
        for group in groups
        if group.get("terms")
    )
    if active_weight <= 0:
        raise ValueError(
            "La suma de pesos de los criterios activos debe ser mayor a cero."
        )

    existing = fetch_one(
        "SELECT id FROM job_scoring_config WHERE job_id = ?",
        (job["id"],),
    )

    params = (
        preset_name,
        json.dumps(groups, ensure_ascii=False),
        json.dumps(thresholds, ensure_ascii=False),
        user["id"],
        now_iso(),
        job["id"],
    )

    if existing:
        execute(
            """
            UPDATE job_scoring_config
            SET preset_name = ?,
                criteria_json = ?,
                thresholds_json = ?,
                updated_by = ?,
                updated_at = ?
            WHERE job_id = ?
            """,
            params,
        )
    else:
        execute(
            """
            INSERT INTO job_scoring_config(
                preset_name,
                criteria_json,
                thresholds_json,
                updated_by,
                updated_at,
                job_id
            )
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            params,
        )

    log_event(
        user["company_id"],
        user["id"],
        "UPDATE_SCORING_CRITERIA",
        "job",
        job["id"],
        {
            "preset": preset_name,
            "groups": len(groups),
            "thresholds": thresholds,
        },
    )


def ensure_company_candidate(
    company_id: int,
    candidate_id: int,
    source: str,
    added_by: int | None,
) -> None:
    execute(
        """
        INSERT OR IGNORE INTO company_candidates(
            company_id,
            candidate_id,
            source,
            added_by,
            created_at
        )
        VALUES (?, ?, ?, ?, ?)
        """,
        (
            company_id,
            candidate_id,
            source,
            added_by,
            now_iso(),
        ),
    )


def assign_candidate_to_job(
    user: dict,
    candidate_id: int,
    job_id: int,
    initial_status: str = "RECIBIDA",
) -> tuple[int, bool]:
    job = fetch_one(
        """
        SELECT id
        FROM jobs
        WHERE id = ? AND company_id = ?
        """,
        (job_id, user["company_id"]),
    )
    if not job:
        raise ValueError("La búsqueda seleccionada no pertenece a la empresa.")

    ensure_company_candidate(
        user["company_id"],
        candidate_id,
        "ASIGNACIÓN",
        user["id"],
    )

    existing = fetch_one(
        """
        SELECT id
        FROM applications
        WHERE candidate_id = ? AND job_id = ?
        """,
        (candidate_id, job_id),
    )

    if existing:
        execute(
            """
            UPDATE applications
            SET status = ?, updated_at = ?
            WHERE id = ?
            """,
            (initial_status, now_iso(), existing["id"]),
        )
        application_id = existing["id"]
        created = False
    else:
        application_id = execute(
            """
            INSERT INTO applications(
                candidate_id,
                job_id,
                status,
                created_at,
                updated_at
            )
            VALUES (?, ?, ?, ?, ?)
            """,
            (
                candidate_id,
                job_id,
                initial_status,
                now_iso(),
                now_iso(),
            ),
        )
        created = True

    score_application(application_id)

    log_event(
        user["company_id"],
        user["id"],
        "ASSIGN_CANDIDATE",
        "application",
        application_id,
        {
            "candidate_id": candidate_id,
            "job_id": job_id,
            "status": initial_status,
            "created": created,
        },
    )

    return application_id, created


def rescore_job_applications(job_id: int) -> int:
    rows = fetch_all(
        "SELECT id FROM applications WHERE job_id = ?",
        (job_id,),
    )

    for row in rows:
        score_application(row["id"])

    return len(rows)


def parse_terms_for_editor(value: str) -> list[str]:
    return extract_scoring_terms(value)


def score_dimension(
    label: str,
    terms: list[str],
    cv_text_normalized: str,
    weight: float,
    required: bool = False,
    key: str = "",
) -> dict:
    unique_terms = []
    seen = set()

    for term in terms:
        normalized = normalize_match_text(term)
        if normalized and normalized not in seen:
            unique_terms.append(normalized)
            seen.add(normalized)

    matched = [
        term
        for term in unique_terms
        if term in cv_text_normalized
    ]
    missing = [
        term
        for term in unique_terms
        if term not in cv_text_normalized
    ]

    coverage = (
        len(matched) / len(unique_terms)
        if unique_terms
        else None
    )

    return {
        "key": key,
        "label": label,
        "weight": float(weight or 0),
        "required": bool(required),
        "criteria_count": len(unique_terms),
        "matched": matched,
        "missing": missing,
        "coverage": coverage,
    }

def calculate_match_score(
    job: dict,
    cv_text: str,
) -> dict:
    cv_normalized = normalize_match_text(cv_text)
    scoring_config = get_job_scoring_config(job)

    dimensions = [
        score_dimension(
            group.get("label") or "Criterio",
            group.get("terms") or [],
            cv_normalized,
            float(group.get("weight") or 0),
            bool(group.get("required")),
            group.get("key") or "",
        )
        for group in scoring_config.get("groups", [])
    ]

    active_dimensions = [
        dimension
        for dimension in dimensions
        if dimension["criteria_count"] > 0
        and dimension["weight"] > 0
    ]

    if not active_dimensions:
        return {
            "total": None,
            "recommendation": "REVISIÓN MANUAL",
            "summary": (
                "La búsqueda no tiene criterios activos suficientes "
                "para calcular una coincidencia."
            ),
            "reasons": [
                "No existen criterios con términos y peso mayor a cero."
            ],
            "strengths": [],
            "gaps": [],
            "dimensions": dimensions,
            "config": scoring_config,
            "notice": SCORING_NOTICE,
        }

    active_weight = sum(
        dimension["weight"]
        for dimension in active_dimensions
    )

    weighted_points = sum(
        dimension["weight"]
        * float(dimension["coverage"])
        for dimension in active_dimensions
    )

    total = round(
        (weighted_points / active_weight) * 100,
        1,
    )

    for dimension in dimensions:
        if dimension["coverage"] is None or active_weight <= 0:
            dimension["points"] = None
            dimension["normalized_weight"] = 0
        else:
            normalized_weight = (
                dimension["weight"] / active_weight
            ) * 100
            dimension["normalized_weight"] = round(
                normalized_weight,
                1,
            )
            dimension["points"] = round(
                normalized_weight
                * float(dimension["coverage"]),
                1,
            )

    thresholds = scoring_config.get("thresholds") or {}
    advance_threshold = float(thresholds.get("advance", 75))
    review_threshold = float(thresholds.get("review", 50))
    minimum_required = (
        float(thresholds.get("minimum_required_coverage", 50))
        / 100
    )

    required_breaches = [
        dimension
        for dimension in active_dimensions
        if dimension["required"]
        and float(dimension["coverage"] or 0) < minimum_required
    ]

    if required_breaches:
        recommendation = "BAJA COINCIDENCIA"
    elif total >= advance_threshold:
        recommendation = "AVANZA"
    elif total >= review_threshold:
        recommendation = "REVISAR"
    else:
        recommendation = "BAJA COINCIDENCIA"

    matched_count = sum(
        len(dimension["matched"])
        for dimension in active_dimensions
    )
    criteria_count = sum(
        dimension["criteria_count"]
        for dimension in active_dimensions
    )

    strengths = []
    gaps = []
    reasons = []

    for dimension in active_dimensions:
        percentage = round(
            float(dimension["coverage"] or 0) * 100
        )

        if dimension["matched"]:
            strengths.append(
                f"{dimension['label']}: coincide con "
                + ", ".join(dimension["matched"][:8])
                + "."
            )

        if dimension["missing"]:
            gaps.append(
                f"{dimension['label']}: no se encontraron "
                + ", ".join(dimension["missing"][:8])
                + "."
            )

        reasons.append(
            f"{dimension['label']}: {percentage}% de cobertura, "
            f"{dimension.get('points') or 0} puntos."
        )

    if required_breaches:
        reasons.insert(
            0,
            "La recomendación baja porque uno o más grupos obligatorios "
            "quedaron por debajo de la cobertura mínima configurada."
        )
    elif recommendation == "AVANZA":
        reasons.insert(
            0,
            "Supera el umbral configurado para avanzar y no presenta "
            "incumplimientos críticos en criterios obligatorios."
        )
    elif recommendation == "REVISAR":
        reasons.insert(
            0,
            "La coincidencia es intermedia: reúne parte de los criterios, "
            "pero necesita validación humana."
        )
    else:
        reasons.insert(
            0,
            "El puntaje quedó por debajo del umbral de revisión."
        )

    return {
        "total": total,
        "recommendation": recommendation,
        "summary": (
            f"Coincide con {matched_count} de {criteria_count} criterios. "
            f"Plantilla activa: {scoring_config.get('preset_name') or 'General'}."
        ),
        "reasons": reasons,
        "strengths": strengths,
        "gaps": gaps,
        "dimensions": dimensions,
        "config": scoring_config,
        "notice": SCORING_NOTICE,
    }

def get_latest_cv(candidate_id: int) -> dict | None:
    return fetch_one(
        """
        SELECT *
        FROM candidate_documents
        WHERE candidate_id = ?
        ORDER BY created_at DESC, id DESC
        LIMIT 1
        """,
        (candidate_id,),
    )


def save_candidate_document(
    candidate_id: int,
    filename: str,
    mime_type: str,
    raw: bytes,
    extracted_text: str,
    parsed: dict,
    uploaded_by: int | None,
) -> int:
    return execute(
        """
        INSERT INTO candidate_documents(
            candidate_id,
            filename,
            mime_type,
            content_blob,
            extracted_text,
            parsed_json,
            uploaded_by,
            created_at
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            candidate_id,
            filename,
            mime_type,
            raw,
            extracted_text,
            json.dumps(parsed, ensure_ascii=False),
            uploaded_by,
            now_iso(),
        ),
    )


def update_candidate_profile(
    candidate_id: int,
    values: dict,
) -> None:
    execute(
        """
        UPDATE candidates
        SET phone = ?,
            birth_date = ?,
            country = ?,
            province = ?,
            city = ?,
            locality = ?,
            headline = ?,
            education_summary = ?,
            experience_summary = ?,
            skills_text = ?,
            languages_text = ?,
            tags_text = ?,
            updated_at = ?
        WHERE id = ?
        """,
        (
            values.get("phone", "").strip(),
            values.get("birth_date", "").strip(),
            values.get("country", "").strip(),
            values.get("province", "").strip(),
            values.get("city", "").strip(),
            values.get("locality", "").strip(),
            values.get("headline", "").strip(),
            values.get("education_summary", "").strip(),
            values.get("experience_summary", "").strip(),
            values.get("skills_text", "").strip(),
            values.get("languages_text", "").strip(),
            values.get("tags_text", "").strip(),
            now_iso(),
            candidate_id,
        ),
    )

def score_application(application_id: int) -> dict:
    application = fetch_one(
        """
        SELECT
            applications.id,
            applications.candidate_id,
            applications.job_id,
            jobs.*
        FROM applications
        JOIN jobs
            ON jobs.id = applications.job_id
        WHERE applications.id = ?
        """,
        (application_id,),
    )

    if not application:
        raise ValueError("No se encontró la postulación.")

    latest_cv = get_latest_cv(application["candidate_id"])

    if not latest_cv or not latest_cv.get("extracted_text"):
        result = {
            "total": None,
            "recommendation": "PENDIENTE DE CV",
            "summary": (
                "No hay un CV con texto disponible para analizar."
            ),
            "dimensions": [],
            "notice": SCORING_NOTICE,
        }
    else:
        result = calculate_match_score(
            application,
            latest_cv["extracted_text"],
        )

    execute(
        """
        UPDATE applications
        SET score_total = ?,
            score_breakdown_json = ?,
            screening_recommendation = ?,
            screening_summary = ?,
            screened_at = ?,
            updated_at = ?
        WHERE id = ?
        """,
        (
            result["total"],
            json.dumps(result, ensure_ascii=False),
            result["recommendation"],
            result["summary"],
            now_iso(),
            now_iso(),
            application_id,
        ),
    )

    return result


def rescore_candidate_applications(candidate_id: int) -> None:
    application_rows = fetch_all(
        """
        SELECT id
        FROM applications
        WHERE candidate_id = ?
        """,
        (candidate_id,),
    )

    for application_row in application_rows:
        score_application(application_row["id"])


def get_or_create_candidate_from_cv(
    parsed: dict,
    raw: bytes,
    mime_type: str,
    extracted_text: str,
    filename: str,
    uploaded_by: int,
) -> tuple[int, int]:
    email = normalize_email(parsed.get("email") or "")
    user = None

    if email:
        existing_user = fetch_one(
            "SELECT * FROM users WHERE email = ?",
            (email,),
        )
        if (
            existing_user
            and existing_user.get("account_type") == "CANDIDATE"
        ):
            user = existing_user

    if not user:
        digest = hashlib.sha256(raw).hexdigest()[:16]
        synthetic_email = (
            email
            if email and not fetch_one(
                "SELECT id FROM users WHERE email = ?",
                (email,),
            )
            else f"cv-{digest}@alba.local"
        )

        random_password = secrets.token_urlsafe(24)
        password_hash, salt = hash_password(random_password)

        user_id = execute(
            """
            INSERT INTO users(
                email,
                password_hash,
                password_salt,
                full_name,
                account_type,
                active,
                created_at
            )
            VALUES (?, ?, ?, ?, 'CANDIDATE', 0, ?)
            """,
            (
                synthetic_email,
                password_hash,
                salt,
                parsed.get("full_name")
                or "Candidato sin identificar",
                now_iso(),
            ),
        )

        user = fetch_one(
            "SELECT * FROM users WHERE id = ?",
            (user_id,),
        )

    candidate = fetch_one(
        "SELECT * FROM candidates WHERE user_id = ?",
        (user["id"],),
    )

    if not candidate:
        candidate_id = execute(
            """
            INSERT INTO candidates(
                user_id,
                phone,
                city,
                source,
                headline,
                education_summary,
                experience_summary,
                skills_text,
                languages_text,
                created_at,
                updated_at
            )
            VALUES (?, ?, ?, 'CARGA_RRHH', ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                user["id"],
                parsed.get("phone", ""),
                "",
                parsed.get("headline", ""),
                parsed.get("education_summary", ""),
                parsed.get("experience_summary", ""),
                parsed.get("skills_text", ""),
                parsed.get("languages_text", ""),
                now_iso(),
                now_iso(),
            ),
        )
    else:
        candidate_id = candidate["id"]

        execute(
            """
            UPDATE candidates
            SET phone = CASE
                    WHEN ? <> '' THEN ?
                    ELSE phone
                END,
                headline = ?,
                education_summary = ?,
                experience_summary = ?,
                skills_text = ?,
                languages_text = ?,
                updated_at = ?
            WHERE id = ?
            """,
            (
                parsed.get("phone", ""),
                parsed.get("phone", ""),
                parsed.get("headline", ""),
                parsed.get("education_summary", ""),
                parsed.get("experience_summary", ""),
                parsed.get("skills_text", ""),
                parsed.get("languages_text", ""),
                now_iso(),
                candidate_id,
            ),
        )

    document_id = save_candidate_document(
        candidate_id,
        filename,
        mime_type,
        raw,
        extracted_text,
        parsed,
        uploaded_by,
    )

    return candidate_id, document_id


def create_or_refresh_application(
    candidate_id: int,
    job_id: int,
) -> tuple[int, dict]:
    existing = fetch_one(
        """
        SELECT id
        FROM applications
        WHERE candidate_id = ?
          AND job_id = ?
        """,
        (candidate_id, job_id),
    )

    if existing:
        application_id = existing["id"]
        execute(
            """
            UPDATE applications
            SET updated_at = ?
            WHERE id = ?
            """,
            (now_iso(), application_id),
        )
    else:
        application_id = execute(
            """
            INSERT INTO applications(
                candidate_id,
                job_id,
                status,
                created_at,
                updated_at
            )
            VALUES (?, ?, 'RECIBIDA', ?, ?)
            """,
            (
                candidate_id,
                job_id,
                now_iso(),
                now_iso(),
            ),
        )

    result = score_application(application_id)
    return application_id, result


def render_score_report(report: dict) -> None:
    total = report.get("total")
    recommendation = report.get(
        "recommendation",
        "REVISIÓN MANUAL",
    )

    c1, c2 = st.columns(2)
    c1.metric(
        "Coincidencia",
        "Pendiente"
        if total is None
        else f"{float(total):.1f}%",
    )
    c2.metric("Recomendación automática", recommendation)

    st.write(report.get("summary") or "")
    st.caption(report.get("notice") or SCORING_NOTICE)

    reasons = report.get("reasons") or []
    if reasons:
        st.markdown("**Razones de la recomendación**")
        for reason in reasons:
            st.write(f"• {reason}")

    strengths = report.get("strengths") or []
    gaps = report.get("gaps") or []

    c3, c4 = st.columns(2)
    with c3:
        st.markdown("**Fortalezas detectadas**")
        if strengths:
            for item in strengths:
                st.write(f"• {item}")
        else:
            st.write("No se detectaron fortalezas suficientes.")

    with c4:
        st.markdown("**Brechas detectadas**")
        if gaps:
            for item in gaps:
                st.write(f"• {item}")
        else:
            st.write("No se detectaron brechas relevantes.")

    for dimension in report.get("dimensions", []):
        if not dimension.get("criteria_count"):
            continue

        required_label = (
            " · obligatorio"
            if dimension.get("required")
            else ""
        )
        with st.expander(
            f"{dimension['label']}{required_label} · "
            f"{round(float(dimension.get('coverage') or 0) * 100)}%",
        ):
            st.write(
                f"**Peso configurado:** "
                f"{dimension.get('weight', 0)}"
            )
            st.write(
                f"**Peso normalizado:** "
                f"{dimension.get('normalized_weight', 0)}%"
            )
            st.write(
                f"**Puntos obtenidos:** "
                f"{dimension.get('points', 0)}"
            )
            st.write(
                "**Coincidencias:** "
                + (
                    ", ".join(dimension.get("matched") or [])
                    or "Ninguna"
                )
            )
            st.write(
                "**Brechas:** "
                + (
                    ", ".join(dimension.get("missing") or [])
                    or "Ninguna"
                )
            )

def render_bulk_cv_upload(user: dict) -> None:
    st.markdown("### Carga masiva de CV")

    if not has_permission("manage_cv_pool"):
        st.warning(
            "No tenés permiso para cargar y analizar CV."
        )
        return

    jobs = fetch_all(
        """
        SELECT id, title, status
        FROM jobs
        WHERE company_id = ?
          AND status <> 'CERRADA'
        ORDER BY created_at DESC
        """,
        (user["company_id"],),
    )

    if not jobs:
        st.info(
            "Primero creá una búsqueda abierta o pausada."
        )
        return

    job_options = {
        f"{job['title']} · {job['status']}": job["id"]
        for job in jobs
    }

    selected_job_label = st.selectbox(
        "Búsqueda a la que se asignarán los CV",
        list(job_options.keys()),
        key="bulk_cv_job",
    )
    selected_job_id = job_options[selected_job_label]

    uploaded_files = st.file_uploader(
        "Subí hasta 50 CV en PDF o DOCX",
        type=["pdf", "docx"],
        accept_multiple_files=True,
        key="bulk_cv_files",
    )

    st.caption(
        "Cada archivo puede pesar hasta 5 MB. "
        + SCORING_NOTICE
    )

    if st.button(
        "Procesar CV y generar ranking",
        type="primary",
        key="process_bulk_cv",
    ):
        if not uploaded_files:
            st.error("Seleccioná al menos un CV.")
            return

        if len(uploaded_files) > MAX_CV_BATCH:
            st.error(
                f"El máximo por lote es {MAX_CV_BATCH} CV."
            )
            return

        results = []

        for uploaded_file in uploaded_files:
            try:
                raw, mime_type, text = extract_cv_text(
                    uploaded_file
                )
                parsed = parse_cv_text(
                    text,
                    uploaded_file.name,
                )
                candidate_id, document_id = (
                    get_or_create_candidate_from_cv(
                        parsed,
                        raw,
                        mime_type,
                        text,
                        uploaded_file.name,
                        user["id"],
                    )
                )

                ensure_company_candidate(
                    user["company_id"],
                    candidate_id,
                    "CARGA_CV",
                    user["id"],
                )

                application_id, score_result = (
                    create_or_refresh_application(
                        candidate_id,
                        selected_job_id,
                    )
                )

                results.append(
                    {
                        "Archivo": uploaded_file.name,
                        "Candidato": parsed.get("full_name"),
                        "Email detectado": (
                            parsed.get("email")
                            or "No detectado"
                        ),
                        "Puntaje": score_result.get("total"),
                        "Recomendación": score_result.get(
                            "recommendation"
                        ),
                        "Estado": "Procesado",
                    }
                )

                log_event(
                    user["company_id"],
                    user["id"],
                    "UPLOAD_CV",
                    "candidate_document",
                    document_id,
                    {
                        "application_id": application_id,
                        "job_id": selected_job_id,
                        "filename": uploaded_file.name,
                    },
                )

            except Exception as exc:
                results.append(
                    {
                        "Archivo": uploaded_file.name,
                        "Candidato": "",
                        "Email detectado": "",
                        "Puntaje": None,
                        "Recomendación": "",
                        "Estado": f"Error: {exc}",
                    }
                )

        st.session_state["bulk_cv_results"] = results
        st.success("El lote terminó de procesarse.")

    results = st.session_state.get(
        "bulk_cv_results",
        [],
    )
    if results:
        st.dataframe(
            pd.DataFrame(results),
            use_container_width=True,
            hide_index=True,
        )


def render_candidate_profile(
    user: dict,
    candidate: dict,
) -> None:
    st.subheader("Mi perfil y CV")

    latest_cv = get_latest_cv(candidate["id"])

    uploaded_file = st.file_uploader(
        "Subir o reemplazar CV",
        type=["pdf", "docx"],
        key=f"candidate_cv_{candidate['id']}",
    )

    if st.button(
        "Analizar CV",
        key=f"analyze_candidate_cv_{candidate['id']}",
    ):
        if uploaded_file is None:
            st.error("Seleccioná un CV.")
        else:
            try:
                raw, mime_type, text = extract_cv_text(
                    uploaded_file
                )
                parsed = parse_cv_text(
                    text,
                    uploaded_file.name,
                )

                st.session_state[
                    f"candidate_cv_draft_{candidate['id']}"
                ] = {
                    "filename": uploaded_file.name,
                    "mime_type": mime_type,
                    "raw": raw,
                    "text": text,
                    "parsed": parsed,
                }

                st.success(
                    "CV analizado. Revisá los datos y guardá el perfil."
                )
            except Exception as exc:
                st.error(str(exc))

    draft = st.session_state.get(
        f"candidate_cv_draft_{candidate['id']}",
        {},
    )
    parsed = draft.get("parsed", {})

    with st.form(
        f"candidate_profile_form_{candidate['id']}"
    ):
        st.text_input(
            "Nombre y apellido",
            value=user["full_name"],
            disabled=True,
        )
        st.text_input(
            "Correo electrónico",
            value=user["email"],
            disabled=True,
        )

        c1, c2 = st.columns(2)
        phone = c1.text_input(
            "Teléfono",
            value=(
                parsed.get("phone")
                or candidate.get("phone")
                or ""
            ),
        )
        birth_date = c2.text_input(
            "Fecha de nacimiento (AAAA-MM-DD)",
            value=candidate.get("birth_date") or "",
        )

        c3, c4 = st.columns(2)
        country = c3.text_input(
            "País",
            value=candidate.get("country") or "Argentina",
        )
        province = c4.text_input(
            "Provincia / Estado",
            value=candidate.get("province") or "",
        )

        c5, c6 = st.columns(2)
        city = c5.text_input(
            "Ciudad",
            value=candidate.get("city") or "",
        )
        locality = c6.text_input(
            "Localidad / Partido",
            value=candidate.get("locality") or "",
        )

        headline = st.text_area(
            "Perfil profesional",
            value=(
                parsed.get("headline")
                or candidate.get("headline")
                or ""
            ),
            height=100,
        )
        education_summary = st.text_area(
            "Formación",
            value=(
                parsed.get("education_summary")
                or candidate.get("education_summary")
                or ""
            ),
            height=130,
        )
        experience_summary = st.text_area(
            "Experiencia",
            value=(
                parsed.get("experience_summary")
                or candidate.get("experience_summary")
                or ""
            ),
            height=170,
        )
        skills_text = st.text_area(
            "Habilidades y herramientas",
            value=(
                parsed.get("skills_text")
                or candidate.get("skills_text")
                or ""
            ),
            height=120,
        )
        languages_text = st.text_area(
            "Idiomas",
            value=(
                parsed.get("languages_text")
                or candidate.get("languages_text")
                or ""
            ),
            height=80,
        )

        submitted = st.form_submit_button(
            "Guardar perfil y CV",
            type="primary",
        )

        if submitted:
            try:
                if birth_date:
                    date.fromisoformat(birth_date)

                update_candidate_profile(
                    candidate["id"],
                    {
                        "phone": phone,
                        "birth_date": birth_date,
                        "country": country,
                        "province": province,
                        "city": city,
                        "locality": locality,
                        "headline": headline,
                        "education_summary": education_summary,
                        "experience_summary": experience_summary,
                        "skills_text": skills_text,
                        "languages_text": languages_text,
                        "tags_text": candidate.get("tags_text") or "",
                    },
                )

                if draft:
                    document_id = save_candidate_document(
                        candidate["id"],
                        draft["filename"],
                        draft["mime_type"],
                        draft["raw"],
                        draft["text"],
                        draft["parsed"],
                        user["id"],
                    )
                    log_event(
                        None,
                        user["id"],
                        "UPLOAD_OWN_CV",
                        "candidate_document",
                        document_id,
                        {"candidate_id": candidate["id"]},
                    )
                    st.session_state.pop(
                        f"candidate_cv_draft_{candidate['id']}",
                        None,
                    )

                rescore_candidate_applications(
                    candidate["id"]
                )

                st.success(
                    "Perfil guardado y postulaciones recalculadas."
                )
                st.rerun()

            except ValueError:
                st.error(
                    "La fecha debe tener formato AAAA-MM-DD."
                )
            except Exception as exc:
                st.error(str(exc))

    if latest_cv:
        st.caption(
            f"Último CV: {latest_cv['filename']} · "
            f"{latest_cv['created_at']}"
        )
        st.download_button(
            "Descargar mi CV",
            data=latest_cv["content_blob"],
            file_name=latest_cv["filename"],
            mime=latest_cv.get("mime_type")
            or "application/octet-stream",
            key=f"download_own_cv_{latest_cv['id']}",
        )

def render_applications(user: dict) -> None:
    st.subheader("Candidatos, CV y banco inteligente")

    if not has_permission("view_candidates"):
        st.warning("No tenés permiso para ver candidatos.")
        return

    st.info(SCORING_NOTICE)

    (
        tab_job_cvs,
        tab_pipeline,
        tab_detail,
        tab_suggestions,
        tab_pool,
        tab_criteria,
        tab_upload,
    ) = st.tabs(
        [
            "CV por búsqueda",
            "Pipeline y filtros",
            "Detalle y decisión",
            "Banco inteligente",
            "Banco general",
            "Criterios de análisis",
            "Carga masiva de CV",
        ]
    )

    jobs = fetch_all(
        """
        SELECT *
        FROM jobs
        WHERE company_id = ?
        ORDER BY created_at DESC
        """,
        (user["company_id"],),
    )

    job_options = {
        f"{job['title']} · {job.get('status') or 'ABIERTA'}": job
        for job in jobs
    }

    # -----------------------------------------------------
    # CV ASIGNADOS A CADA BÚSQUEDA
    # -----------------------------------------------------
    with tab_job_cvs:
        if not jobs:
            st.info("No hay búsquedas registradas.")
        else:
            selected_job_label = st.selectbox(
                "Seleccionar búsqueda",
                list(job_options.keys()),
                key="job_cv_search",
            )
            selected_job = job_options[selected_job_label]

            assigned = fetch_all(
                """
                SELECT
                    applications.*,
                    candidates.id AS candidate_id,
                    candidates.country,
                    candidates.province,
                    candidates.city,
                    candidates.locality,
                    candidates.headline,
                    users.full_name,
                    users.email,
                    candidate_documents.id AS document_id,
                    candidate_documents.filename,
                    candidate_documents.mime_type,
                    candidate_documents.content_blob,
                    candidate_documents.created_at AS cv_created_at
                FROM applications
                JOIN candidates
                    ON candidates.id = applications.candidate_id
                JOIN users
                    ON users.id = candidates.user_id
                LEFT JOIN candidate_documents
                    ON candidate_documents.id = (
                        SELECT cd.id
                        FROM candidate_documents cd
                        WHERE cd.candidate_id = candidates.id
                        ORDER BY cd.created_at DESC, cd.id DESC
                        LIMIT 1
                    )
                WHERE applications.job_id = ?
                ORDER BY
                    applications.starred DESC,
                    CASE
                        WHEN applications.score_total IS NULL THEN 1
                        ELSE 0
                    END,
                    applications.score_total DESC,
                    applications.created_at DESC
                """,
                (selected_job["id"],),
            )

            c1, c2, c3 = st.columns(3)
            c1.metric("CV asignados", len(assigned))
            c2.metric(
                "Con recomendación AVANZA",
                sum(
                    1
                    for row in assigned
                    if row.get("screening_recommendation") == "AVANZA"
                ),
            )
            c3.metric(
                "Pendientes de revisión",
                sum(
                    1
                    for row in assigned
                    if row.get("screening_recommendation")
                    in {"REVISAR", "REVISIÓN MANUAL", None}
                ),
            )

            if not assigned:
                st.info(
                    "Todavía no hay candidatos asignados a esta búsqueda."
                )
            else:
                for row in assigned:
                    with st.container(border=True):
                        h1, h2, h3 = st.columns([3, 2, 2])

                        with h1:
                            st.markdown(
                                f"### {row['full_name']}"
                            )
                            st.write(
                                f"**Perfil:** "
                                f"{row.get('headline') or 'Sin resumen'}"
                            )
                            st.write(
                                f"**Ubicación:** "
                                f"{build_location_label(row)}"
                            )

                        with h2:
                            st.metric(
                                "Coincidencia",
                                (
                                    "Pendiente"
                                    if row.get("score_total") is None
                                    else f"{float(row['score_total']):.1f}%"
                                ),
                            )
                            st.write(
                                f"**Recomendación:** "
                                f"{row.get('screening_recommendation') or 'PENDIENTE'}"
                            )
                            st.write(
                                f"**Etapa:** {row.get('status')}"
                            )

                        with h3:
                            st.write(
                                f"**Decisión manual:** "
                                f"{row.get('manual_recommendation') or 'Sin definir'}"
                            )
                            if row.get("filename"):
                                st.download_button(
                                    "Descargar CV",
                                    data=row["content_blob"],
                                    file_name=row["filename"],
                                    mime=row.get("mime_type")
                                    or "application/octet-stream",
                                    key=f"job_cv_download_{row['id']}",
                                )
                            else:
                                st.warning("No tiene CV almacenado.")

                        try:
                            report = json.loads(
                                row.get("score_breakdown_json")
                                or "{}"
                            )
                        except json.JSONDecodeError:
                            report = {}

                        if report:
                            explanation = recommendation_explanation(report)
                            st.write(
                                f"**Por qué:** {explanation['headline']}"
                            )
                            with st.expander(
                                "Ver explicación completa"
                            ):
                                render_detailed_recommendation(report)

    # -----------------------------------------------------
    # PIPELINE Y FILTROS
    # -----------------------------------------------------
    with tab_pipeline:
        rows = fetch_all(
            """
            SELECT
                applications.*,
                users.full_name AS candidate,
                users.email,
                candidates.id AS candidate_id,
                candidates.birth_date,
                candidates.country,
                candidates.province,
                candidates.city,
                candidates.locality,
                candidates.tags_text,
                jobs.title AS job
            FROM applications
            JOIN candidates
                ON candidates.id = applications.candidate_id
            JOIN users
                ON users.id = candidates.user_id
            JOIN jobs
                ON jobs.id = applications.job_id
            WHERE jobs.company_id = ?
            ORDER BY
                applications.starred DESC,
                CASE
                    WHEN applications.score_total IS NULL THEN 1
                    ELSE 0
                END,
                applications.score_total DESC,
                applications.created_at DESC
            """,
            (user["company_id"],),
        )

        filter_job_options = {"Todas las búsquedas": None}
        filter_job_options.update(
            {
                label: job["id"]
                for label, job in job_options.items()
            }
        )

        f1, f2, f3 = st.columns(3)
        selected_job_label = f1.selectbox(
            "Búsqueda",
            list(filter_job_options.keys()),
            key="pipeline_job_filter",
        )
        selected_job_id = filter_job_options[
            selected_job_label
        ]
        selected_stage = f2.selectbox(
            "Etapa",
            ["Todas"] + APPLICATION_STATUSES,
            key="pipeline_stage_filter",
        )
        selected_recommendation = f3.selectbox(
            "Recomendación",
            [
                "Todas",
                "AVANZA",
                "REVISAR",
                "BAJA COINCIDENCIA",
                "PENDIENTE DE CV",
                "REVISIÓN MANUAL",
            ],
            key="pipeline_rec_filter",
        )

        countries = sorted(
            {
                row.get("country")
                for row in rows
                if row.get("country")
            }
        )
        cities = sorted(
            {
                row.get("city")
                for row in rows
                if row.get("city")
            }
        )
        localities = sorted(
            {
                row.get("locality")
                for row in rows
                if row.get("locality")
            }
        )

        l1, l2, l3 = st.columns(3)
        selected_country = l1.selectbox(
            "País",
            ["Todos"] + countries,
            key="pipeline_country",
        )
        selected_city = l2.selectbox(
            "Ciudad",
            ["Todas"] + cities,
            key="pipeline_city",
        )
        selected_locality = l3.selectbox(
            "Localidad",
            ["Todas"] + localities,
            key="pipeline_locality",
        )

        use_age_filter = st.checkbox(
            "Habilitar filtro administrativo de edad",
            key="pipeline_age_enabled",
        )
        min_age, max_age = 18, 80
        if use_age_filter:
            st.warning(AGE_FILTER_NOTICE)
            min_age, max_age = st.slider(
                "Rango de edad",
                16,
                90,
                (18, 65),
                key="pipeline_age_range",
            )

        filtered = []
        for row in rows:
            age = calculate_age(row.get("birth_date"))

            if (
                selected_job_id is not None
                and row["job_id"] != selected_job_id
            ):
                continue
            if (
                selected_stage != "Todas"
                and row["status"] != selected_stage
            ):
                continue
            if (
                selected_recommendation != "Todas"
                and (
                    row.get("screening_recommendation")
                    or "PENDIENTE DE CV"
                ) != selected_recommendation
            ):
                continue
            if (
                selected_country != "Todos"
                and row.get("country") != selected_country
            ):
                continue
            if (
                selected_city != "Todas"
                and row.get("city") != selected_city
            ):
                continue
            if (
                selected_locality != "Todas"
                and row.get("locality") != selected_locality
            ):
                continue
            if use_age_filter and (
                age is None
                or age < min_age
                or age > max_age
            ):
                continue

            item = dict(row)
            item["age"] = age
            filtered.append(item)

        display_rows = [
            {
                "⭐": "Sí" if row.get("starred") else "",
                "Postulación": row["id"],
                "Candidato": row["candidate"],
                "Búsqueda": row["job"],
                "Etapa": row["status"],
                "Puntaje": row.get("score_total"),
                "Recomendación": (
                    row.get("screening_recommendation")
                    or "PENDIENTE"
                ),
                "Decisión manual": (
                    row.get("manual_recommendation")
                    or "Sin definir"
                ),
                "Edad": row.get("age"),
                "País": row.get("country") or "",
                "Ciudad": row.get("city") or "",
                "Localidad": row.get("locality") or "",
                "Etiquetas": row.get("tags_text") or "",
            }
            for row in filtered
        ]

        if display_rows:
            dataframe = pd.DataFrame(display_rows)
            st.dataframe(
                dataframe,
                use_container_width=True,
                hide_index=True,
            )
            st.download_button(
                "Exportar resultados a CSV",
                data=dataframe.to_csv(
                    index=False
                ).encode("utf-8-sig"),
                file_name="pipeline_candidatos.csv",
                mime="text/csv",
                key="export_pipeline_csv",
            )
        else:
            st.info("No hay candidatos para los filtros seleccionados.")

    # -----------------------------------------------------
    # DETALLE Y DECISIÓN
    # -----------------------------------------------------
    with tab_detail:
        detail_rows = fetch_all(
            """
            SELECT
                applications.*,
                users.full_name AS candidate_name,
                users.email,
                candidates.id AS candidate_id,
                candidates.phone,
                candidates.birth_date,
                candidates.country,
                candidates.province,
                candidates.city,
                candidates.locality,
                candidates.tags_text,
                candidates.source,
                candidates.headline,
                candidates.education_summary,
                candidates.experience_summary,
                candidates.skills_text,
                candidates.languages_text,
                jobs.title AS job_title
            FROM applications
            JOIN candidates
                ON candidates.id = applications.candidate_id
            JOIN users
                ON users.id = candidates.user_id
            JOIN jobs
                ON jobs.id = applications.job_id
            WHERE jobs.company_id = ?
            ORDER BY applications.created_at DESC
            """,
            (user["company_id"],),
        )

        if not detail_rows:
            st.info("No hay postulaciones para mostrar.")
        else:
            options = {
                (
                    f"#{row['id']} · {row['candidate_name']} · "
                    f"{row['job_title']}"
                ): row
                for row in detail_rows
            }
            selected_label = st.selectbox(
                "Seleccionar postulación",
                list(options.keys()),
                key="detail_application",
            )
            selected = options[selected_label]
            selected_age = calculate_age(
                selected.get("birth_date")
            )

            h1, h2 = st.columns(2)
            with h1:
                st.markdown(
                    f"### {selected['candidate_name']}"
                )
                st.write(
                    f"**Búsqueda:** {selected['job_title']}"
                )
                st.write(
                    f"**Edad administrativa:** "
                    f"{selected_age if selected_age is not None else 'No informada'}"
                )
                st.write(
                    f"**Ubicación:** {build_location_label(selected)}"
                )
                st.write(
                    f"**Etiquetas:** "
                    f"{selected.get('tags_text') or 'Sin etiquetas'}"
                )

            with h2:
                st.metric(
                    "Recomendación automática",
                    selected.get("screening_recommendation")
                    or "PENDIENTE",
                )
                st.metric(
                    "Decisión humana",
                    selected.get("manual_recommendation")
                    or "Sin definir",
                )

            try:
                report = json.loads(
                    selected.get("score_breakdown_json")
                    or "{}"
                )
            except json.JSONDecodeError:
                report = {}

            if report:
                render_detailed_recommendation(report)
                with st.expander(
                    "Ver cálculo técnico completo"
                ):
                    render_score_report(report)
            else:
                st.info(
                    "La postulación todavía no tiene análisis automático."
                )

            if has_permission("manage_candidates"):
                with st.form(
                    f"application_management_{selected['id']}"
                ):
                    m1, m2 = st.columns(2)
                    current_status = (
                        selected.get("status")
                        if selected.get("status")
                        in APPLICATION_STATUSES
                        else "RECIBIDA"
                    )
                    new_status = m1.selectbox(
                        "Etapa del pipeline",
                        APPLICATION_STATUSES,
                        index=APPLICATION_STATUSES.index(
                            current_status
                        ),
                    )

                    manual_options = [
                        "Sin decisión manual",
                        "AVANZA",
                        "REVISAR",
                        "NO AVANZA",
                        "DESCARTAR",
                        "SELECCIONAR",
                    ]
                    current_manual = (
                        selected.get("manual_recommendation")
                        or "Sin decisión manual"
                    )
                    if current_manual not in manual_options:
                        current_manual = "Sin decisión manual"

                    manual_recommendation = m2.selectbox(
                        "Decisión manual",
                        manual_options,
                        index=manual_options.index(
                            current_manual
                        ),
                    )
                    manual_reason = st.text_area(
                        "Justificación de la decisión humana",
                        value=selected.get("manual_reason") or "",
                        help=(
                            "Describí las evidencias concretas que respaldan "
                            "la decisión. Es obligatoria cuando se define "
                            "una decisión manual."
                        ),
                    )
                    recruiter_notes = st.text_area(
                        "Notas internas",
                        value=selected.get("recruiter_notes") or "",
                    )
                    starred = st.checkbox(
                        "Marcar como favorito",
                        value=bool(selected.get("starred")),
                    )

                    st.markdown("**Datos administrativos**")
                    d1, d2 = st.columns(2)
                    birth_date_value = d1.text_input(
                        "Fecha de nacimiento (AAAA-MM-DD)",
                        value=selected.get("birth_date") or "",
                    )
                    tags_text = d2.text_input(
                        "Etiquetas",
                        value=selected.get("tags_text") or "",
                    )

                    d3, d4 = st.columns(2)
                    country = d3.text_input(
                        "País",
                        value=selected.get("country") or "",
                    )
                    province = d4.text_input(
                        "Provincia / Estado",
                        value=selected.get("province") or "",
                    )

                    d5, d6 = st.columns(2)
                    city = d5.text_input(
                        "Ciudad",
                        value=selected.get("city") or "",
                    )
                    locality = d6.text_input(
                        "Localidad / Partido",
                        value=selected.get("locality") or "",
                    )

                    submitted = st.form_submit_button(
                        "Guardar seguimiento",
                        type="primary",
                    )

                    if submitted:
                        try:
                            if birth_date_value:
                                date.fromisoformat(
                                    birth_date_value
                                )

                            if (
                                manual_recommendation
                                != "Sin decisión manual"
                                and not manual_reason.strip()
                            ):
                                raise ValueError(
                                    "Detallá la razón de la decisión manual."
                                )

                            execute(
                                """
                                UPDATE applications
                                SET status = ?,
                                    manual_recommendation = ?,
                                    manual_reason = ?,
                                    recruiter_notes = ?,
                                    starred = ?,
                                    assigned_recruiter_id = ?,
                                    updated_at = ?
                                WHERE id = ?
                                """,
                                (
                                    new_status,
                                    (
                                        ""
                                        if manual_recommendation
                                        == "Sin decisión manual"
                                        else manual_recommendation
                                    ),
                                    manual_reason.strip(),
                                    recruiter_notes.strip(),
                                    int(starred),
                                    user["id"],
                                    now_iso(),
                                    selected["id"],
                                ),
                            )

                            execute(
                                """
                                UPDATE candidates
                                SET birth_date = ?,
                                    country = ?,
                                    province = ?,
                                    city = ?,
                                    locality = ?,
                                    tags_text = ?,
                                    updated_at = ?
                                WHERE id = ?
                                """,
                                (
                                    birth_date_value.strip(),
                                    country.strip(),
                                    province.strip(),
                                    city.strip(),
                                    locality.strip(),
                                    tags_text.strip(),
                                    now_iso(),
                                    selected["candidate_id"],
                                ),
                            )

                            log_event(
                                user["company_id"],
                                user["id"],
                                "UPDATE_CANDIDATE_FOLLOWUP",
                                "application",
                                selected["id"],
                                {
                                    "status": new_status,
                                    "manual_recommendation": manual_recommendation,
                                    "starred": starred,
                                },
                            )
                            st.success("Seguimiento actualizado.")
                            st.rerun()

                        except Exception as exc:
                            st.error(str(exc))

            latest_cv = get_latest_cv(
                selected["candidate_id"]
            )
            if latest_cv:
                st.download_button(
                    "Descargar CV",
                    data=latest_cv["content_blob"],
                    file_name=latest_cv["filename"],
                    mime=latest_cv.get("mime_type")
                    or "application/octet-stream",
                    key=f"download_cv_{latest_cv['id']}",
                )

            if st.button(
                "Recalcular análisis",
                key=f"rescore_{selected['id']}",
            ):
                try:
                    score_application(selected["id"])
                    st.success("Análisis recalculado.")
                    st.rerun()
                except Exception as exc:
                    st.error(str(exc))

    # -----------------------------------------------------
    # BANCO INTELIGENTE
    # -----------------------------------------------------
    with tab_suggestions:
        if not jobs:
            st.info("No hay búsquedas para generar sugerencias.")
        else:
            open_jobs = {
                label: job
                for label, job in job_options.items()
                if job.get("status") == "ABIERTA"
            }

            if not open_jobs:
                st.info("No hay búsquedas abiertas.")
            else:
                suggestion_job_label = st.selectbox(
                    "Búsqueda para recibir sugerencias",
                    list(open_jobs.keys()),
                    key="suggestion_job",
                )
                suggestion_job = open_jobs[
                    suggestion_job_label
                ]

                st.write(
                    "ALBA compara los CV del banco con los criterios de "
                    "la búsqueda y propone perfiles que todavía no están asignados."
                )

                if st.button(
                    "Actualizar sugerencias",
                    type="primary",
                    key="generate_suggestions",
                ):
                    result = generate_candidate_job_suggestions(
                        user["company_id"],
                        suggestion_job["id"],
                    )
                    st.success(
                        f"Sugerencias evaluadas: {result['generated']}. "
                        f"CV omitidos por falta de texto: "
                        f"{result['skipped_without_cv']}."
                    )
                    st.rerun()

                suggestions = get_job_suggestions(
                    user["company_id"],
                    suggestion_job["id"],
                )

                if not suggestions:
                    st.info(
                        "No hay sugerencias activas. Actualizá el banco "
                        "o revisá los criterios de la búsqueda."
                    )
                else:
                    for suggestion in suggestions:
                        with st.container(border=True):
                            s1, s2, s3 = st.columns([3, 2, 2])

                            with s1:
                                st.markdown(
                                    f"### {suggestion['full_name']}"
                                )
                                st.write(
                                    f"**Perfil:** "
                                    f"{suggestion.get('headline') or 'Sin resumen'}"
                                )
                                st.write(
                                    f"**Ubicación:** "
                                    f"{build_location_label(suggestion)}"
                                )
                                st.write(
                                    f"**Etiquetas:** "
                                    f"{suggestion.get('tags_text') or 'Sin etiquetas'}"
                                )

                            with s2:
                                st.metric(
                                    "Coincidencia sugerida",
                                    (
                                        "Pendiente"
                                        if suggestion.get("score_total") is None
                                        else f"{float(suggestion['score_total']):.1f}%"
                                    ),
                                )
                                st.write(
                                    f"**Recomendación:** "
                                    f"{suggestion.get('recommendation') or 'REVISAR'}"
                                )

                            with s3:
                                if st.button(
                                    "Asignar a la búsqueda",
                                    type="primary",
                                    key=f"assign_suggestion_{suggestion['id']}",
                                ):
                                    assign_candidate_to_job(
                                        user,
                                        suggestion["candidate_id"],
                                        suggestion_job["id"],
                                        "EN REVISIÓN",
                                    )
                                    execute(
                                        """
                                        UPDATE candidate_job_suggestions
                                        SET status = 'ASIGNADA'
                                        WHERE id = ?
                                        """,
                                        (suggestion["id"],),
                                    )
                                    st.success(
                                        "Candidato asignado al pipeline."
                                    )
                                    st.rerun()

                            try:
                                report = json.loads(
                                    suggestion.get("report_json")
                                    or "{}"
                                )
                            except json.JSONDecodeError:
                                report = {}

                            if report:
                                explanation = recommendation_explanation(
                                    report
                                )
                                st.write(
                                    f"**Motivo de la sugerencia:** "
                                    f"{explanation['headline']}"
                                )
                                with st.expander(
                                    "Ver análisis detallado"
                                ):
                                    render_detailed_recommendation(report)

    # -----------------------------------------------------
    # BANCO GENERAL
    # -----------------------------------------------------
    with tab_pool:
        pool_rows = fetch_all(
            """
            SELECT
                candidates.*,
                users.full_name,
                users.email,
                company_candidates.source AS pool_source,
                company_candidates.created_at AS pool_created_at
            FROM company_candidates
            JOIN candidates
                ON candidates.id = company_candidates.candidate_id
            JOIN users
                ON users.id = candidates.user_id
            WHERE company_candidates.company_id = ?
            ORDER BY company_candidates.created_at DESC
            """,
            (user["company_id"],),
        )

        if not pool_rows:
            st.info("El banco está vacío.")
        else:
            search_text = st.text_input(
                "Buscar por nombre, email, etiquetas o perfil",
                key="pool_search",
            )
            normalized_search = normalize_match_text(search_text)

            countries = sorted(
                {
                    row.get("country")
                    for row in pool_rows
                    if row.get("country")
                }
            )
            cities = sorted(
                {
                    row.get("city")
                    for row in pool_rows
                    if row.get("city")
                }
            )
            localities = sorted(
                {
                    row.get("locality")
                    for row in pool_rows
                    if row.get("locality")
                }
            )

            p1, p2, p3 = st.columns(3)
            pool_country = p1.selectbox(
                "País",
                ["Todos"] + countries,
                key="pool_country",
            )
            pool_city = p2.selectbox(
                "Ciudad",
                ["Todas"] + cities,
                key="pool_city",
            )
            pool_locality = p3.selectbox(
                "Localidad",
                ["Todas"] + localities,
                key="pool_locality",
            )

            filtered_pool = []
            for row in pool_rows:
                searchable = normalize_match_text(
                    " ".join(
                        [
                            row.get("full_name") or "",
                            row.get("email") or "",
                            row.get("headline") or "",
                            row.get("skills_text") or "",
                            row.get("tags_text") or "",
                        ]
                    )
                )

                if (
                    normalized_search
                    and normalized_search not in searchable
                ):
                    continue
                if (
                    pool_country != "Todos"
                    and row.get("country") != pool_country
                ):
                    continue
                if (
                    pool_city != "Todas"
                    and row.get("city") != pool_city
                ):
                    continue
                if (
                    pool_locality != "Todas"
                    and row.get("locality") != pool_locality
                ):
                    continue

                filtered_pool.append(row)

            candidate_options = {
                (
                    f"#{row['id']} · {row['full_name']} · "
                    f"{build_location_label(row)}"
                ): row["id"]
                for row in filtered_pool
            }

            selected_candidates = st.multiselect(
                "Seleccionar candidatos para asignar",
                list(candidate_options.keys()),
                key="pool_candidates_assign",
            )

            if jobs:
                a1, a2 = st.columns(2)
                target_job_label = a1.selectbox(
                    "Asignar a la búsqueda",
                    list(job_options.keys()),
                    key="pool_target_job",
                )
                target_status = a2.selectbox(
                    "Etapa inicial",
                    APPLICATION_STATUSES,
                    key="pool_target_status",
                )

                if st.button(
                    "Asignar seleccionados",
                    type="primary",
                    key="assign_pool_candidates",
                ):
                    if not selected_candidates:
                        st.error(
                            "Seleccioná al menos un candidato."
                        )
                    else:
                        target_job = job_options[
                            target_job_label
                        ]
                        created_count = 0
                        updated_count = 0

                        for label in selected_candidates:
                            _, created = assign_candidate_to_job(
                                user,
                                candidate_options[label],
                                target_job["id"],
                                target_status,
                            )
                            if created:
                                created_count += 1
                            else:
                                updated_count += 1

                        st.success(
                            f"Nuevas asignaciones: {created_count}. "
                            f"Actualizadas: {updated_count}."
                        )
                        st.rerun()

            pool_display = pd.DataFrame(
                [
                    {
                        "ID": row["id"],
                        "Candidato": row["full_name"],
                        "Email": (
                            "Sin cuenta activa"
                            if str(row["email"]).endswith(
                                "@alba.local"
                            )
                            else row["email"]
                        ),
                        "País": row.get("country") or "",
                        "Ciudad": row.get("city") or "",
                        "Localidad": row.get("locality") or "",
                        "Perfil": row.get("headline") or "",
                        "Etiquetas": row.get("tags_text") or "",
                        "Origen": row.get("pool_source") or "",
                    }
                    for row in filtered_pool
                ]
            )

            if not pool_display.empty:
                st.dataframe(
                    pool_display,
                    use_container_width=True,
                    hide_index=True,
                )

    # -----------------------------------------------------
    # CRITERIOS
    # -----------------------------------------------------
    with tab_criteria:
        if not has_permission("manage_jobs"):
            st.warning(
                "No tenés permiso para editar criterios."
            )
        elif not jobs:
            st.info("No hay búsquedas.")
        else:
            criteria_job_label = st.selectbox(
                "Búsqueda",
                list(job_options.keys()),
                key="criteria_job",
            )
            criteria_job = job_options[criteria_job_label]
            saved_config = get_job_scoring_config(
                criteria_job
            )

            preset_options = [
                "Automático",
                *ANALYSIS_PRESETS.keys(),
            ]
            current_preset = (
                saved_config.get("preset_name")
                or detect_analysis_preset(criteria_job)
            )
            if current_preset not in preset_options:
                current_preset = "Automático"

            selected_preset = st.selectbox(
                "Plantilla predeterminada",
                preset_options,
                index=preset_options.index(current_preset),
                key=f"criteria_preset_{criteria_job['id']}",
            )

            if st.button(
                "Cargar criterios sugeridos",
                key=f"load_preset_{criteria_job['id']}",
            ):
                st.session_state[
                    f"criteria_draft_{criteria_job['id']}"
                ] = default_scoring_config(
                    criteria_job,
                    selected_preset,
                )
                st.rerun()

            config = st.session_state.get(
                f"criteria_draft_{criteria_job['id']}",
                saved_config,
            )
            groups_by_key = {
                group.get("key"): group
                for group in config.get("groups", [])
            }

            with st.form(
                f"criteria_form_{criteria_job['id']}"
            ):
                st.caption(
                    "Los criterios profesionales determinan el puntaje. "
                    "La ubicación y la edad solo se utilizan como filtros administrativos."
                )

                edited_groups = []

                definitions = [
                    (
                        "must_have",
                        "Requisitos excluyentes",
                        criteria_job.get("must_have") or "",
                        50.0,
                        True,
                    ),
                    (
                        "desirable",
                        "Requisitos deseables",
                        criteria_job.get("desirable") or "",
                        20.0,
                        False,
                    ),
                    (
                        "competencies",
                        "Competencias",
                        criteria_job.get("competencies") or "",
                        15.0,
                        False,
                    ),
                    (
                        "context",
                        "Contexto del puesto",
                        " | ".join(
                            [
                                criteria_job.get("title") or "",
                                criteria_job.get("area") or "",
                                criteria_job.get("seniority") or "",
                            ]
                        ),
                        10.0,
                        False,
                    ),
                    (
                        "preset",
                        "Criterios sugeridos por plantilla",
                        ", ".join(
                            ANALYSIS_PRESETS.get(
                                (
                                    detect_analysis_preset(
                                        criteria_job
                                    )
                                    if selected_preset
                                    == "Automático"
                                    else selected_preset
                                ),
                                [],
                            )
                        ),
                        5.0,
                        False,
                    ),
                    (
                        "custom",
                        "Criterios personalizados",
                        "",
                        0.0,
                        False,
                    ),
                ]

                for (
                    key,
                    label,
                    fallback_text,
                    fallback_weight,
                    fallback_required,
                ) in definitions:
                    existing_group = groups_by_key.get(key, {})
                    st.markdown(f"**{label}**")
                    terms_text = st.text_area(
                        "Términos separados por coma o renglón",
                        value=", ".join(
                            existing_group.get("terms")
                            or extract_scoring_terms(
                                fallback_text
                            )
                        ),
                        key=f"terms_{criteria_job['id']}_{key}",
                    )
                    w1, w2 = st.columns(2)
                    weight = w1.number_input(
                        "Peso",
                        min_value=0.0,
                        max_value=100.0,
                        value=float(
                            existing_group.get(
                                "weight",
                                fallback_weight,
                            )
                        ),
                        step=1.0,
                        key=f"weight_{criteria_job['id']}_{key}",
                    )
                    required = w2.checkbox(
                        "Obligatorio",
                        value=bool(
                            existing_group.get(
                                "required",
                                fallback_required,
                            )
                        ),
                        key=f"required_{criteria_job['id']}_{key}",
                    )

                    edited_groups.append(
                        {
                            "key": key,
                            "label": label,
                            "terms": parse_terms_for_editor(
                                terms_text
                            ),
                            "weight": float(weight),
                            "required": bool(required),
                        }
                    )

                thresholds = config.get("thresholds") or {}
                st.markdown("**Umbrales**")
                t1, t2, t3 = st.columns(3)
                advance = t1.number_input(
                    "AVANZA desde",
                    0.0,
                    100.0,
                    float(thresholds.get("advance", 75)),
                    1.0,
                )
                review = t2.number_input(
                    "REVISAR desde",
                    0.0,
                    100.0,
                    float(thresholds.get("review", 50)),
                    1.0,
                )
                minimum_required = t3.number_input(
                    "Cobertura mínima obligatoria",
                    0.0,
                    100.0,
                    float(
                        thresholds.get(
                            "minimum_required_coverage",
                            50,
                        )
                    ),
                    1.0,
                )

                save_and_rescore = st.form_submit_button(
                    "Guardar y recalcular",
                    type="primary",
                )

                if save_and_rescore:
                    try:
                        if review > advance:
                            raise ValueError(
                                "El umbral de REVISAR no puede "
                                "ser mayor que el de AVANZA."
                            )

                        resolved_preset = (
                            detect_analysis_preset(criteria_job)
                            if selected_preset == "Automático"
                            else selected_preset
                        )

                        save_job_scoring_config(
                            user,
                            criteria_job,
                            resolved_preset,
                            edited_groups,
                            {
                                "advance": float(advance),
                                "review": float(review),
                                "minimum_required_coverage": float(
                                    minimum_required
                                ),
                            },
                        )
                        rescored = rescore_job_applications(
                            criteria_job["id"]
                        )
                        generate_candidate_job_suggestions(
                            user["company_id"],
                            criteria_job["id"],
                        )
                        st.session_state.pop(
                            f"criteria_draft_{criteria_job['id']}",
                            None,
                        )
                        st.success(
                            f"Criterios guardados. "
                            f"Postulaciones recalculadas: {rescored}."
                        )
                        st.rerun()

                    except Exception as exc:
                        st.error(str(exc))

    with tab_upload:
        render_bulk_cv_upload(user)


# =========================================================
# ENTREVISTAS ESTRUCTURADAS CON ALBA
# =========================================================


# =========================================================
# AVATAR ALBA · D-ID
# =========================================================

def get_secret_value(key: str, default: str = "") -> str:
    try:
        value = st.secrets.get(key)
        if value is not None:
            return str(value).strip()
    except Exception:
        pass

    for section_name in (
        "did",
        "smtp",
        "email",
        "app",
    ):
        try:
            section = st.secrets.get(section_name, {})
            value = section.get(key)
            if value is not None:
                return str(value).strip()
        except Exception:
            continue

    return default


def get_did_api_key() -> str:
    return get_secret_value("DID_API_KEY")


def did_credentials_configured() -> bool:
    key = get_did_api_key()
    return bool(key and ":" in key)


def get_avatar_settings(company_id: int) -> dict:
    saved = fetch_one(
        "SELECT * FROM avatar_settings WHERE company_id = ?",
        (company_id,),
    )

    defaults = {
        "company_id": company_id,
        "enabled": 0,
        "display_name": "ALBA",
        "source_url": get_secret_value("DID_SOURCE_URL"),
        "voice_provider": "microsoft",
        "voice_id": get_secret_value(
            "DID_VOICE_ID",
            "es-AR-ElenaNeural",
        ),
        "welcome_text": (
            "Hola, soy ALBA. Voy a acompañarte durante "
            "esta entrevista estructurada."
        ),
    }

    if not saved:
        return defaults

    result = dict(defaults)
    result.update(saved)

    if not result.get("source_url"):
        result["source_url"] = defaults["source_url"]
    if not result.get("voice_id"):
        result["voice_id"] = defaults["voice_id"]

    return result


def save_avatar_settings(
    user: dict,
    enabled: bool,
    display_name: str,
    source_url: str,
    voice_provider: str,
    voice_id: str,
    welcome_text: str,
) -> None:
    source_url = source_url.strip()

    if enabled:
        if not source_url:
            raise ValueError(
                "Configurá una URL pública para la imagen de ALBA."
            )
        if not re.match(
            r"^https://.+\.(?:jpg|jpeg|png)(?:\?.*)?$",
            source_url,
            flags=re.I,
        ):
            raise ValueError(
                "La imagen debe tener una URL HTTPS pública "
                "que termine en JPG, JPEG o PNG."
            )

    existing = fetch_one(
        "SELECT id FROM avatar_settings WHERE company_id = ?",
        (user["company_id"],),
    )

    values = (
        int(enabled),
        display_name.strip() or "ALBA",
        source_url,
        voice_provider.strip() or "microsoft",
        voice_id.strip(),
        welcome_text.strip(),
        now_iso(),
        user["company_id"],
    )

    if existing:
        execute(
            """
            UPDATE avatar_settings
            SET enabled = ?,
                display_name = ?,
                source_url = ?,
                voice_provider = ?,
                voice_id = ?,
                welcome_text = ?,
                updated_at = ?
            WHERE company_id = ?
            """,
            values,
        )
    else:
        execute(
            """
            INSERT INTO avatar_settings(
                enabled,
                display_name,
                source_url,
                voice_provider,
                voice_id,
                welcome_text,
                updated_at,
                company_id
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            values,
        )

    log_event(
        user["company_id"],
        user["id"],
        "UPDATE_AVATAR_SETTINGS",
        "avatar_settings",
        user["company_id"],
        {
            "enabled": bool(enabled),
            "display_name": display_name,
            "source_url_configured": bool(source_url),
            "voice_provider": voice_provider,
            "voice_id": voice_id,
        },
    )


def did_headers() -> dict:
    api_key = get_did_api_key()

    if not api_key:
        raise RuntimeError(
            "Falta DID_API_KEY en Streamlit Secrets."
        )
    if ":" not in api_key:
        raise RuntimeError(
            "DID_API_KEY debe tener el formato "
            "API_USERNAME:API_PASSWORD."
        )

    encoded = base64.b64encode(
        api_key.encode("utf-8")
    ).decode("ascii")

    return {
        "Authorization": f"Basic {encoded}",
        "Content-Type": "application/json",
        "Accept": "application/json",
    }


def did_create_talk(
    text: str,
    source_url: str,
    voice_provider: str,
    voice_id: str,
) -> str:
    try:
        import requests
    except ImportError as exc:
        raise RuntimeError("Falta instalar requests.") from exc

    script = {
        "type": "text",
        "input": text.strip(),
    }

    if voice_id.strip():
        script["provider"] = {
            "type": voice_provider.strip() or "microsoft",
            "voice_id": voice_id.strip(),
        }

    response = requests.post(
        f"{DID_API_BASE}/talks",
        headers=did_headers(),
        json={
            "source_url": source_url,
            "script": script,
            "name": "ALBA interview question",
        },
        timeout=40,
    )

    if response.status_code not in {200, 201}:
        raise RuntimeError(
            f"D-ID rechazó la creación ({response.status_code}): "
            f"{response.text[:1000]}"
        )

    payload = response.json()
    talk_id = payload.get("id")
    if not talk_id:
        raise RuntimeError(
            "D-ID no devolvió un identificador."
        )

    return str(talk_id)


def did_get_talk(talk_id: str) -> dict:
    try:
        import requests
    except ImportError as exc:
        raise RuntimeError("Falta instalar requests.") from exc

    response = requests.get(
        f"{DID_API_BASE}/talks/{talk_id}",
        headers=did_headers(),
        timeout=30,
    )

    if response.status_code != 200:
        raise RuntimeError(
            f"No se pudo consultar D-ID ({response.status_code}): "
            f"{response.text[:1000]}"
        )

    return response.json()


def avatar_content_hash(
    question: dict,
    settings: dict,
) -> str:
    payload = json.dumps(
        {
            "question": question.get("question_text") or "",
            "source_url": settings.get("source_url") or "",
            "voice_provider": settings.get("voice_provider") or "",
            "voice_id": settings.get("voice_id") or "",
        },
        sort_keys=True,
        ensure_ascii=False,
    )

    return hashlib.sha256(
        payload.encode("utf-8")
    ).hexdigest()


def get_cached_avatar_video(
    company_id: int,
    question: dict,
    settings: dict,
) -> dict | None:
    return fetch_one(
        """
        SELECT *
        FROM avatar_question_videos
        WHERE company_id = ?
          AND question_id = ?
          AND content_hash = ?
        ORDER BY id DESC
        LIMIT 1
        """,
        (
            company_id,
            question["id"],
            avatar_content_hash(question, settings),
        ),
    )


def generate_avatar_question_video(
    company_id: int,
    question: dict,
    settings: dict,
    force: bool = False,
) -> dict:
    if not bool(settings.get("enabled")):
        raise ValueError(
            "El avatar ALBA no está habilitado."
        )
    if not did_credentials_configured():
        raise RuntimeError(
            "D-ID no está configurado en Streamlit Secrets."
        )

    source_url = settings.get("source_url") or ""
    if not source_url:
        raise ValueError(
            "Falta la URL pública de la imagen de ALBA."
        )

    content_hash = avatar_content_hash(
        question,
        settings,
    )
    cached = get_cached_avatar_video(
        company_id,
        question,
        settings,
    )

    if (
        cached
        and not force
        and cached.get("status") == "LISTO"
        and cached.get("result_url")
    ):
        return cached

    if cached:
        cache_id = cached["id"]
        execute(
            """
            UPDATE avatar_question_videos
            SET status = 'GENERANDO',
                result_url = '',
                error_text = '',
                updated_at = ?
            WHERE id = ?
            """,
            (now_iso(), cache_id),
        )
    else:
        cache_id = execute(
            """
            INSERT INTO avatar_question_videos(
                company_id,
                question_id,
                content_hash,
                status,
                created_at,
                updated_at
            )
            VALUES (?, ?, ?, 'GENERANDO', ?, ?)
            """,
            (
                company_id,
                question["id"],
                content_hash,
                now_iso(),
                now_iso(),
            ),
        )

    try:
        talk_id = did_create_talk(
            question["question_text"],
            source_url,
            settings.get("voice_provider") or "microsoft",
            settings.get("voice_id") or "",
        )

        execute(
            """
            UPDATE avatar_question_videos
            SET talk_id = ?,
                status = 'GENERANDO',
                updated_at = ?
            WHERE id = ?
            """,
            (talk_id, now_iso(), cache_id),
        )

        last_payload = {}

        for _ in range(DID_POLL_ATTEMPTS):
            payload = did_get_talk(talk_id)
            last_payload = payload
            result_url = payload.get("result_url")
            status = str(
                payload.get("status") or ""
            ).lower()

            if result_url:
                execute(
                    """
                    UPDATE avatar_question_videos
                    SET status = 'LISTO',
                        result_url = ?,
                        error_text = '',
                        updated_at = ?
                    WHERE id = ?
                    """,
                    (result_url, now_iso(), cache_id),
                )
                return fetch_one(
                    "SELECT * FROM avatar_question_videos WHERE id = ?",
                    (cache_id,),
                )

            if status in {"error", "failed", "rejected"}:
                raise RuntimeError(
                    "D-ID informó un error: "
                    + json.dumps(
                        payload,
                        ensure_ascii=False,
                    )[:1500]
                )

            time.sleep(DID_POLL_SECONDS)

        execute(
            """
            UPDATE avatar_question_videos
            SET status = 'PROCESANDO',
                error_text = ?,
                updated_at = ?
            WHERE id = ?
            """,
            (
                "El video continúa procesándose. "
                + json.dumps(
                    last_payload,
                    ensure_ascii=False,
                )[:700],
                now_iso(),
                cache_id,
            ),
        )

        return fetch_one(
            "SELECT * FROM avatar_question_videos WHERE id = ?",
            (cache_id,),
        )

    except Exception as exc:
        execute(
            """
            UPDATE avatar_question_videos
            SET status = 'ERROR',
                error_text = ?,
                updated_at = ?
            WHERE id = ?
            """,
            (str(exc)[:1800], now_iso(), cache_id),
        )
        raise


def render_avatar_identity(settings: dict) -> None:
    c1, c2 = st.columns([1, 4])

    with c1:
        if settings.get("source_url"):
            st.image(
                settings["source_url"],
                width=110,
            )
        else:
            st.markdown("## 🤖")

    with c2:
        st.markdown(
            f"### {settings.get('display_name') or 'ALBA'}"
        )
        st.write(
            settings.get("welcome_text")
            or "Voy a acompañarte durante la entrevista."
        )
        st.caption(AVATAR_NOTICE)


def render_avatar_settings(user: dict) -> None:
    st.subheader("Avatar ALBA")

    settings = get_avatar_settings(
        user["company_id"]
    )

    c1, c2 = st.columns(2)
    c1.metric(
        "Credenciales D-ID",
        "Configuradas"
        if did_credentials_configured()
        else "Pendientes",
    )
    c2.metric(
        "Imagen de ALBA",
        "Configurada"
        if settings.get("source_url")
        else "Pendiente",
    )

    st.info(
        "La clave se configura en Streamlit Secrets "
        "y nunca se guarda en GitHub ni en la base."
    )

    with st.expander(
        "Ejemplo de Streamlit Secrets",
        expanded=not did_credentials_configured(),
    ):
        st.code(
            'DID_API_KEY = "API_USERNAME:API_PASSWORD"\n'
            'DID_SOURCE_URL = "https://tu-dominio.com/alba.png"\n'
            'DID_VOICE_ID = "es-AR-ElenaNeural"',
            language="toml",
        )

    with st.form("avatar_settings_form"):
        enabled = st.checkbox(
            "Habilitar el avatar parlante",
            value=bool(settings.get("enabled")),
        )
        display_name = st.text_input(
            "Nombre visible",
            value=settings.get("display_name") or "ALBA",
        )
        source_url = st.text_input(
            "URL pública de la imagen de ALBA",
            value=settings.get("source_url") or "",
            help="Debe ser HTTPS y terminar en JPG, JPEG o PNG.",
        )
        providers = [
            "microsoft",
            "amazon",
            "elevenlabs",
        ]
        current_provider = (
            settings.get("voice_provider")
            if settings.get("voice_provider") in providers
            else "microsoft"
        )
        voice_provider = st.selectbox(
            "Proveedor de voz",
            providers,
            index=providers.index(current_provider),
        )
        voice_id = st.text_input(
            "Identificador de voz",
            value=settings.get("voice_id")
            or "es-AR-ElenaNeural",
        )
        welcome_text = st.text_area(
            "Presentación de ALBA",
            value=settings.get("welcome_text")
            or (
                "Hola, soy ALBA. Voy a acompañarte "
                "durante esta entrevista estructurada."
            ),
        )
        submitted = st.form_submit_button(
            "Guardar configuración",
            type="primary",
        )

        if submitted:
            try:
                save_avatar_settings(
                    user,
                    enabled,
                    display_name,
                    source_url,
                    voice_provider,
                    voice_id,
                    welcome_text,
                )
                st.success("Configuración guardada.")
                st.rerun()
            except Exception as exc:
                st.error(str(exc))

    settings = get_avatar_settings(
        user["company_id"]
    )
    render_avatar_identity(settings)

    templates = get_interview_templates(
        user["company_id"],
        active_only=True,
    )

    st.markdown("#### Generar videos para las preguntas")

    if not templates:
        st.info(
            "Primero creá o activá una plantilla."
        )
        return

    template_options = {
        f"#{item['id']} · {item['name']}": item
        for item in templates
    }
    selected_label = st.selectbox(
        "Plantilla",
        list(template_options.keys()),
        key="avatar_template_select",
    )
    selected_template = template_options[selected_label]
    questions = get_template_questions(
        selected_template["id"],
        active_only=True,
    )

    status_rows = []
    for question in questions:
        cached = get_cached_avatar_video(
            user["company_id"],
            question,
            settings,
        )
        status_rows.append(
            {
                "Orden": question["position"],
                "Competencia": question.get("competency"),
                "Estado": (
                    cached.get("status")
                    if cached
                    else "NO GENERADO"
                ),
                "Video disponible": (
                    "Sí"
                    if cached and cached.get("result_url")
                    else "No"
                ),
            }
        )

    if status_rows:
        st.dataframe(
            pd.DataFrame(status_rows),
            use_container_width=True,
            hide_index=True,
        )

    force = st.checkbox(
        "Regenerar videos existentes",
        value=False,
        key="avatar_force",
    )

    if st.button(
        "Generar videos de ALBA",
        type="primary",
        key="generate_avatar_videos",
    ):
        if not bool(settings.get("enabled")):
            st.error("Habilitá el avatar primero.")
        elif not did_credentials_configured():
            st.error(
                "Configurá DID_API_KEY en Streamlit Secrets."
            )
        elif not questions:
            st.error(
                "La plantilla no tiene preguntas activas."
            )
        else:
            progress = st.progress(0)
            results = []

            for index, question in enumerate(
                questions,
                start=1,
            ):
                try:
                    video = generate_avatar_question_video(
                        user["company_id"],
                        question,
                        settings,
                        force=force,
                    )
                    results.append(
                        {
                            "Pregunta": question["position"],
                            "Estado": video.get("status"),
                            "Detalle": (
                                video.get("error_text")
                                or "Video disponible"
                            ),
                        }
                    )
                except Exception as exc:
                    results.append(
                        {
                            "Pregunta": question["position"],
                            "Estado": "ERROR",
                            "Detalle": str(exc),
                        }
                    )

                progress.progress(
                    index / len(questions)
                )

            st.session_state[
                "avatar_generation_results"
            ] = results
            st.success("Finalizó la generación.")

    results = st.session_state.get(
        "avatar_generation_results",
        [],
    )
    if results:
        st.dataframe(
            pd.DataFrame(results),
            use_container_width=True,
            hide_index=True,
        )

    st.markdown("#### Vista previa")

    for question in questions:
        cached = get_cached_avatar_video(
            user["company_id"],
            question,
            settings,
        )

        with st.expander(
            f"{question['position']}. "
            f"{question.get('competency') or 'Pregunta'}"
        ):
            st.write(question["question_text"])

            if cached and cached.get("result_url"):
                st.video(cached["result_url"])
            elif cached and cached.get("error_text"):
                st.warning(cached["error_text"])
            else:
                st.caption(
                    "Todavía no hay un video disponible."
                )


def render_question_avatar(
    company_id: int,
    question: dict,
    settings: dict,
) -> None:
    cached = get_cached_avatar_video(
        company_id,
        question,
        settings,
    )

    c1, c2 = st.columns([2, 3])

    with c1:
        if (
            bool(settings.get("enabled"))
            and cached
            and cached.get("result_url")
        ):
            st.video(cached["result_url"])
        elif settings.get("source_url"):
            st.image(
                settings["source_url"],
                width=220,
            )
            st.caption(
                "ALBA presenta la pregunta por escrito."
            )
        else:
            st.markdown("## 🤖 ALBA")

    with c2:
        st.markdown(
            f"### {question['position']}. "
            f"{question.get('competency') or 'Pregunta'}"
        )
        st.write(question["question_text"])


# =========================================================
# INVITACIONES DE ENTREVISTA POR EMAIL
# =========================================================

def secret_bool(
    key: str,
    default: bool = False,
) -> bool:
    value = get_secret_value(
        key,
        "true" if default else "false",
    )
    return str(value).strip().lower() in {
        "1",
        "true",
        "yes",
        "si",
        "sí",
        "on",
    }


def smtp_configuration() -> dict:
    host = get_secret_value("SMTP_HOST")
    username = get_secret_value("SMTP_USERNAME")
    password = get_secret_value("SMTP_PASSWORD")
    from_email = get_secret_value(
        "SMTP_FROM_EMAIL",
        username,
    )

    try:
        port = int(
            get_secret_value(
                "SMTP_PORT",
                "465",
            )
        )
    except ValueError:
        port = 465

    return {
        "host": host,
        "port": port,
        "username": username,
        "password": password,
        "from_email": from_email,
        "from_name": get_secret_value(
            "SMTP_FROM_NAME",
            "ALBA Selección",
        ),
        "reply_to": get_secret_value(
            "SMTP_REPLY_TO",
        ),
        "use_ssl": secret_bool(
            "SMTP_USE_SSL",
            default=port == 465,
        ),
        "use_tls": secret_bool(
            "SMTP_USE_TLS",
            default=port != 465,
        ),
        "app_base_url": get_secret_value(
            "APP_BASE_URL",
        ).rstrip("/"),
    }


def smtp_configured() -> bool:
    config = smtp_configuration()
    return bool(
        config["host"]
        and config["from_email"]
        and config["app_base_url"]
    )


def parse_cv_email(
    candidate_id: int,
) -> str:
    document = fetch_one(
        """
        SELECT parsed_json, extracted_text
        FROM candidate_documents
        WHERE candidate_id = ?
        ORDER BY created_at DESC, id DESC
        LIMIT 1
        """,
        (candidate_id,),
    )

    if not document:
        return ""

    try:
        parsed = json.loads(
            document.get("parsed_json") or "{}"
        )
    except json.JSONDecodeError:
        parsed = {}

    parsed_email = normalize_email(
        parsed.get("email") or ""
    )
    if validate_email(parsed_email):
        return parsed_email

    text = document.get("extracted_text") or ""
    match = re.search(
        r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}",
        text,
    )
    if match:
        detected = normalize_email(
            match.group(0)
        )
        if validate_email(detected):
            return detected

    return ""


def get_application_email_details(
    application_id: int,
) -> dict:
    details = fetch_one(
        """
        SELECT
            applications.id AS application_id,
            candidates.id AS candidate_id,
            candidates.source AS candidate_source,
            users.id AS candidate_user_id,
            users.full_name AS candidate_name,
            users.email AS account_email,
            users.active AS user_active,
            jobs.id AS job_id,
            jobs.title AS job_title,
            jobs.company_id,
            companies.name AS company_name
        FROM applications
        JOIN candidates
            ON candidates.id = applications.candidate_id
        JOIN users
            ON users.id = candidates.user_id
        JOIN jobs
            ON jobs.id = applications.job_id
        JOIN companies
            ON companies.id = jobs.company_id
        WHERE applications.id = ?
        """,
        (application_id,),
    )

    if not details:
        raise ValueError(
            "No se encontró la postulación."
        )

    cv_email = parse_cv_email(
        details["candidate_id"]
    )
    account_email = normalize_email(
        details.get("account_email") or ""
    )
    account_is_real = (
        validate_email(account_email)
        and not account_email.endswith(
            "@alba.local"
        )
    )

    if not bool(details.get("user_active")) and cv_email:
        recommended_email = cv_email
        email_source = "Detectado en el CV"
    elif account_is_real:
        recommended_email = account_email
        email_source = "Cuenta del candidato"
    elif cv_email:
        recommended_email = cv_email
        email_source = "Detectado en el CV"
    else:
        recommended_email = ""
        email_source = "No detectado"

    details["cv_email"] = cv_email
    details["recommended_email"] = recommended_email
    details["email_source"] = email_source
    return details


def update_candidate_email_for_invitation(
    candidate_user_id: int,
    recipient_email: str,
) -> None:
    recipient_email = normalize_email(
        recipient_email
    )

    if not validate_email(recipient_email):
        raise ValueError(
            "Ingresá un correo electrónico válido."
        )

    conflict = fetch_one(
        """
        SELECT id
        FROM users
        WHERE email = ?
          AND id <> ?
        """,
        (
            recipient_email,
            candidate_user_id,
        ),
    )
    if conflict:
        raise ValueError(
            "Ese correo ya pertenece a otra cuenta."
        )

    execute(
        """
        UPDATE users
        SET email = ?
        WHERE id = ?
        """,
        (
            recipient_email,
            candidate_user_id,
        ),
    )


def invitation_expiration(
    due_date: str,
) -> datetime:
    if due_date:
        try:
            due_day = date.fromisoformat(
                due_date[:10]
            )
        except ValueError as exc:
            raise ValueError(
                "La fecha límite debe tener formato AAAA-MM-DD."
            ) from exc

        expiration = datetime.combine(
            due_day,
            datetime_time(
                23,
                59,
                59,
            ),
        )
        if expiration <= datetime.now():
            raise ValueError(
                "La fecha límite debe ser futura."
            )
        return expiration

    return datetime.now() + timedelta(
        days=INVITATION_DEFAULT_DAYS
    )


def invitation_link(
    token: str,
) -> str:
    base_url = smtp_configuration()[
        "app_base_url"
    ]
    if not base_url:
        raise RuntimeError(
            "Falta APP_BASE_URL en Streamlit Secrets."
        )

    return (
        f"{base_url}/?"
        + urlencode({"invite": token})
    )


def create_interview_invitation(
    interview_id: int,
    recipient_email: str,
) -> dict:
    details = fetch_one(
        """
        SELECT
            interviews.id AS interview_id,
            interviews.due_date,
            users.id AS candidate_user_id
        FROM interviews
        JOIN applications
            ON applications.id = interviews.application_id
        JOIN candidates
            ON candidates.id = applications.candidate_id
        JOIN users
            ON users.id = candidates.user_id
        WHERE interviews.id = ?
        """,
        (interview_id,),
    )

    if not details:
        raise ValueError(
            "No se encontró la entrevista."
        )

    token = secrets.token_urlsafe(32)
    token_hash = hashlib.sha256(
        token.encode("utf-8")
    ).hexdigest()
    expires_at = invitation_expiration(
        details.get("due_date") or ""
    ).isoformat(timespec="seconds")

    existing = fetch_one(
        """
        SELECT id
        FROM interview_invitations
        WHERE interview_id = ?
        """,
        (interview_id,),
    )

    params = (
        details["candidate_user_id"],
        normalize_email(recipient_email),
        token_hash,
        expires_at,
        "",
        "",
        "PENDIENTE",
        "",
        now_iso(),
        interview_id,
    )

    if existing:
        execute(
            """
            UPDATE interview_invitations
            SET candidate_user_id = ?,
                recipient_email = ?,
                token_hash = ?,
                expires_at = ?,
                used_at = ?,
                sent_at = ?,
                send_status = ?,
                send_error = ?,
                updated_at = ?
            WHERE interview_id = ?
            """,
            params,
        )
    else:
        execute(
            """
            INSERT INTO interview_invitations(
                candidate_user_id,
                recipient_email,
                token_hash,
                expires_at,
                used_at,
                sent_at,
                send_status,
                send_error,
                updated_at,
                interview_id
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            params,
        )

    return {
        "token": token,
        "token_hash": token_hash,
        "expires_at": expires_at,
        "url": invitation_link(token),
    }


def send_smtp_message(
    recipient_email: str,
    subject: str,
    text_body: str,
    html_body: str,
) -> None:
    config = smtp_configuration()

    if not config["host"]:
        raise RuntimeError(
            "Falta SMTP_HOST en Streamlit Secrets."
        )
    if not config["from_email"]:
        raise RuntimeError(
            "Falta SMTP_FROM_EMAIL o SMTP_USERNAME."
        )

    message = EmailMessage()
    message["Subject"] = subject
    message["From"] = (
        f"{config['from_name']} "
        f"<{config['from_email']}>"
    )
    message["To"] = recipient_email

    if config["reply_to"]:
        message["Reply-To"] = config[
            "reply_to"
        ]

    message.set_content(text_body)
    message.add_alternative(
        html_body,
        subtype="html",
    )

    context = ssl.create_default_context()

    if config["use_ssl"]:
        with smtplib.SMTP_SSL(
            config["host"],
            config["port"],
            timeout=30,
            context=context,
        ) as server:
            if config["username"]:
                server.login(
                    config["username"],
                    config["password"],
                )
            server.send_message(message)
    else:
        with smtplib.SMTP(
            config["host"],
            config["port"],
            timeout=30,
        ) as server:
            server.ehlo()

            if config["use_tls"]:
                server.starttls(
                    context=context
                )
                server.ehlo()

            if config["username"]:
                server.login(
                    config["username"],
                    config["password"],
                )

            server.send_message(message)


def build_interview_invitation_email(
    details: dict,
    invitation: dict,
    intro_message: str,
) -> tuple[str, str, str]:
    candidate_name = (
        details.get("candidate_name")
        or "Candidato/a"
    )
    company_name = (
        details.get("company_name")
        or "la empresa"
    )
    job_title = (
        details.get("job_title")
        or "la búsqueda"
    )
    due_date = (
        details.get("due_date")
        or "Sin fecha límite informada"
    )
    url = invitation["url"]

    subject = (
        f"Entrevista con ALBA – "
        f"{job_title} | {company_name}"
    )

    text_body = f"""Hola {candidate_name}:

{company_name} te invita a completar una entrevista estructurada con ALBA para la posición {job_title}.

{intro_message.strip()}

Fecha límite: {due_date}

Ingresá desde este enlace seguro:
{url}

Si todavía no tenés una cuenta activa, el enlace te permitirá crear tu contraseña.

Por seguridad, el enlace vence el {invitation['expires_at']} y no debe compartirse.

Saludos,
{company_name}
"""

    html_body = f"""
    <html>
      <body style="font-family:Arial,sans-serif;line-height:1.5;color:#222;">
        <div style="max-width:640px;margin:auto;padding:24px;">
          <h2>Entrevista con ALBA</h2>
          <p>Hola <strong>{html.escape(candidate_name)}</strong>:</p>
          <p>
            <strong>{html.escape(company_name)}</strong> te invita a completar
            una entrevista estructurada con ALBA para la posición
            <strong>{html.escape(job_title)}</strong>.
          </p>
          <p>{html.escape(intro_message.strip())}</p>
          <p><strong>Fecha límite:</strong> {html.escape(due_date)}</p>
          <p style="margin:28px 0;">
            <a href="{html.escape(url)}"
               style="background:#111827;color:white;padding:14px 22px;
                      text-decoration:none;border-radius:8px;display:inline-block;">
              Ingresar a la entrevista
            </a>
          </p>
          <p>
            Si todavía no tenés una cuenta activa, el enlace te permitirá
            crear tu contraseña.
          </p>
          <p style="font-size:13px;color:#666;">
            El enlace vence el {html.escape(invitation['expires_at'])}
            y no debe compartirse.
          </p>
        </div>
      </body>
    </html>
    """

    return subject, text_body, html_body


def mark_invitation_delivery(
    interview_id: int,
    status: str,
    error_text: str = "",
) -> None:
    execute(
        """
        UPDATE interview_invitations
        SET send_status = ?,
            sent_at = ?,
            send_error = ?,
            updated_at = ?
        WHERE interview_id = ?
        """,
        (
            status,
            now_iso() if status == "ENVIADA" else "",
            error_text[:2000],
            now_iso(),
            interview_id,
        ),
    )


def send_interview_invitation(
    user: dict,
    interview_id: int,
    recipient_email: str,
) -> dict:
    details = fetch_one(
        """
        SELECT
            interviews.id AS interview_id,
            interviews.application_id,
            interviews.due_date,
            interviews.intro_message,
            users.full_name AS candidate_name,
            jobs.title AS job_title,
            companies.name AS company_name
        FROM interviews
        JOIN applications
            ON applications.id = interviews.application_id
        JOIN candidates
            ON candidates.id = applications.candidate_id
        JOIN users
            ON users.id = candidates.user_id
        JOIN jobs
            ON jobs.id = applications.job_id
        JOIN companies
            ON companies.id = jobs.company_id
        WHERE interviews.id = ?
          AND jobs.company_id = ?
        """,
        (
            interview_id,
            user["company_id"],
        ),
    )

    if not details:
        raise ValueError(
            "La entrevista no pertenece a la empresa."
        )

    try:
        invitation = create_interview_invitation(
            interview_id,
            recipient_email,
        )
    except Exception as exc:
        return {
            "status": "ERROR",
            "error": str(exc),
            "recipient_email": recipient_email,
            "url": "",
            "expires_at": "",
        }

    subject, text_body, html_body = (
        build_interview_invitation_email(
            details,
            invitation,
            details.get("intro_message")
            or (
                "Te invitamos a completar una entrevista "
                "estructurada con ALBA."
            ),
        )
    )

    try:
        send_smtp_message(
            recipient_email,
            subject,
            text_body,
            html_body,
        )
        mark_invitation_delivery(
            interview_id,
            "ENVIADA",
        )
        delivery_status = "ENVIADA"
        delivery_error = ""

        log_event(
            user["company_id"],
            user["id"],
            "SEND_INTERVIEW_INVITATION",
            "interview",
            interview_id,
            {
                "recipient_email": recipient_email,
                "status": delivery_status,
            },
        )

    except Exception as exc:
        delivery_status = "ERROR"
        delivery_error = str(exc)
        mark_invitation_delivery(
            interview_id,
            delivery_status,
            delivery_error,
        )

    return {
        **invitation,
        "status": delivery_status,
        "error": delivery_error,
        "recipient_email": recipient_email,
    }


def get_interview_invitation_status(
    interview_id: int,
) -> dict | None:
    return fetch_one(
        """
        SELECT *
        FROM interview_invitations
        WHERE interview_id = ?
        """,
        (interview_id,),
    )


def lookup_interview_invitation(
    token: str,
) -> dict | None:
    token_hash = hashlib.sha256(
        token.encode("utf-8")
    ).hexdigest()

    return fetch_one(
        """
        SELECT
            interview_invitations.*,
            interviews.status AS interview_status,
            interviews.due_date,
            interview_templates.name AS template_name,
            applications.id AS application_id,
            candidates.id AS candidate_id,
            users.full_name AS candidate_name,
            users.email AS account_email,
            users.active AS user_active,
            jobs.title AS job_title,
            companies.name AS company_name
        FROM interview_invitations
        JOIN interviews
            ON interviews.id = interview_invitations.interview_id
        JOIN interview_templates
            ON interview_templates.id = interviews.template_id
        JOIN applications
            ON applications.id = interviews.application_id
        JOIN candidates
            ON candidates.id = applications.candidate_id
        JOIN users
            ON users.id = candidates.user_id
        JOIN jobs
            ON jobs.id = applications.job_id
        JOIN companies
            ON companies.id = jobs.company_id
        WHERE interview_invitations.token_hash = ?
        """,
        (token_hash,),
    )


def clear_invitation_query() -> None:
    try:
        st.query_params.clear()
    except Exception:
        try:
            st.query_params.from_dict({})
        except Exception:
            pass


def complete_invitation_login(
    invitation: dict,
    password: str,
) -> bool:
    user = authenticate(
        invitation["account_email"],
        password,
    )

    if not user:
        return False

    if user["id"] != invitation[
        "candidate_user_id"
    ]:
        return False

    execute(
        """
        UPDATE interview_invitations
        SET used_at = ?,
            updated_at = ?
        WHERE id = ?
        """,
        (
            now_iso(),
            now_iso(),
            invitation["id"],
        ),
    )

    st.session_state["auth_user"] = user
    st.session_state[
        "candidate_portal_menu"
    ] = "Mis entrevistas"
    clear_invitation_query()
    return True


def render_interview_invitation_access(
    token: str,
) -> None:
    invitation = lookup_interview_invitation(
        token,
    )

    st.title("Invitación a entrevista con ALBA")

    if not invitation:
        st.error(
            "El enlace es inválido o ya no existe."
        )
        return

    try:
        expiration = datetime.fromisoformat(
            invitation["expires_at"]
        )
    except (TypeError, ValueError):
        st.error(
            "La invitación tiene una fecha de vencimiento inválida."
        )
        return

    if expiration < datetime.now():
        st.error(
            "El enlace de invitación venció. "
            "Solicitá a RR. HH. un nuevo envío."
        )
        return

    if invitation.get("used_at"):
        st.warning(
            "Este enlace ya fue utilizado. "
            "Ingresá normalmente con tu correo y contraseña."
        )
        return

    st.subheader(
        f"{invitation['company_name']} · "
        f"{invitation['job_title']}"
    )
    st.write(
        f"**Candidato/a:** "
        f"{invitation['candidate_name']}"
    )
    st.write(
        f"**Plantilla:** "
        f"{invitation['template_name']}"
    )
    st.write(
        f"**Correo:** "
        f"{invitation['recipient_email']}"
    )
    st.caption(
        f"El enlace vence el "
        f"{invitation['expires_at']}."
    )

    session_user = current_user()
    if (
        session_user
        and session_user["id"]
        == invitation["candidate_user_id"]
    ):
        execute(
            """
            UPDATE interview_invitations
            SET used_at = ?,
                updated_at = ?
            WHERE id = ?
            """,
            (
                now_iso(),
                now_iso(),
                invitation["id"],
            ),
        )
        st.session_state[
            "candidate_portal_menu"
        ] = "Mis entrevistas"
        clear_invitation_query()
        st.rerun()

    if bool(invitation.get("user_active")):
        st.info(
            "Ya tenés una cuenta activa. "
            "Ingresá tu contraseña para abrir la entrevista."
        )

        with st.form("invitation_existing_user_login"):
            st.text_input(
                "Correo",
                value=invitation[
                    "account_email"
                ],
                disabled=True,
            )
            password = st.text_input(
                "Contraseña",
                type="password",
            )
            submitted = st.form_submit_button(
                "Ingresar a la entrevista",
                type="primary",
                use_container_width=True,
            )

            if submitted:
                if complete_invitation_login(
                    invitation,
                    password,
                ):
                    st.rerun()
                else:
                    st.error(
                        "La contraseña es incorrecta."
                    )

    else:
        st.info(
            "Creá una contraseña para activar tu cuenta "
            "y acceder a la entrevista."
        )

        with st.form("invitation_activate_candidate"):
            password = st.text_input(
                "Nueva contraseña",
                type="password",
            )
            repeat = st.text_input(
                "Repetir contraseña",
                type="password",
            )
            consent = st.checkbox(
                "Acepto el tratamiento de mis datos "
                "para este proceso de selección."
            )
            submitted = st.form_submit_button(
                "Activar cuenta e ingresar",
                type="primary",
                use_container_width=True,
            )

            if submitted:
                try:
                    if password != repeat:
                        raise ValueError(
                            "Las contraseñas no coinciden."
                        )

                    password_error = validate_password(
                        password
                    )
                    if password_error:
                        raise ValueError(
                            password_error
                        )

                    if not consent:
                        raise ValueError(
                            "Tenés que aceptar el consentimiento."
                        )

                    password_hash, salt = hash_password(
                        password
                    )

                    execute(
                        """
                        UPDATE users
                        SET password_hash = ?,
                            password_salt = ?,
                            active = 1,
                            email = ?
                        WHERE id = ?
                        """,
                        (
                            password_hash,
                            salt,
                            invitation[
                                "recipient_email"
                            ],
                            invitation[
                                "candidate_user_id"
                            ],
                        ),
                    )

                    refreshed = lookup_interview_invitation(
                        token
                    )
                    user = authenticate(
                        refreshed["account_email"],
                        password,
                    )
                    if not user:
                        raise RuntimeError(
                            "La cuenta se activó, pero no "
                            "se pudo iniciar sesión."
                        )

                    execute(
                        """
                        UPDATE interview_invitations
                        SET used_at = ?,
                            updated_at = ?
                        WHERE id = ?
                        """,
                        (
                            now_iso(),
                            now_iso(),
                            invitation["id"],
                        ),
                    )

                    st.session_state[
                        "auth_user"
                    ] = user
                    st.session_state[
                        "candidate_portal_menu"
                    ] = "Mis entrevistas"
                    clear_invitation_query()
                    st.rerun()

                except Exception as exc:
                    st.error(str(exc))


def render_email_configuration_status() -> None:
    config = smtp_configuration()

    c1, c2 = st.columns(2)
    c1.metric(
        "Servidor de correo",
        (
            "Configurado"
            if config["host"]
            and config["from_email"]
            else "Pendiente"
        ),
    )
    c2.metric(
        "URL pública de ALBA",
        (
            "Configurada"
            if config["app_base_url"]
            else "Pendiente"
        ),
    )

    st.caption(INVITATION_EMAIL_NOTICE)

    with st.expander(
        "Ejemplo de configuración en Streamlit Secrets",
        expanded=not smtp_configured(),
    ):
        example = (
            'APP_BASE_URL = "https://tu-app.streamlit.app"\n\n'
            'SMTP_HOST = "smtp.gmail.com"\n'
            'SMTP_PORT = "465"\n'
            'SMTP_USERNAME = "tu-cuenta@gmail.com"\n'
            'SMTP_PASSWORD = "CONTRASEÑA_DE_APLICACIÓN"\n'
            'SMTP_FROM_EMAIL = "tu-cuenta@gmail.com"\n'
            'SMTP_FROM_NAME = "Selección · Tu empresa"\n'
            'SMTP_USE_SSL = "true"\n'
            'SMTP_USE_TLS = "false"'
        )
        st.code(
            example,
            language="toml",
        )
        st.caption(
            "Para Gmail se debe utilizar una contraseña "
            "de aplicación, no la contraseña habitual. "
            "Nunca subas estos datos a GitHub."
        )


def validate_audio_file(uploaded_file) -> tuple[bytes, str, str]:
    raw = uploaded_file.getvalue()
    if not raw:
        raise ValueError("El archivo de audio está vacío.")

    if len(raw) > MAX_AUDIO_SIZE_BYTES:
        raise ValueError(
            "El audio supera el máximo permitido de 10 MB."
        )

    extension = Path(uploaded_file.name).suffix.lower()
    if extension not in {".mp3", ".wav", ".m4a", ".ogg", ".webm"}:
        raise ValueError(
            "El audio debe estar en MP3, WAV, M4A, OGG o WEBM."
        )

    return (
        raw,
        uploaded_file.type or "application/octet-stream",
        uploaded_file.name,
    )


def create_default_interview_templates(
    company_id: int,
    user_id: int,
) -> int:
    created = 0

    for template_name, template_data in (
        DEFAULT_INTERVIEW_TEMPLATES.items()
    ):
        existing = fetch_one(
            """
            SELECT id
            FROM interview_templates
            WHERE company_id = ? AND name = ?
            """,
            (company_id, template_name),
        )
        if existing:
            continue

        template_id = execute(
            """
            INSERT INTO interview_templates(
                company_id,
                name,
                description,
                category,
                active,
                created_by,
                created_at,
                updated_at
            )
            VALUES (?, ?, ?, ?, 1, ?, ?, ?)
            """,
            (
                company_id,
                template_name,
                template_data["description"],
                template_data["category"],
                user_id,
                now_iso(),
                now_iso(),
            ),
        )

        for position, question in enumerate(
            template_data["questions"],
            start=1,
        ):
            execute(
                """
                INSERT INTO interview_questions(
                    template_id,
                    position,
                    competency,
                    question_text,
                    indicators_text,
                    max_score,
                    required,
                    response_mode,
                    active,
                    created_at,
                    updated_at
                )
                VALUES (?, ?, ?, ?, ?, 5, 1, 'TEXTO_O_AUDIO', 1, ?, ?)
                """,
                (
                    template_id,
                    position,
                    question["competency"],
                    question["question"],
                    question["indicators"],
                    now_iso(),
                    now_iso(),
                ),
            )

        created += 1

    log_event(
        company_id,
        user_id,
        "CREATE_DEFAULT_INTERVIEW_TEMPLATES",
        "interview_template",
        None,
        {"created": created},
    )
    return created


def get_interview_templates(
    company_id: int,
    active_only: bool = False,
) -> list[dict]:
    query = """
        SELECT *
        FROM interview_templates
        WHERE company_id = ?
    """
    params: tuple = (company_id,)

    if active_only:
        query += " AND active = 1"

    query += " ORDER BY name"
    return fetch_all(query, params)


def get_template_questions(
    template_id: int,
    active_only: bool = False,
) -> list[dict]:
    query = """
        SELECT *
        FROM interview_questions
        WHERE template_id = ?
    """
    params: tuple = (template_id,)

    if active_only:
        query += " AND active = 1"

    query += " ORDER BY position, id"
    return fetch_all(query, params)


def assign_interview(
    user: dict,
    application_id: int,
    template_id: int,
    due_date: str,
    intro_message: str,
) -> int:
    application = fetch_one(
        """
        SELECT
            applications.id,
            jobs.company_id
        FROM applications
        JOIN jobs
            ON jobs.id = applications.job_id
        WHERE applications.id = ?
          AND jobs.company_id = ?
        """,
        (application_id, user["company_id"]),
    )
    if not application:
        raise ValueError(
            "La postulación no pertenece a la empresa."
        )

    template = fetch_one(
        """
        SELECT id
        FROM interview_templates
        WHERE id = ?
          AND company_id = ?
          AND active = 1
        """,
        (template_id, user["company_id"]),
    )
    if not template:
        raise ValueError(
            "La plantilla no pertenece a la empresa o está inactiva."
        )

    existing = fetch_one(
        """
        SELECT id
        FROM interviews
        WHERE application_id = ?
          AND template_id = ?
        """,
        (application_id, template_id),
    )

    if existing:
        execute(
            """
            UPDATE interviews
            SET status = 'INVITADA',
                due_date = ?,
                intro_message = ?,
                assigned_by = ?,
                assigned_at = ?,
                updated_at = ?
            WHERE id = ?
            """,
            (
                due_date or "",
                intro_message.strip(),
                user["id"],
                now_iso(),
                now_iso(),
                existing["id"],
            ),
        )
        interview_id = existing["id"]
    else:
        interview_id = execute(
            """
            INSERT INTO interviews(
                application_id,
                template_id,
                status,
                due_date,
                intro_message,
                assigned_by,
                assigned_at,
                updated_at
            )
            VALUES (?, ?, 'INVITADA', ?, ?, ?, ?, ?)
            """,
            (
                application_id,
                template_id,
                due_date or "",
                intro_message.strip(),
                user["id"],
                now_iso(),
                now_iso(),
            ),
        )

    execute(
        """
        UPDATE applications
        SET status = 'ENTREVISTA',
            updated_at = ?
        WHERE id = ?
        """,
        (now_iso(), application_id),
    )

    log_event(
        user["company_id"],
        user["id"],
        "ASSIGN_INTERVIEW",
        "interview",
        interview_id,
        {
            "application_id": application_id,
            "template_id": template_id,
            "due_date": due_date,
        },
    )

    return interview_id


def evaluate_interview_answer(
    question: dict,
    answer_text: str,
) -> dict:
    answer = normalize_match_text(answer_text or "")
    maximum = float(question.get("max_score") or 5)

    if not answer:
        return {
            "score": None,
            "maximum": maximum,
            "summary": (
                "La respuesta no tiene texto o transcripción. "
                "El audio requiere revisión humana."
            ),
            "matched_indicators": [],
            "missing_indicators": extract_scoring_terms(
                question.get("indicators_text") or ""
            ),
            "star_elements": [],
            "word_count": 0,
            "strengths": [],
            "gaps": [
                "No existe texto suficiente para una asistencia automática."
            ],
        }

    indicators = extract_scoring_terms(
        question.get("indicators_text") or ""
    )
    matched = [
        term for term in indicators if term in answer
    ]
    missing = [
        term for term in indicators if term not in answer
    ]

    indicator_coverage = (
        len(matched) / len(indicators)
        if indicators
        else 0.5
    )

    star_markers = {
        "Situación": [
            "situacion",
            "contexto",
            "momento",
            "ocasion",
            "proyecto",
        ],
        "Tarea": [
            "responsabilidad",
            "objetivo",
            "tarea",
            "debia",
            "tenia que",
        ],
        "Acción": [
            "hice",
            "realice",
            "implemente",
            "propuse",
            "decidi",
            "organice",
            "analice",
        ],
        "Resultado": [
            "resultado",
            "logre",
            "mejoro",
            "redujo",
            "aumento",
            "impacto",
            "aprendi",
        ],
    }

    star_elements = [
        label
        for label, terms in star_markers.items()
        if any(term in answer for term in terms)
    ]
    star_coverage = len(star_elements) / 4

    word_count = len(answer.split())
    if word_count >= 120:
        depth_coverage = 1.0
    elif word_count >= 70:
        depth_coverage = 0.8
    elif word_count >= 40:
        depth_coverage = 0.6
    elif word_count >= 20:
        depth_coverage = 0.4
    else:
        depth_coverage = 0.2

    raw_percentage = (
        indicator_coverage * 0.60
        + star_coverage * 0.25
        + depth_coverage * 0.15
    )
    score = round(maximum * raw_percentage, 1)

    strengths = []
    gaps = []

    if matched:
        strengths.append(
            "Incluye evidencias relacionadas con: "
            + ", ".join(matched[:10])
            + "."
        )
    if len(star_elements) >= 3:
        strengths.append(
            "La respuesta presenta una estructura conductual clara: "
            + ", ".join(star_elements)
            + "."
        )
    if word_count >= 70:
        strengths.append(
            "La respuesta tiene un nivel de desarrollo suficiente."
        )

    if missing:
        gaps.append(
            "No se identificaron evidencias claras de: "
            + ", ".join(missing[:10])
            + "."
        )
    missing_star = [
        item
        for item in ["Situación", "Tarea", "Acción", "Resultado"]
        if item not in star_elements
    ]
    if missing_star:
        gaps.append(
            "Sería conveniente profundizar: "
            + ", ".join(missing_star)
            + "."
        )
    if word_count < 40:
        gaps.append(
            "La respuesta es breve y necesita ejemplos más concretos."
        )

    return {
        "score": score,
        "maximum": maximum,
        "summary": (
            f"Se detectaron {len(matched)} de {len(indicators)} "
            f"indicadores y {len(star_elements)} de 4 componentes STAR."
        ),
        "matched_indicators": matched,
        "missing_indicators": missing,
        "star_elements": star_elements,
        "word_count": word_count,
        "strengths": strengths,
        "gaps": gaps,
    }


def save_automatic_interview_evaluation(
    interview_id: int,
) -> dict:
    questions = fetch_all(
        """
        SELECT
            interview_questions.*,
            interview_responses.answer_text,
            interview_responses.audio_blob
        FROM interviews
        JOIN interview_questions
            ON interview_questions.template_id = interviews.template_id
        LEFT JOIN interview_responses
            ON interview_responses.interview_id = interviews.id
           AND interview_responses.question_id = interview_questions.id
        WHERE interviews.id = ?
          AND interview_questions.active = 1
        ORDER BY interview_questions.position, interview_questions.id
        """,
        (interview_id,),
    )

    question_reports = []
    obtained = 0.0
    possible = 0.0
    pending_audio = 0

    for question in questions:
        report = evaluate_interview_answer(
            question,
            question.get("answer_text") or "",
        )
        question_reports.append(
            {
                "question_id": question["id"],
                "competency": question.get("competency") or "",
                "question_text": question["question_text"],
                **report,
            }
        )

        if report["score"] is not None:
            obtained += float(report["score"])
            possible += float(report["maximum"])
        elif question.get("audio_blob"):
            pending_audio += 1

        existing = fetch_one(
            """
            SELECT id
            FROM interview_evaluations
            WHERE interview_id = ?
              AND question_id = ?
              AND evaluator_type = 'AUTO'
            """,
            (interview_id, question["id"]),
        )

        params = (
            report["score"],
            report["summary"],
            json.dumps(
                report["strengths"],
                ensure_ascii=False,
            ),
            json.dumps(
                report["gaps"],
                ensure_ascii=False,
            ),
            json.dumps(report, ensure_ascii=False),
            now_iso(),
            interview_id,
            question["id"],
        )

        if existing:
            execute(
                """
                UPDATE interview_evaluations
                SET score = ?,
                    evidence_text = ?,
                    strengths_text = ?,
                    gaps_text = ?,
                    notes = ?,
                    evaluated_at = ?
                WHERE interview_id = ?
                  AND question_id = ?
                  AND evaluator_type = 'AUTO'
                """,
                params,
            )
        else:
            execute(
                """
                INSERT INTO interview_evaluations(
                    score,
                    evidence_text,
                    strengths_text,
                    gaps_text,
                    notes,
                    evaluated_at,
                    evaluator_type,
                    interview_id,
                    question_id
                )
                VALUES (?, ?, ?, ?, ?, ?, 'AUTO', ?, ?)
                """,
                params,
            )

    percentage = (
        round((obtained / possible) * 100, 1)
        if possible > 0
        else None
    )

    if percentage is None:
        recommendation = "REVISIÓN HUMANA"
    elif pending_audio > 0:
        recommendation = "REVISIÓN HUMANA"
    elif percentage >= 75:
        recommendation = "AVANZA"
    elif percentage >= 50:
        recommendation = "REVISAR"
    else:
        recommendation = "NO AVANZA"

    strengths = []
    gaps = []

    for item in question_reports:
        strengths.extend(item.get("strengths") or [])
        gaps.extend(item.get("gaps") or [])

    report = {
        "score": percentage,
        "recommendation": recommendation,
        "pending_audio_responses": pending_audio,
        "questions": question_reports,
        "strengths": strengths[:20],
        "gaps": gaps[:20],
        "notice": INTERVIEW_NOTICE,
    }

    execute(
        """
        UPDATE interviews
        SET auto_score = ?,
            auto_recommendation = ?,
            auto_report_json = ?,
            updated_at = ?
        WHERE id = ?
        """,
        (
            percentage,
            recommendation,
            json.dumps(report, ensure_ascii=False),
            now_iso(),
            interview_id,
        ),
    )

    return report


def render_interview_auto_report(report: dict) -> None:
    c1, c2, c3 = st.columns(3)
    c1.metric(
        "Puntaje asistido",
        (
            "Pendiente"
            if report.get("score") is None
            else f"{float(report['score']):.1f}%"
        ),
    )
    c2.metric(
        "Recomendación asistida",
        report.get("recommendation") or "PENDIENTE",
    )
    c3.metric(
        "Audios pendientes",
        int(report.get("pending_audio_responses") or 0),
    )

    st.caption(report.get("notice") or INTERVIEW_NOTICE)

    strengths = report.get("strengths") or []
    gaps = report.get("gaps") or []

    a1, a2 = st.columns(2)
    with a1:
        st.markdown("**Fortalezas detectadas**")
        if strengths:
            for item in strengths:
                st.write(f"• {item}")
        else:
            st.write("No hay fortalezas suficientes para mostrar.")

    with a2:
        st.markdown("**Aspectos para profundizar**")
        if gaps:
            for item in gaps:
                st.write(f"• {item}")
        else:
            st.write("No se detectaron brechas relevantes.")

    st.markdown("#### Explicación por pregunta")
    for item in report.get("questions") or []:
        with st.expander(
            f"{item.get('competency') or 'Competencia'} · "
            f"{'Pendiente' if item.get('score') is None else str(item.get('score')) + '/' + str(item.get('maximum'))}"
        ):
            st.write(f"**Pregunta:** {item.get('question_text')}")
            st.write(f"**Análisis:** {item.get('summary')}")
            st.write(
                "**Indicadores encontrados:** "
                + (
                    ", ".join(item.get("matched_indicators") or [])
                    or "Ninguno"
                )
            )
            st.write(
                "**Indicadores no encontrados:** "
                + (
                    ", ".join(item.get("missing_indicators") or [])
                    or "Ninguno"
                )
            )
            st.write(
                "**Componentes STAR detectados:** "
                + (
                    ", ".join(item.get("star_elements") or [])
                    or "Ninguno"
                )
            )
            st.write(
                f"**Extensión de la respuesta:** "
                f"{item.get('word_count') or 0} palabras"
            )


def render_template_management(user: dict) -> None:
    templates = get_interview_templates(
        user["company_id"],
        active_only=False,
    )

    if not templates:
        st.info(
            "Todavía no hay plantillas. Podés crear el banco "
            "predeterminado para comenzar."
        )

    if st.button(
        "Crear plantillas predeterminadas",
        key="create_default_interview_templates",
    ):
        created = create_default_interview_templates(
            user["company_id"],
            user["id"],
        )
        st.success(f"Plantillas creadas: {created}.")
        st.rerun()

    with st.expander("Crear plantilla personalizada"):
        with st.form("create_interview_template"):
            name = st.text_input("Nombre de la plantilla")
            description = st.text_area("Descripción")
            category = st.text_input(
                "Categoría",
                value="PERSONALIZADA",
            )
            submitted = st.form_submit_button(
                "Crear plantilla",
                type="primary",
            )

            if submitted:
                if not name.strip():
                    st.error("Ingresá un nombre.")
                else:
                    template_id = execute(
                        """
                        INSERT INTO interview_templates(
                            company_id,
                            name,
                            description,
                            category,
                            active,
                            created_by,
                            created_at,
                            updated_at
                        )
                        VALUES (?, ?, ?, ?, 1, ?, ?, ?)
                        """,
                        (
                            user["company_id"],
                            name.strip(),
                            description.strip(),
                            category.strip(),
                            user["id"],
                            now_iso(),
                            now_iso(),
                        ),
                    )
                    log_event(
                        user["company_id"],
                        user["id"],
                        "CREATE_INTERVIEW_TEMPLATE",
                        "interview_template",
                        template_id,
                        {"name": name},
                    )
                    st.success("Plantilla creada.")
                    st.rerun()

    templates = get_interview_templates(
        user["company_id"],
        active_only=False,
    )
    if not templates:
        return

    template_options = {
        f"#{item['id']} · {item['name']}": item
        for item in templates
    }
    selected_label = st.selectbox(
        "Plantilla",
        list(template_options.keys()),
        key="template_management_select",
    )
    selected_template = template_options[selected_label]

    with st.form(
        f"edit_template_{selected_template['id']}"
    ):
        template_name = st.text_input(
            "Nombre",
            value=selected_template["name"],
        )
        template_description = st.text_area(
            "Descripción",
            value=selected_template.get("description") or "",
        )
        template_category = st.text_input(
            "Categoría",
            value=selected_template.get("category") or "",
        )
        template_active = st.checkbox(
            "Plantilla activa",
            value=bool(selected_template.get("active")),
        )
        update_template = st.form_submit_button(
            "Guardar plantilla",
        )

        if update_template:
            execute(
                """
                UPDATE interview_templates
                SET name = ?,
                    description = ?,
                    category = ?,
                    active = ?,
                    updated_at = ?
                WHERE id = ?
                  AND company_id = ?
                """,
                (
                    template_name.strip(),
                    template_description.strip(),
                    template_category.strip(),
                    int(template_active),
                    now_iso(),
                    selected_template["id"],
                    user["company_id"],
                ),
            )
            st.success("Plantilla actualizada.")
            st.rerun()

    questions = get_template_questions(
        selected_template["id"],
        active_only=False,
    )

    st.markdown("#### Preguntas de la plantilla")
    if questions:
        st.dataframe(
            pd.DataFrame(
                [
                    {
                        "ID": item["id"],
                        "Orden": item["position"],
                        "Competencia": item.get("competency"),
                        "Pregunta": item["question_text"],
                        "Indicadores": item.get("indicators_text"),
                        "Puntaje máximo": item.get("max_score"),
                        "Obligatoria": (
                            "Sí" if item.get("required") else "No"
                        ),
                        "Activa": (
                            "Sí" if item.get("active") else "No"
                        ),
                    }
                    for item in questions
                ]
            ),
            use_container_width=True,
            hide_index=True,
        )

    with st.expander("Agregar una pregunta", expanded=not bool(questions)):
        with st.form(
            f"add_question_{selected_template['id']}"
        ):
            q1, q2 = st.columns(2)
            position = q1.number_input(
                "Orden",
                min_value=1,
                value=len(questions) + 1,
                step=1,
            )
            competency = q2.text_input("Competencia")
            question_text = st.text_area("Pregunta")
            indicators = st.text_area(
                "Indicadores esperados",
                help=(
                    "Separalos por coma. Se usan para explicar la "
                    "asistencia automática."
                ),
            )
            q3, q4, q5 = st.columns(3)
            max_score = q3.number_input(
                "Puntaje máximo",
                min_value=1.0,
                max_value=10.0,
                value=5.0,
                step=1.0,
            )
            required = q4.checkbox(
                "Obligatoria",
                value=True,
            )
            response_mode = q5.selectbox(
                "Modo de respuesta",
                [
                    "TEXTO_O_AUDIO",
                    "SOLO_TEXTO",
                    "SOLO_AUDIO",
                ],
            )
            add_question = st.form_submit_button(
                "Agregar pregunta",
                type="primary",
            )

            if add_question:
                if not question_text.strip():
                    st.error("Ingresá la pregunta.")
                else:
                    question_id = execute(
                        """
                        INSERT INTO interview_questions(
                            template_id,
                            position,
                            competency,
                            question_text,
                            indicators_text,
                            max_score,
                            required,
                            response_mode,
                            active,
                            created_at,
                            updated_at
                        )
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, 1, ?, ?)
                        """,
                        (
                            selected_template["id"],
                            int(position),
                            competency.strip(),
                            question_text.strip(),
                            indicators.strip(),
                            float(max_score),
                            int(required),
                            response_mode,
                            now_iso(),
                            now_iso(),
                        ),
                    )
                    log_event(
                        user["company_id"],
                        user["id"],
                        "CREATE_INTERVIEW_QUESTION",
                        "interview_question",
                        question_id,
                        {
                            "template_id": selected_template["id"],
                            "competency": competency,
                        },
                    )
                    st.success("Pregunta agregada.")
                    st.rerun()

    if questions:
        question_options = {
            f"#{item['id']} · {item.get('competency') or 'Sin competencia'}": item
            for item in questions
        }
        selected_question_label = st.selectbox(
            "Editar pregunta existente",
            list(question_options.keys()),
            key=f"edit_question_select_{selected_template['id']}",
        )
        selected_question = question_options[
            selected_question_label
        ]

        with st.form(
            f"edit_question_{selected_question['id']}"
        ):
            e1, e2 = st.columns(2)
            edit_position = e1.number_input(
                "Orden",
                min_value=1,
                value=int(selected_question["position"]),
                step=1,
            )
            edit_competency = e2.text_input(
                "Competencia",
                value=selected_question.get("competency") or "",
            )
            edit_question_text = st.text_area(
                "Pregunta",
                value=selected_question["question_text"],
            )
            edit_indicators = st.text_area(
                "Indicadores",
                value=selected_question.get("indicators_text") or "",
            )
            e3, e4, e5 = st.columns(3)
            edit_max_score = e3.number_input(
                "Puntaje máximo",
                min_value=1.0,
                max_value=10.0,
                value=float(
                    selected_question.get("max_score") or 5
                ),
                step=1.0,
            )
            edit_required = e4.checkbox(
                "Obligatoria",
                value=bool(selected_question.get("required")),
            )
            edit_active = e5.checkbox(
                "Activa",
                value=bool(selected_question.get("active")),
            )
            modes = [
                "TEXTO_O_AUDIO",
                "SOLO_TEXTO",
                "SOLO_AUDIO",
            ]
            current_mode = (
                selected_question.get("response_mode")
                if selected_question.get("response_mode") in modes
                else "TEXTO_O_AUDIO"
            )
            edit_mode = st.selectbox(
                "Modo de respuesta",
                modes,
                index=modes.index(current_mode),
            )
            save_question = st.form_submit_button(
                "Guardar pregunta",
            )

            if save_question:
                execute(
                    """
                    UPDATE interview_questions
                    SET position = ?,
                        competency = ?,
                        question_text = ?,
                        indicators_text = ?,
                        max_score = ?,
                        required = ?,
                        response_mode = ?,
                        active = ?,
                        updated_at = ?
                    WHERE id = ?
                    """,
                    (
                        int(edit_position),
                        edit_competency.strip(),
                        edit_question_text.strip(),
                        edit_indicators.strip(),
                        float(edit_max_score),
                        int(edit_required),
                        edit_mode,
                        int(edit_active),
                        now_iso(),
                        selected_question["id"],
                    ),
                )
                st.success("Pregunta actualizada.")
                st.rerun()


def render_interview_assignment(user: dict) -> None:
    render_email_configuration_status()

    templates = get_interview_templates(
        user["company_id"],
        active_only=True,
    )
    if not templates:
        st.info(
            "Creá o activá una plantilla antes de asignar entrevistas."
        )
        return

    applications = fetch_all(
        """
        SELECT
            applications.id,
            applications.status,
            candidates.id AS candidate_id,
            users.id AS candidate_user_id,
            users.full_name,
            users.email,
            users.active AS user_active,
            jobs.title AS job_title
        FROM applications
        JOIN candidates
            ON candidates.id = applications.candidate_id
        JOIN users
            ON users.id = candidates.user_id
        JOIN jobs
            ON jobs.id = applications.job_id
        WHERE jobs.company_id = ?
          AND applications.status <> 'RECHAZADA'
        ORDER BY applications.created_at DESC
        """,
        (user["company_id"],),
    )

    if not applications:
        st.info("No hay postulaciones disponibles.")
        return

    application_options = {
        (
            f"#{item['id']} · {item['full_name']} · "
            f"{item['job_title']} · {item['status']}"
        ): item
        for item in applications
    }
    template_options = {
        f"#{item['id']} · {item['name']}": item["id"]
        for item in templates
    }

    application_label = st.selectbox(
        "Candidato y búsqueda",
        list(application_options.keys()),
        key="interview_assignment_application",
    )
    selected_application = application_options[
        application_label
    ]
    email_details = get_application_email_details(
        selected_application["id"]
    )

    if email_details["cv_email"]:
        st.success(
            f"Correo detectado en el CV: "
            f"{email_details['cv_email']}"
        )
    else:
        st.warning(
            "No se detectó un correo válido en el CV. "
            "Ingresalo manualmente antes de enviar."
        )

    with st.form("assign_interview_form"):
        template_label = st.selectbox(
            "Plantilla",
            list(template_options.keys()),
        )
        recipient_email = st.text_input(
            "Correo de invitación",
            value=email_details[
                "recommended_email"
            ],
            help=(
                "ALBA intenta completarlo con el correo "
                "extraído del CV. Podés corregirlo."
            ),
        )
        st.caption(
            f"Fuente sugerida: "
            f"{email_details['email_source']}."
        )

        due_date = st.text_input(
            "Fecha límite (AAAA-MM-DD, opcional)",
        )
        intro_message = st.text_area(
            "Mensaje para el candidato",
            value=(
                "Te invitamos a completar una entrevista estructurada con "
                "ALBA. Podés responder por escrito o adjuntar un audio cuando "
                "la pregunta lo permita."
            ),
        )

        b1, b2 = st.columns(2)
        assign_and_send = b1.form_submit_button(
            "Asignar y enviar por email",
            type="primary",
            use_container_width=True,
        )
        assign_only = b2.form_submit_button(
            "Asignar sin enviar",
            use_container_width=True,
        )

        if assign_and_send or assign_only:
            try:
                if due_date:
                    date.fromisoformat(
                        due_date[:10]
                    )

                normalized_recipient = normalize_email(
                    recipient_email
                )

                if assign_and_send:
                    if not validate_email(
                        normalized_recipient
                    ):
                        raise ValueError(
                            "Ingresá un correo válido "
                            "para enviar la invitación."
                        )

                    update_candidate_email_for_invitation(
                        email_details[
                            "candidate_user_id"
                        ],
                        normalized_recipient,
                    )

                interview_id = assign_interview(
                    user,
                    selected_application["id"],
                    template_options[
                        template_label
                    ],
                    due_date,
                    intro_message,
                )

                result = {
                    "interview_id": interview_id,
                    "assigned": True,
                    "sent": False,
                    "recipient": (
                        normalized_recipient
                        if assign_and_send
                        else ""
                    ),
                    "url": "",
                    "error": "",
                }

                if assign_and_send:
                    delivery = send_interview_invitation(
                        user,
                        interview_id,
                        normalized_recipient,
                    )
                    result.update(
                        {
                            "sent": (
                                delivery["status"]
                                == "ENVIADA"
                            ),
                            "url": delivery.get(
                                "url",
                                "",
                            ),
                            "error": delivery.get(
                                "error",
                                "",
                            ),
                            "expires_at": delivery.get(
                                "expires_at",
                                "",
                            ),
                        }
                    )

                st.session_state[
                    "last_interview_assignment"
                ] = result

            except Exception as exc:
                st.error(str(exc))

    last_result = st.session_state.pop(
        "last_interview_assignment",
        None,
    )

    if last_result:
        st.success(
            f"Entrevista asignada correctamente. "
            f"ID: {last_result['interview_id']}."
        )

        if last_result["recipient"]:
            if last_result["sent"]:
                st.success(
                    f"Invitación enviada a "
                    f"{last_result['recipient']}."
                )
            else:
                st.warning(
                    "La entrevista fue asignada, pero el correo "
                    "no pudo enviarse."
                )
                if last_result.get("error"):
                    st.code(
                        last_result["error"]
                    )

            if last_result.get("url"):
                st.markdown(
                    "**Enlace seguro de respaldo**"
                )
                st.code(
                    last_result["url"]
                )
                st.caption(
                    "Podés copiar este enlace y enviarlo "
                    "manualmente si fuera necesario."
                )

def render_interview_monitoring(user: dict) -> None:
    interviews = fetch_all(
        """
        SELECT
            interviews.*,
            interview_templates.name AS template_name,
            applications.status AS application_status,
            users.full_name AS candidate_name,
            users.email,
            jobs.title AS job_title,
            companies.name AS company_name
        FROM interviews
        JOIN interview_templates
            ON interview_templates.id = interviews.template_id
        JOIN applications
            ON applications.id = interviews.application_id
        JOIN candidates
            ON candidates.id = applications.candidate_id
        JOIN users
            ON users.id = candidates.user_id
        JOIN jobs
            ON jobs.id = applications.job_id
        JOIN companies
            ON companies.id = jobs.company_id
        WHERE jobs.company_id = ?
        ORDER BY interviews.assigned_at DESC
        """,
        (user["company_id"],),
    )

    if not interviews:
        st.info("Todavía no hay entrevistas asignadas.")
        return

    st.dataframe(
        pd.DataFrame(
            [
                {
                    "ID": item["id"],
                    "Candidato": item["candidate_name"],
                    "Búsqueda": item["job_title"],
                    "Plantilla": item["template_name"],
                    "Estado": item["status"],
                    "Puntaje asistido": item.get("auto_score"),
                    "Recomendación asistida": (
                        item.get("auto_recommendation")
                    ),
                    "Puntaje humano": item.get("human_score"),
                    "Decisión final": (
                        item.get("final_recommendation")
                    ),
                    "Fecha límite": item.get("due_date"),
                }
                for item in interviews
            ]
        ),
        use_container_width=True,
        hide_index=True,
    )

    options = {
        (
            f"#{item['id']} · {item['candidate_name']} · "
            f"{item['job_title']} · {item['status']}"
        ): item
        for item in interviews
    }
    selected_label = st.selectbox(
        "Abrir entrevista",
        list(options.keys()),
        key="monitor_interview_select",
    )
    selected = options[selected_label]

    st.markdown(f"### {selected['candidate_name']}")
    st.write(f"**Búsqueda:** {selected['job_title']}")
    st.write(f"**Plantilla:** {selected['template_name']}")
    st.write(f"**Estado:** {selected['status']}")
    st.write(
        f"**Fecha límite:** "
        f"{selected.get('due_date') or 'Sin fecha límite'}"
    )

    invitation_status = get_interview_invitation_status(
        selected["id"]
    )
    if invitation_status:
        st.write(
            f"**Invitación por email:** "
            f"{invitation_status.get('send_status') or 'PENDIENTE'}"
        )
        st.write(
            f"**Destinatario:** "
            f"{invitation_status.get('recipient_email') or 'No informado'}"
        )
        if invitation_status.get("sent_at"):
            st.caption(
                f"Último envío: "
                f"{invitation_status['sent_at']}"
            )
        if invitation_status.get("send_error"):
            st.warning(
                invitation_status["send_error"]
            )
    else:
        st.caption(
            "Todavía no se envió una invitación por email."
        )

    if has_permission("manage_interviews"):
        if st.button(
            "Reenviar invitación por email",
            key=f"resend_interview_invitation_{selected['id']}",
        ):
            try:
                email_details = get_application_email_details(
                    selected["application_id"]
                )
                recipient = email_details[
                    "recommended_email"
                ]
                if not validate_email(recipient):
                    raise ValueError(
                        "No hay un correo válido. "
                        "Corregilo desde Asignar entrevistas."
                    )

                delivery = send_interview_invitation(
                    user,
                    selected["id"],
                    recipient,
                )
                if delivery["status"] == "ENVIADA":
                    st.success(
                        f"Invitación reenviada a {recipient}."
                    )
                else:
                    st.warning(
                        "No se pudo enviar el correo."
                    )
                    if delivery.get("error"):
                        st.code(delivery["error"])
                    if delivery.get("url"):
                        st.code(delivery["url"])
            except Exception as exc:
                st.error(str(exc))

    responses = fetch_all(
        """
        SELECT
            interview_questions.*,
            interview_responses.answer_text,
            interview_responses.audio_blob,
            interview_responses.audio_mime,
            interview_responses.audio_filename,
            interview_responses.submitted_at
        FROM interview_questions
        LEFT JOIN interview_responses
            ON interview_responses.question_id = interview_questions.id
           AND interview_responses.interview_id = ?
        WHERE interview_questions.template_id = ?
          AND interview_questions.active = 1
        ORDER BY interview_questions.position, interview_questions.id
        """,
        (
            selected["id"],
            selected["template_id"],
        ),
    )

    st.markdown("#### Respuestas")
    for response in responses:
        with st.expander(
            f"{response['position']}. "
            f"{response.get('competency') or 'Pregunta'}",
            expanded=False,
        ):
            st.write(f"**Pregunta:** {response['question_text']}")
            st.write(
                f"**Respuesta escrita/transcripción:** "
                f"{response.get('answer_text') or 'Sin texto'}"
            )
            if response.get("audio_blob"):
                st.audio(
                    response["audio_blob"],
                    format=response.get("audio_mime")
                    or "audio/mpeg",
                )
                st.caption(
                    response.get("audio_filename")
                    or "Respuesta de audio"
                )

    if selected["status"] in {"COMPLETADA", "EVALUADA"}:
        if st.button(
            "Recalcular asistencia automática",
            key=f"auto_eval_interview_{selected['id']}",
        ):
            report = save_automatic_interview_evaluation(
                selected["id"]
            )
            st.success("Asistencia recalculada.")
            st.session_state[
                f"interview_report_{selected['id']}"
            ] = report
            st.rerun()

        report = st.session_state.get(
            f"interview_report_{selected['id']}"
        )
        if not report:
            try:
                report = json.loads(
                    selected.get("auto_report_json") or "{}"
                )
            except json.JSONDecodeError:
                report = {}

        if report:
            st.markdown("#### Asistencia automática")
            render_interview_auto_report(report)
        else:
            st.info(
                "Todavía no existe una asistencia automática."
            )

    if not has_permission("manage_interviews"):
        return

    st.markdown("#### Evaluación humana")
    with st.form(
        f"human_interview_evaluation_{selected['id']}"
    ):
        human_scores = {}
        human_evidence = {}
        human_notes = {}

        for response in responses:
            question_id = response["id"]
            existing_human = fetch_one(
                """
                SELECT *
                FROM interview_evaluations
                WHERE interview_id = ?
                  AND question_id = ?
                  AND evaluator_type = 'HUMAN'
                """,
                (selected["id"], question_id),
            )

            st.markdown(
                f"**{response['position']}. "
                f"{response.get('competency') or 'Competencia'}**"
            )
            human_scores[question_id] = st.slider(
                "Puntaje",
                min_value=0.0,
                max_value=float(response.get("max_score") or 5),
                value=float(
                    existing_human.get("score")
                    if existing_human
                    and existing_human.get("score") is not None
                    else 0
                ),
                step=0.5,
                key=f"human_score_{selected['id']}_{question_id}",
            )
            human_evidence[question_id] = st.text_area(
                "Evidencia observada",
                value=(
                    existing_human.get("evidence_text") or ""
                    if existing_human
                    else ""
                ),
                key=f"human_evidence_{selected['id']}_{question_id}",
            )
            human_notes[question_id] = st.text_area(
                "Notas",
                value=(
                    existing_human.get("notes") or ""
                    if existing_human
                    else ""
                ),
                key=f"human_notes_{selected['id']}_{question_id}",
            )

        final_recommendation = st.selectbox(
            "Decisión final",
            INTERVIEW_RECOMMENDATIONS,
        )
        final_reason = st.text_area(
            "Justificación final",
            help=(
                "Explicá las evidencias que respaldan la decisión. "
                "La recomendación automática es solo un insumo."
            ),
        )
        save_human = st.form_submit_button(
            "Guardar evaluación humana",
            type="primary",
        )

        if save_human:
            if not final_reason.strip():
                st.error("La justificación final es obligatoria.")
            else:
                obtained = 0.0
                possible = 0.0

                for response in responses:
                    question_id = response["id"]
                    score = float(human_scores[question_id])
                    maximum = float(
                        response.get("max_score") or 5
                    )
                    obtained += score
                    possible += maximum

                    existing_human = fetch_one(
                        """
                        SELECT id
                        FROM interview_evaluations
                        WHERE interview_id = ?
                          AND question_id = ?
                          AND evaluator_type = 'HUMAN'
                        """,
                        (selected["id"], question_id),
                    )

                    params = (
                        score,
                        human_evidence[question_id].strip(),
                        human_notes[question_id].strip(),
                        user["id"],
                        now_iso(),
                        selected["id"],
                        question_id,
                    )

                    if existing_human:
                        execute(
                            """
                            UPDATE interview_evaluations
                            SET score = ?,
                                evidence_text = ?,
                                notes = ?,
                                evaluated_by = ?,
                                evaluated_at = ?
                            WHERE interview_id = ?
                              AND question_id = ?
                              AND evaluator_type = 'HUMAN'
                            """,
                            params,
                        )
                    else:
                        execute(
                            """
                            INSERT INTO interview_evaluations(
                                score,
                                evidence_text,
                                notes,
                                evaluated_by,
                                evaluated_at,
                                evaluator_type,
                                interview_id,
                                question_id
                            )
                            VALUES (?, ?, ?, ?, ?, 'HUMAN', ?, ?)
                            """,
                            params,
                        )

                human_percentage = (
                    round((obtained / possible) * 100, 1)
                    if possible > 0
                    else None
                )

                execute(
                    """
                    UPDATE interviews
                    SET status = 'EVALUADA',
                        human_score = ?,
                        final_recommendation = ?,
                        final_reason = ?,
                        evaluated_by = ?,
                        evaluated_at = ?,
                        updated_at = ?
                    WHERE id = ?
                    """,
                    (
                        human_percentage,
                        final_recommendation,
                        final_reason.strip(),
                        user["id"],
                        now_iso(),
                        now_iso(),
                        selected["id"],
                    ),
                )

                log_event(
                    user["company_id"],
                    user["id"],
                    "EVALUATE_INTERVIEW",
                    "interview",
                    selected["id"],
                    {
                        "human_score": human_percentage,
                        "final_recommendation": (
                            final_recommendation
                        ),
                    },
                )
                st.success("Evaluación humana guardada.")
                st.rerun()

    if selected.get("final_recommendation"):
        st.markdown("#### Decisión final registrada")
        st.write(
            f"**Decisión:** "
            f"{selected['final_recommendation']}"
        )
        st.write(
            f"**Justificación:** "
            f"{selected.get('final_reason') or 'Sin detalle'}"
        )


def render_interviews(user: dict) -> None:
    render_interviews_redesigned(user)


def save_candidate_interview(
    user: dict,
    candidate: dict,
    interview: dict,
    questions: list[dict],
    text_answers: dict[int, str],
    audio_files: dict[int, object],
) -> None:
    for question in questions:
        question_id = question["id"]
        answer_text = (
            text_answers.get(question_id) or ""
        ).strip()
        audio_file = audio_files.get(question_id)

        audio_blob = None
        audio_mime = None
        audio_filename = None

        if audio_file is not None:
            (
                audio_blob,
                audio_mime,
                audio_filename,
            ) = validate_audio_file(audio_file)

        response_mode = (
            question.get("response_mode")
            or "TEXTO_O_AUDIO"
        )
        required = bool(question.get("required"))

        if required:
            if response_mode == "SOLO_TEXTO" and not answer_text:
                raise ValueError(
                    f"La pregunta {question['position']} requiere "
                    "una respuesta escrita."
                )
            if response_mode == "SOLO_AUDIO" and audio_blob is None:
                raise ValueError(
                    f"La pregunta {question['position']} requiere "
                    "un audio."
                )
            if (
                response_mode == "TEXTO_O_AUDIO"
                and not answer_text
                and audio_blob is None
            ):
                raise ValueError(
                    f"Completá la pregunta {question['position']} "
                    "por escrito o mediante audio."
                )

        existing = fetch_one(
            """
            SELECT id
            FROM interview_responses
            WHERE interview_id = ?
              AND question_id = ?
            """,
            (interview["id"], question_id),
        )

        if existing:
            if audio_blob is None:
                execute(
                    """
                    UPDATE interview_responses
                    SET answer_text = ?,
                        submitted_at = ?,
                        updated_at = ?
                    WHERE interview_id = ?
                      AND question_id = ?
                    """,
                    (
                        answer_text,
                        now_iso(),
                        now_iso(),
                        interview["id"],
                        question_id,
                    ),
                )
            else:
                execute(
                    """
                    UPDATE interview_responses
                    SET answer_text = ?,
                        audio_blob = ?,
                        audio_mime = ?,
                        audio_filename = ?,
                        submitted_at = ?,
                        updated_at = ?
                    WHERE interview_id = ?
                      AND question_id = ?
                    """,
                    (
                        answer_text,
                        audio_blob,
                        audio_mime,
                        audio_filename,
                        now_iso(),
                        now_iso(),
                        interview["id"],
                        question_id,
                    ),
                )
        else:
            execute(
                """
                INSERT INTO interview_responses(
                    interview_id,
                    question_id,
                    answer_text,
                    audio_blob,
                    audio_mime,
                    audio_filename,
                    submitted_at,
                    updated_at
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    interview["id"],
                    question_id,
                    answer_text,
                    audio_blob,
                    audio_mime,
                    audio_filename,
                    now_iso(),
                    now_iso(),
                ),
            )

    execute(
        """
        UPDATE interviews
        SET status = 'COMPLETADA',
            completed_at = ?,
            updated_at = ?
        WHERE id = ?
        """,
        (
            now_iso(),
            now_iso(),
            interview["id"],
        ),
    )

    save_automatic_interview_evaluation(
        interview["id"]
    )

    log_event(
        None,
        user["id"],
        "COMPLETE_INTERVIEW",
        "interview",
        interview["id"],
        {
            "candidate_id": candidate["id"],
            "questions": len(questions),
        },
    )


def render_candidate_interviews(
    user: dict,
    candidate: dict,
) -> None:
    st.subheader("Mis entrevistas con ALBA")
    st.info(INTERVIEW_NOTICE)

    interviews = fetch_all(
        """
        SELECT
            interviews.*,
            interview_templates.name AS template_name,
            interview_templates.description AS template_description,
            jobs.title AS job_title,
            jobs.company_id,
            companies.name AS company_name
        FROM interviews
        JOIN interview_templates
            ON interview_templates.id = interviews.template_id
        JOIN applications
            ON applications.id = interviews.application_id
        JOIN jobs
            ON jobs.id = applications.job_id
        JOIN companies
            ON companies.id = jobs.company_id
        WHERE applications.candidate_id = ?
        ORDER BY interviews.assigned_at DESC
        """,
        (candidate["id"],),
    )

    if not interviews:
        st.info("No tenés entrevistas asignadas.")
        return

    options = {
        (
            f"#{item['id']} · {item['company_name']} · "
            f"{item['job_title']} · {item['status']}"
        ): item
        for item in interviews
    }
    selected_label = st.selectbox(
        "Entrevista",
        list(options.keys()),
        key="candidate_interview_select",
    )
    selected = options[selected_label]
    avatar_settings = get_avatar_settings(
        selected["company_id"]
    )

    render_avatar_identity(
        avatar_settings
    )

    st.markdown(f"### {selected['job_title']}")
    st.write(f"**Empresa:** {selected['company_name']}")
    st.write(f"**Plantilla:** {selected['template_name']}")
    st.write(f"**Estado:** {selected['status']}")
    st.write(
        f"**Fecha límite:** "
        f"{selected.get('due_date') or 'Sin fecha límite'}"
    )
    st.write(
        selected.get("intro_message")
        or selected.get("template_description")
        or ""
    )

    if selected["status"] in {
        "COMPLETADA",
        "EVALUADA",
    }:
        st.success(
            "La entrevista ya fue enviada. "
            "RR. HH. realizará la revisión final."
        )
        return

    if selected["status"] in {"VENCIDA", "CANCELADA"}:
        st.warning(
            "Esta entrevista ya no está disponible."
        )
        return

    if selected["status"] == "INVITADA":
        if st.button(
            "Comenzar entrevista",
            type="primary",
            key=f"start_interview_{selected['id']}",
        ):
            execute(
                """
                UPDATE interviews
                SET status = 'EN CURSO',
                    started_at = ?,
                    updated_at = ?
                WHERE id = ?
                """,
                (
                    now_iso(),
                    now_iso(),
                    selected["id"],
                ),
            )
            st.rerun()
        return

    st.markdown("### Sala de videollamada")
    live_agent_available = render_live_alba_agent(
        f"candidate_live_agent_{selected['id']}"
    )

    with st.expander("Vista previa de mi cámara"):
        render_camera_preview(
            f"candidate_camera_{selected['id']}"
        )

    if live_agent_available:
        st.info(
            "La videollamada con el avatar es complementaria. "
            "Las respuestas estructuradas y su consentimiento se guardan debajo."
        )

    questions = get_template_questions(
        selected["template_id"],
        active_only=True,
    )
    if not questions:
        st.error(
            "La plantilla no tiene preguntas activas."
        )
        return

    existing_responses = fetch_all(
        """
        SELECT *
        FROM interview_responses
        WHERE interview_id = ?
        """,
        (selected["id"],),
    )
    responses_by_question = {
        item["question_id"]: item
        for item in existing_responses
    }

    text_answers = {}
    audio_files = {}

    with st.form(
        f"candidate_interview_form_{selected['id']}"
    ):
        st.caption(
            "Respondé con ejemplos concretos. Cuando sea posible, "
            "explicá la situación, tu responsabilidad, las acciones "
            "realizadas y el resultado."
        )

        for question in questions:
            question_id = question["id"]
            existing = responses_by_question.get(
                question_id,
                {},
            )

            render_question_avatar(
                selected["company_id"],
                question,
                avatar_settings,
            )

            response_mode = (
                question.get("response_mode")
                or "TEXTO_O_AUDIO"
            )

            if response_mode != "SOLO_AUDIO":
                text_answers[question_id] = st.text_area(
                    "Respuesta escrita o transcripción",
                    value=existing.get("answer_text") or "",
                    height=160,
                    key=(
                        f"candidate_interview_text_"
                        f"{selected['id']}_{question_id}"
                    ),
                )
            else:
                text_answers[question_id] = ""

            if response_mode != "SOLO_TEXTO":
                audio_files[question_id] = st.file_uploader(
                    "Audio opcional"
                    if response_mode == "TEXTO_O_AUDIO"
                    else "Audio",
                    type=["mp3", "wav", "m4a", "ogg", "webm"],
                    key=(
                        f"candidate_interview_audio_"
                        f"{selected['id']}_{question_id}"
                    ),
                )
                if existing.get("audio_blob"):
                    st.caption(
                        "Ya existe un audio guardado para esta pregunta."
                    )
            else:
                audio_files[question_id] = None

            st.divider()

        consent = st.checkbox(
            "Confirmo que las respuestas son propias y autorizo "
            "su revisión para este proceso de selección."
        )

        submitted = st.form_submit_button(
            "Enviar entrevista",
            type="primary",
            use_container_width=True,
        )

        if submitted:
            if not consent:
                st.error(
                    "Tenés que confirmar la declaración antes de enviar."
                )
            else:
                try:
                    save_candidate_interview(
                        user,
                        candidate,
                        selected,
                        questions,
                        text_answers,
                        audio_files,
                    )
                    st.success(
                        "Entrevista enviada correctamente."
                    )
                    st.rerun()
                except Exception as exc:
                    st.error(str(exc))



# =========================================================
# EXPERIENCIA DE USUARIO
# =========================================================

def inject_ui_styles() -> None:
    st.markdown(
        """
        <style>
        .block-container {
            max-width: 1450px;
            padding-top: 1.2rem;
            padding-bottom: 3rem;
        }
        [data-testid="stSidebar"] {
            min-width: 285px;
        }
        .alba-card {
            border: 1px solid rgba(128,128,128,.24);
            border-radius: 16px;
            padding: 18px;
            min-height: 150px;
            background: rgba(128,128,128,.035);
        }
        .alba-help {
            padding: 12px 16px;
            border-radius: 10px;
            background: rgba(59,130,246,.08);
            border: 1px solid rgba(59,130,246,.18);
        }
        div[data-testid="stMetric"] {
            border: 1px solid rgba(128,128,128,.20);
            padding: 12px;
            border-radius: 14px;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def navigate_company(section: str) -> None:
    st.session_state["company_menu"] = section
    st.rerun()


def render_process_guide() -> None:
    st.markdown("### Flujo recomendado")
    steps = [
        ("1", "Crear la búsqueda", "Definí requisitos, criterios y competencias."),
        ("2", "Agregar candidatos", "Subí CV, usá el banco o elegí sugerencias."),
        ("3", "Revisar el ranking", "Validá fortalezas, brechas y requisitos críticos."),
        ("4", "Entrevistar con ALBA", "Asigná la entrevista, grabá y tomá notas."),
        ("5", "Evaluar y decidir", "Aplicá la escala 1–5 y documentá la decisión."),
    ]
    columns = st.columns(5)
    for column, (number, title, detail) in zip(columns, steps):
        with column:
            st.markdown(
                f"""
                <div class="alba-card">
                    <h3>{number}</h3>
                    <strong>{title}</strong>
                    <p>{detail}</p>
                </div>
                """,
                unsafe_allow_html=True,
            )


def render_intuitive_dashboard(user: dict, company: dict) -> None:
    jobs_total = fetch_one(
        "SELECT COUNT(*) AS total FROM jobs WHERE company_id = ?",
        (company["id"],),
    )["total"]
    open_jobs = fetch_one(
        "SELECT COUNT(*) AS total FROM jobs WHERE company_id = ? AND status = 'ABIERTA'",
        (company["id"],),
    )["total"]
    applications_total = fetch_one(
        """
        SELECT COUNT(*) AS total
        FROM applications
        JOIN jobs ON jobs.id = applications.job_id
        WHERE jobs.company_id = ?
        """,
        (company["id"],),
    )["total"]
    pending_interviews = fetch_one(
        """
        SELECT COUNT(*) AS total
        FROM interviews
        JOIN applications ON applications.id = interviews.application_id
        JOIN jobs ON jobs.id = applications.job_id
        WHERE jobs.company_id = ?
          AND interviews.status IN ('INVITADA', 'EN CURSO', 'COMPLETADA')
        """,
        (company["id"],),
    )["total"]

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Búsquedas", jobs_total)
    c2.metric("Abiertas", open_jobs)
    c3.metric("Candidatos asignados", applications_total)
    c4.metric("Entrevistas pendientes", pending_interviews)

    render_process_guide()

    st.markdown("### Acciones rápidas")
    a1, a2, a3, a4 = st.columns(4)
    if a1.button("➕ Crear búsqueda", use_container_width=True):
        navigate_company("💼 Búsquedas")
    if a2.button("⚡ Agregar candidatos", use_container_width=True):
        navigate_company("⚡ Asignación rápida")
    if a3.button("🎥 Gestionar entrevistas", use_container_width=True):
        navigate_company("🎥 Entrevistas")
    if a4.button("👥 Revisar candidatos", use_container_width=True):
        navigate_company("👥 Candidatos")

    st.markdown(
        """
        <div class="alba-help">
            <strong>Consejo:</strong> comenzá por una búsqueda y usá
            “Asignación rápida” para subir CV, elegir candidatos del banco
            o aceptar sugerencias sin recorrer varias pantallas.
        </div>
        """,
        unsafe_allow_html=True,
    )


# =========================================================
# ASIGNACIÓN RÁPIDA DE CANDIDATOS
# =========================================================

def render_quick_candidate_assignment(user: dict) -> None:
    st.subheader("Asignación rápida de candidatos")
    st.caption(
        "Unificamos en una sola pantalla la carga de CV, el banco y las sugerencias."
    )

    jobs = fetch_all(
        """
        SELECT *
        FROM jobs
        WHERE company_id = ?
          AND status <> 'CERRADA'
        ORDER BY created_at DESC
        """,
        (user["company_id"],),
    )

    if not jobs:
        st.info("Primero creá una búsqueda abierta o pausada.")
        return

    job_options = {
        f"{job['title']} · {job.get('status') or 'ABIERTA'}": job
        for job in jobs
    }
    selected_job_label = st.selectbox(
        "1. Elegí la búsqueda",
        list(job_options.keys()),
        key="quick_assignment_job",
    )
    selected_job = job_options[selected_job_label]

    tab_upload, tab_bank, tab_suggestions = st.tabs(
        ["Subir CV y asignar", "Elegir del banco", "Sugerencias de ALBA"]
    )

    with tab_upload:
        uploaded_files = st.file_uploader(
            "2. Subí uno o varios CV",
            type=["pdf", "docx"],
            accept_multiple_files=True,
            key="quick_assignment_files",
        )
        initial_status = st.selectbox(
            "3. Etapa inicial",
            APPLICATION_STATUSES,
            index=APPLICATION_STATUSES.index("EN REVISIÓN"),
            key="quick_upload_status",
        )

        if st.button(
            "Procesar y asignar ahora",
            type="primary",
            use_container_width=True,
            key="quick_upload_assign",
        ):
            if not uploaded_files:
                st.error("Seleccioná al menos un CV.")
            elif len(uploaded_files) > MAX_CV_BATCH:
                st.error(f"El máximo por lote es {MAX_CV_BATCH}.")
            else:
                results = []
                for uploaded_file in uploaded_files:
                    try:
                        raw, mime_type, text = extract_cv_text(uploaded_file)
                        parsed = parse_cv_text(text, uploaded_file.name)
                        candidate_id, document_id = get_or_create_candidate_from_cv(
                            parsed,
                            raw,
                            mime_type,
                            text,
                            uploaded_file.name,
                            user["id"],
                        )
                        ensure_company_candidate(
                            user["company_id"],
                            candidate_id,
                            "ASIGNACIÓN_RÁPIDA",
                            user["id"],
                        )
                        application_id, _ = assign_candidate_to_job(
                            user,
                            candidate_id,
                            selected_job["id"],
                            initial_status,
                        )
                        results.append(
                            {
                                "Archivo": uploaded_file.name,
                                "Candidato": parsed.get("full_name"),
                                "Estado": "Asignado",
                                "Postulación": application_id,
                            }
                        )
                        log_event(
                            user["company_id"],
                            user["id"],
                            "QUICK_UPLOAD_ASSIGN",
                            "candidate_document",
                            document_id,
                            {
                                "job_id": selected_job["id"],
                                "candidate_id": candidate_id,
                            },
                        )
                    except Exception as exc:
                        results.append(
                            {
                                "Archivo": uploaded_file.name,
                                "Candidato": "",
                                "Estado": f"Error: {exc}",
                                "Postulación": "",
                            }
                        )

                st.dataframe(
                    pd.DataFrame(results),
                    use_container_width=True,
                    hide_index=True,
                )
                successful = sum(
                    1 for item in results if item["Estado"] == "Asignado"
                )
                st.success(f"Candidatos asignados: {successful}.")

    with tab_bank:
        rows = fetch_all(
            """
            SELECT
                candidates.*,
                users.full_name,
                users.email
            FROM company_candidates
            JOIN candidates
                ON candidates.id = company_candidates.candidate_id
            JOIN users
                ON users.id = candidates.user_id
            WHERE company_candidates.company_id = ?
              AND NOT EXISTS (
                  SELECT 1
                  FROM applications
                  WHERE applications.candidate_id = candidates.id
                    AND applications.job_id = ?
              )
            ORDER BY users.full_name
            """,
            (
                user["company_id"],
                selected_job["id"],
            ),
        )

        search_text = st.text_input(
            "Buscar por nombre, perfil, habilidades o ubicación",
            key="quick_bank_search",
        )
        normalized_search = normalize_match_text(search_text)

        filtered_rows = []
        for row in rows:
            searchable = normalize_match_text(
                " ".join(
                    [
                        row.get("full_name") or "",
                        row.get("headline") or "",
                        row.get("skills_text") or "",
                        row.get("tags_text") or "",
                        row.get("country") or "",
                        row.get("city") or "",
                        row.get("locality") or "",
                    ]
                )
            )
            if normalized_search and normalized_search not in searchable:
                continue
            filtered_rows.append(row)

        candidate_options = {
            (
                f"#{row['id']} · {row['full_name']} · "
                f"{row.get('headline') or 'Sin resumen'}"
            ): row["id"]
            for row in filtered_rows
        }

        selected_candidates = st.multiselect(
            "Seleccionar candidatos",
            list(candidate_options.keys()),
            key="quick_bank_candidates",
        )
        bank_status = st.selectbox(
            "Etapa inicial",
            APPLICATION_STATUSES,
            index=APPLICATION_STATUSES.index("EN REVISIÓN"),
            key="quick_bank_status",
        )

        if st.button(
            "Asignar seleccionados",
            type="primary",
            use_container_width=True,
            key="quick_bank_assign",
        ):
            if not selected_candidates:
                st.error("Seleccioná al menos un candidato.")
            else:
                for label in selected_candidates:
                    assign_candidate_to_job(
                        user,
                        candidate_options[label],
                        selected_job["id"],
                        bank_status,
                    )
                st.success(
                    f"Candidatos asignados: {len(selected_candidates)}."
                )
                st.rerun()

        if filtered_rows:
            st.dataframe(
                pd.DataFrame(
                    [
                        {
                            "Candidato": row["full_name"],
                            "Perfil": row.get("headline") or "",
                            "País": row.get("country") or "",
                            "Ciudad": row.get("city") or "",
                            "Localidad": row.get("locality") or "",
                            "Etiquetas": row.get("tags_text") or "",
                        }
                        for row in filtered_rows
                    ]
                ),
                use_container_width=True,
                hide_index=True,
            )
        else:
            st.info("No hay candidatos disponibles para esta búsqueda.")

    with tab_suggestions:
        if st.button(
            "Actualizar sugerencias",
            key="quick_generate_suggestions",
        ):
            result = generate_candidate_job_suggestions(
                user["company_id"],
                selected_job["id"],
            )
            st.success(
                f"Comparaciones realizadas: {result['generated']}."
            )
            st.rerun()

        suggestions = get_job_suggestions(
            user["company_id"],
            selected_job["id"],
        )
        suggestion_options = {
            (
                f"#{item['candidate_id']} · {item['full_name']} · "
                f"{float(item.get('score_total') or 0):.1f}%"
            ): item
            for item in suggestions
        }
        selected_suggestions = st.multiselect(
            "Seleccionar sugerencias",
            list(suggestion_options.keys()),
            key="quick_suggestions_select",
        )

        if st.button(
            "Aceptar sugerencias y asignar",
            type="primary",
            use_container_width=True,
            key="quick_suggestions_assign",
        ):
            if not selected_suggestions:
                st.error("Seleccioná al menos una sugerencia.")
            else:
                for label in selected_suggestions:
                    item = suggestion_options[label]
                    assign_candidate_to_job(
                        user,
                        item["candidate_id"],
                        selected_job["id"],
                        "EN REVISIÓN",
                    )
                    execute(
                        """
                        UPDATE candidate_job_suggestions
                        SET status = 'ASIGNADA'
                        WHERE id = ?
                        """,
                        (item["id"],),
                    )
                st.success(
                    f"Sugerencias asignadas: {len(selected_suggestions)}."
                )
                st.rerun()

        for item in suggestions:
            with st.container(border=True):
                c1, c2 = st.columns([3, 1])
                with c1:
                    st.markdown(f"**{item['full_name']}**")
                    st.write(item.get("headline") or "Sin resumen profesional.")
                    st.caption(build_location_label(item))
                with c2:
                    st.metric(
                        "Coincidencia",
                        f"{float(item.get('score_total') or 0):.1f}%",
                    )


# =========================================================
# GUÍA Y METODOLOGÍA POR COMPETENCIAS
# =========================================================

def ensure_default_competency_guides(company_id: int) -> int:
    created = 0
    for name, data in COMPETENCY_GUIDE_DEFAULTS.items():
        existing = fetch_one(
            """
            SELECT id
            FROM competency_guides
            WHERE company_id = ? AND name = ?
            """,
            (company_id, name),
        )
        if existing:
            continue

        execute(
            """
            INSERT INTO competency_guides(
                company_id,
                name,
                definition,
                keywords_json,
                questions_json,
                critical_default,
                active,
                created_at,
                updated_at
            )
            VALUES (?, ?, ?, ?, ?, 0, 1, ?, ?)
            """,
            (
                company_id,
                name,
                data["definition"],
                json.dumps(data["keywords"], ensure_ascii=False),
                json.dumps(data["questions"], ensure_ascii=False),
                now_iso(),
                now_iso(),
            ),
        )
        created += 1
    return created


def get_competency_guides(company_id: int) -> list[dict]:
    ensure_default_competency_guides(company_id)
    return fetch_all(
        """
        SELECT *
        FROM competency_guides
        WHERE company_id = ?
          AND active = 1
        ORDER BY name
        """,
        (company_id,),
    )


def render_competency_methodology() -> None:
    st.markdown("### Metodología de evaluación por competencias")

    scale_rows = [
        {
            "Puntaje": score,
            "Nivel": data["level"],
            "Descripción": data["description"],
        }
        for score, data in COMPETENCY_SCALE.items()
    ]
    st.dataframe(
        pd.DataFrame(scale_rows),
        use_container_width=True,
        hide_index=True,
    )

    st.markdown(
        """
        **Cálculo del resultado final**

        `Promedio final = suma de los puntajes / cantidad de competencias`

        **Semáforo**

        - 🟢 **Verde (4,0 a 5,0):** recomendada/o.
        - 🟡 **Amarillo (3,0 a 3,9):** recomendada/o con observaciones.
        - 🔴 **Rojo (1,0 a 2,9):** no recomendada/o.

        **Competencias críticas**

        - Una competencia crítica con **1** determina resultado **rojo**.
        - Una competencia crítica con **2** limita el resultado máximo a **amarillo**.
        - Con competencias críticas de **3 o más**, se aplica el semáforo del promedio.
        """
    )

    example = pd.DataFrame(
        [
            {"Competencia": "Comunicación", "Puntaje": 4},
            {"Competencia": "Trabajo en Equipo", "Puntaje": 5},
            {"Competencia": "Orientación a Resultados", "Puntaje": 4},
            {"Competencia": "Adaptabilidad", "Puntaje": 3},
            {"Competencia": "Resolución de Problemas", "Puntaje": 4},
        ]
    )
    st.markdown("**Ejemplo: promedio final 4,0**")
    st.dataframe(example, use_container_width=True, hide_index=True)


def render_competency_guide(user: dict) -> None:
    render_competency_methodology()

    st.markdown("### Guía por competencia")
    guides = get_competency_guides(user["company_id"])

    for guide in guides:
        try:
            keywords = json.loads(guide.get("keywords_json") or "[]")
            questions = json.loads(guide.get("questions_json") or "[]")
        except json.JSONDecodeError:
            keywords, questions = [], []

        with st.expander(
            f"{'⚠️ ' if guide.get('critical_default') else ''}{guide['name']}"
        ):
            st.write(guide.get("definition") or "")
            st.write(
                "**Indicadores para observar:** "
                + (", ".join(keywords) or "No configurados")
            )
            st.markdown("**Preguntas sugeridas**")
            for question in questions:
                st.write(f"• {question}")

            if has_permission("manage_interviews"):
                critical = st.checkbox(
                    "Marcar como crítica por defecto",
                    value=bool(guide.get("critical_default")),
                    key=f"guide_critical_{guide['id']}",
                )
                if st.button(
                    "Guardar configuración",
                    key=f"save_guide_{guide['id']}",
                ):
                    execute(
                        """
                        UPDATE competency_guides
                        SET critical_default = ?,
                            updated_at = ?
                        WHERE id = ?
                          AND company_id = ?
                        """,
                        (
                            int(critical),
                            now_iso(),
                            guide["id"],
                            user["company_id"],
                        ),
                    )
                    st.success("Guía actualizada.")
                    st.rerun()


def calculate_competency_result(scores: list[dict]) -> dict:
    if not scores:
        return {
            "average": None,
            "traffic_light": "PENDIENTE",
            "recommendation": "Pendiente de evaluación",
            "critical_alerts": [],
        }

    average = round(
        sum(float(item["score"]) for item in scores) / len(scores),
        2,
    )

    if average >= 4:
        traffic_light = "VERDE"
        recommendation = "Recomendada/o"
    elif average >= 3:
        traffic_light = "AMARILLO"
        recommendation = "Recomendada/o con observaciones"
    else:
        traffic_light = "ROJO"
        recommendation = "No recomendada/o"

    critical_score_1 = [
        item for item in scores
        if item.get("critical") and int(item["score"]) == 1
    ]
    critical_score_2 = [
        item for item in scores
        if item.get("critical") and int(item["score"]) == 2
    ]

    critical_alerts = []

    if critical_score_1:
        traffic_light = "ROJO"
        recommendation = "No recomendada/o"
        critical_alerts.append(
            "Existe al menos una competencia crítica con puntaje 1."
        )
    elif critical_score_2 and traffic_light == "VERDE":
        traffic_light = "AMARILLO"
        recommendation = "Recomendada/o con observaciones"
        critical_alerts.append(
            "Una competencia crítica con puntaje 2 limita el resultado máximo a amarillo."
        )
    elif critical_score_2:
        critical_alerts.append(
            "Existe una competencia crítica con puntaje 2."
        )

    return {
        "average": average,
        "traffic_light": traffic_light,
        "recommendation": recommendation,
        "critical_alerts": critical_alerts,
    }


def build_competency_report(
    candidate_name: str,
    job_title: str,
    scores: list[dict],
    transcript: str,
    notes: str,
) -> dict:
    result = calculate_competency_result(scores)

    strengths = [
        item for item in scores if int(item["score"]) >= 4
    ]
    expected = [
        item for item in scores if int(item["score"]) == 3
    ]
    development = [
        item for item in scores if int(item["score"]) <= 2
    ]

    conclusion = (
        f"{candidate_name} obtiene un promedio final de "
        f"{result['average'] if result['average'] is not None else 'pendiente'} "
        f"para la búsqueda {job_title}. Resultado: "
        f"{result['traffic_light']} – {result['recommendation']}."
    )

    return {
        "candidate_name": candidate_name,
        "job_title": job_title,
        "average": result["average"],
        "traffic_light": result["traffic_light"],
        "recommendation": result["recommendation"],
        "critical_alerts": result["critical_alerts"],
        "strengths": strengths,
        "expected": expected,
        "development": development,
        "scores": scores,
        "transcript_excerpt": (transcript or "")[:2500],
        "interviewer_notes": notes or "",
        "conclusion": conclusion,
        "generated_at": now_iso(),
    }


def report_to_markdown(report: dict) -> str:
    lines = [
        "# Informe de entrevista por competencias",
        "",
        f"**Candidato/a:** {report.get('candidate_name')}",
        f"**Búsqueda:** {report.get('job_title')}",
        f"**Promedio final:** {report.get('average')}",
        f"**Semáforo:** {report.get('traffic_light')}",
        f"**Recomendación:** {report.get('recommendation')}",
        "",
        "## Conclusión",
        report.get("conclusion") or "",
        "",
        "## Competencias evaluadas",
    ]

    for item in report.get("scores") or []:
        critical_label = " · CRÍTICA" if item.get("critical") else ""
        lines.append(
            f"- **{item['competency']}**: {item['score']}/5{critical_label}. "
            f"Evidencia: {item.get('evidence') or 'Sin detalle'}"
        )

    if report.get("critical_alerts"):
        lines.extend(["", "## Alertas críticas"])
        for item in report["critical_alerts"]:
            lines.append(f"- {item}")

    lines.extend(
        [
            "",
            "## Notas del entrevistador",
            report.get("interviewer_notes") or "Sin notas.",
        ]
    )

    return "\n".join(lines)


def save_competency_scores(
    user: dict,
    interview: dict,
    scores: list[dict],
    transcript: str,
    notes: str,
) -> dict:
    for item in scores:
        existing = fetch_one(
            """
            SELECT id
            FROM interview_competency_scores
            WHERE interview_id = ?
              AND competency_name = ?
            """,
            (
                interview["id"],
                item["competency"],
            ),
        )
        params = (
            int(item["score"]),
            int(bool(item.get("critical"))),
            item.get("evidence", "").strip(),
            item.get("notes", "").strip(),
            user["id"],
            now_iso(),
            interview["id"],
            item["competency"],
        )

        if existing:
            execute(
                """
                UPDATE interview_competency_scores
                SET score = ?,
                    critical = ?,
                    evidence_text = ?,
                    notes = ?,
                    evaluated_by = ?,
                    evaluated_at = ?
                WHERE interview_id = ?
                  AND competency_name = ?
                """,
                params,
            )
        else:
            execute(
                """
                INSERT INTO interview_competency_scores(
                    score,
                    critical,
                    evidence_text,
                    notes,
                    evaluated_by,
                    evaluated_at,
                    interview_id,
                    competency_name
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                params,
            )

    report = build_competency_report(
        interview["candidate_name"],
        interview["job_title"],
        scores,
        transcript,
        notes,
    )

    execute(
        """
        UPDATE interviews
        SET human_score = ?,
            traffic_light = ?,
            final_recommendation = ?,
            final_reason = ?,
            competency_result_json = ?,
            status = 'EVALUADA',
            evaluated_by = ?,
            evaluated_at = ?,
            updated_at = ?
        WHERE id = ?
        """,
        (
            report["average"],
            report["traffic_light"],
            report["recommendation"],
            report["conclusion"],
            json.dumps(report, ensure_ascii=False),
            user["id"],
            now_iso(),
            now_iso(),
            interview["id"],
        ),
    )

    return report


# =========================================================
# SALA, GRABACIÓN, NOTAS Y ANÁLISIS
# =========================================================

def get_interview_session(interview_id: int) -> dict:
    saved = fetch_one(
        """
        SELECT *
        FROM interview_sessions
        WHERE interview_id = ?
        """,
        (interview_id,),
    )
    return saved or {
        "interview_id": interview_id,
        "transcript_text": "",
        "interviewer_notes": "",
        "analysis_json": "",
    }


def save_interview_session_text(
    user: dict,
    interview_id: int,
    transcript: str,
    notes: str,
    analysis: dict | None = None,
) -> None:
    existing = fetch_one(
        "SELECT id FROM interview_sessions WHERE interview_id = ?",
        (interview_id,),
    )
    analysis_json = json.dumps(
        analysis or {},
        ensure_ascii=False,
    )

    if existing:
        execute(
            """
            UPDATE interview_sessions
            SET transcript_text = ?,
                interviewer_notes = ?,
                analysis_json = ?,
                created_by = ?,
                updated_at = ?
            WHERE interview_id = ?
            """,
            (
                transcript.strip(),
                notes.strip(),
                analysis_json,
                user["id"],
                now_iso(),
                interview_id,
            ),
        )
    else:
        execute(
            """
            INSERT INTO interview_sessions(
                interview_id,
                transcript_text,
                interviewer_notes,
                analysis_json,
                created_by,
                created_at,
                updated_at
            )
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (
                interview_id,
                transcript.strip(),
                notes.strip(),
                analysis_json,
                user["id"],
                now_iso(),
                now_iso(),
            ),
        )


def save_interview_recording(
    user: dict,
    interview_id: int,
    raw: bytes,
    mime_type: str,
    filename: str,
) -> None:
    if not raw:
        raise ValueError("La grabación está vacía.")

    existing = fetch_one(
        "SELECT id FROM interview_sessions WHERE interview_id = ?",
        (interview_id,),
    )

    if existing:
        execute(
            """
            UPDATE interview_sessions
            SET recording_blob = ?,
                recording_mime = ?,
                recording_filename = ?,
                created_by = ?,
                updated_at = ?
            WHERE interview_id = ?
            """,
            (
                raw,
                mime_type,
                filename,
                user["id"],
                now_iso(),
                interview_id,
            ),
        )
    else:
        execute(
            """
            INSERT INTO interview_sessions(
                interview_id,
                recording_blob,
                recording_mime,
                recording_filename,
                created_by,
                created_at,
                updated_at
            )
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (
                interview_id,
                raw,
                mime_type,
                filename,
                user["id"],
                now_iso(),
                now_iso(),
            ),
        )


def analyze_interview_text(
    company_id: int,
    transcript: str,
    notes: str,
) -> dict:
    combined = normalize_match_text(
        f"{transcript}\n{notes}"
    )
    guides = get_competency_guides(company_id)
    results = []

    for guide in guides:
        try:
            keywords = json.loads(
                guide.get("keywords_json") or "[]"
            )
        except json.JSONDecodeError:
            keywords = []

        normalized_keywords = [
            normalize_match_text(item)
            for item in keywords
            if normalize_match_text(item)
        ]
        matched = [
            item for item in normalized_keywords
            if item in combined
        ]
        coverage = (
            len(matched) / len(normalized_keywords)
            if normalized_keywords
            else 0
        )

        if not combined:
            suggested_score = None
        elif coverage >= 0.8:
            suggested_score = 5
        elif coverage >= 0.6:
            suggested_score = 4
        elif coverage >= 0.4:
            suggested_score = 3
        elif coverage >= 0.2:
            suggested_score = 2
        else:
            suggested_score = 1

        results.append(
            {
                "competency": guide["name"],
                "matched_keywords": matched,
                "coverage": round(coverage * 100, 1),
                "suggested_score": suggested_score,
                "critical_default": bool(
                    guide.get("critical_default")
                ),
            }
        )

    return {
        "competencies": results,
        "notice": (
            "Las sugerencias se basan únicamente en palabras y evidencias "
            "presentes en la transcripción o en las notas. Deben ser validadas "
            "por la persona entrevistadora."
        ),
    }


def render_live_alba_agent(context_key: str) -> bool:
    public_url = get_secret_value("DID_AGENT_PUBLIC_URL")
    embed_html = get_secret_value("DID_AGENT_EMBED_HTML")

    st.markdown("#### Videollamada con ALBA")

    if public_url:
        st.link_button(
            "Abrir videollamada con ALBA",
            public_url,
            use_container_width=True,
        )
        st.caption(
            "Se abre en una pestaña independiente para que el navegador pueda "
            "gestionar correctamente cámara y micrófono."
        )

    if embed_html:
        with st.expander("Vista integrada del agente", expanded=True):
            components.html(
                embed_html,
                height=650,
                scrolling=False,
            )
        return True

    if public_url:
        return True

    st.info(
        "Para activar la videollamada en tiempo real, agregá en Streamlit "
        "Secrets la URL pública o el código oficial de D-ID Agent Embed."
    )
    return False


def render_camera_preview(key: str) -> None:
    try:
        from streamlit_webrtc import WebRtcMode, webrtc_streamer
    except ImportError:
        st.info(
            "La vista de cámara estará disponible cuando Streamlit termine "
            "de instalar streamlit-webrtc."
        )
        return

    st.caption(
        "Vista previa local. La imagen no participa de la evaluación."
    )
    try:
        webrtc_streamer(
            key=key,
            mode=WebRtcMode.SENDONLY,
            rtc_configuration={
                "iceServers": [
                    {"urls": ["stun:stun.l.google.com:19302"]}
                ]
            },
            media_stream_constraints={
                "video": True,
                "audio": False,
            },
            async_processing=True,
        )
    except Exception as exc:
        st.warning(f"No se pudo iniciar la cámara: {exc}")


def render_browser_recording(
    user: dict,
    interview_id: int,
) -> None:
    st.markdown("#### Grabar desde este dispositivo")
    st.caption(INTERVIEW_RECORDING_NOTICE)

    try:
        from aiortc.contrib.media import MediaRecorder
        from streamlit_webrtc import WebRtcMode, webrtc_streamer
    except ImportError:
        st.warning(
            "La grabación directa no está disponible todavía. "
            "Usá la carga manual de archivo."
        )
        return

    path_key = f"interview_recording_path_{interview_id}"
    if path_key not in st.session_state:
        st.session_state[path_key] = str(
            RECORDINGS_DIR
            / f"interview_{interview_id}_{uuid.uuid4().hex}.webm"
        )

    recording_path = Path(st.session_state[path_key])

    def recorder_factory():
        return MediaRecorder(str(recording_path))

    try:
        context = webrtc_streamer(
            key=f"interview_recorder_{interview_id}",
            mode=WebRtcMode.SENDONLY,
            rtc_configuration={
                "iceServers": [
                    {"urls": ["stun:stun.l.google.com:19302"]}
                ]
            },
            media_stream_constraints={
                "video": True,
                "audio": True,
            },
            in_recorder_factory=recorder_factory,
            async_processing=True,
        )

        st.caption(
            "Presioná START para grabar y STOP para finalizar. "
            "Luego guardá la grabación."
        )

        if (
            not context.state.playing
            and recording_path.exists()
            and recording_path.stat().st_size > 0
        ):
            size_mb = recording_path.stat().st_size / (1024 * 1024)
            st.success(f"Grabación disponible: {size_mb:.1f} MB.")

            if st.button(
                "Guardar grabación en la entrevista",
                type="primary",
                key=f"save_browser_recording_{interview_id}",
            ):
                raw = recording_path.read_bytes()
                save_interview_recording(
                    user,
                    interview_id,
                    raw,
                    "video/webm",
                    recording_path.name,
                )
                st.success("Grabación guardada.")
    except Exception as exc:
        st.warning(
            "La grabación directa no pudo iniciarse. "
            f"Detalle: {exc}"
        )


def render_interviewer_workspace(user: dict) -> None:
    interviews = fetch_all(
        """
        SELECT
            interviews.*,
            users.full_name AS candidate_name,
            jobs.title AS job_title,
            jobs.company_id,
            companies.name AS company_name
        FROM interviews
        JOIN applications
            ON applications.id = interviews.application_id
        JOIN candidates
            ON candidates.id = applications.candidate_id
        JOIN users
            ON users.id = candidates.user_id
        JOIN jobs
            ON jobs.id = applications.job_id
        JOIN companies
            ON companies.id = jobs.company_id
        WHERE jobs.company_id = ?
        ORDER BY interviews.assigned_at DESC
        """,
        (user["company_id"],),
    )

    if not interviews:
        st.info("No hay entrevistas asignadas.")
        return

    options = {
        (
            f"#{item['id']} · {item['candidate_name']} · "
            f"{item['job_title']} · {item['status']}"
        ): item
        for item in interviews
    }
    label = st.selectbox(
        "Seleccionar entrevista",
        list(options.keys()),
        key="workspace_interview",
    )
    interview = options[label]
    session = get_interview_session(interview["id"])

    st.markdown(
        f"### {interview['candidate_name']} · {interview['job_title']}"
    )

    room_tab, notes_tab, evaluation_tab, report_tab = st.tabs(
        [
            "Sala y grabación",
            "Notas y transcripción",
            "Evaluación por competencias",
            "Informe final",
        ]
    )

    with room_tab:
        left, right = st.columns([1, 1])
        with left:
            render_live_alba_agent(
                f"interviewer_agent_{interview['id']}"
            )
        with right:
            render_browser_recording(
                user,
                interview["id"],
            )

        st.markdown("#### Cargar una grabación completa")
        uploaded_recording = st.file_uploader(
            "Video o audio exportado de Meet, Teams, Zoom u otra plataforma",
            type=["mp4", "webm", "mov", "mp3", "wav", "m4a"],
            key=f"upload_interview_recording_{interview['id']}",
        )
        if uploaded_recording is not None and st.button(
            "Guardar archivo cargado",
            key=f"save_uploaded_recording_{interview['id']}",
        ):
            save_interview_recording(
                user,
                interview["id"],
                uploaded_recording.getvalue(),
                uploaded_recording.type or "application/octet-stream",
                uploaded_recording.name,
            )
            st.success("Archivo guardado en la entrevista.")
            st.rerun()

        current_session = get_interview_session(interview["id"])
        if current_session.get("recording_blob"):
            st.video(current_session["recording_blob"])
            st.download_button(
                "Descargar grabación",
                data=current_session["recording_blob"],
                file_name=current_session.get("recording_filename")
                or f"entrevista_{interview['id']}.webm",
                mime=current_session.get("recording_mime")
                or "video/webm",
                key=f"download_recording_{interview['id']}",
            )

    with notes_tab:
        transcript = st.text_area(
            "Transcripción de la entrevista",
            value=session.get("transcript_text") or "",
            height=280,
            help=(
                "Pegá la transcripción del audio o video. El análisis automático "
                "se realiza sobre este texto y las notas."
            ),
            key=f"workspace_transcript_{interview['id']}",
        )
        notes = st.text_area(
            "Notas del entrevistador",
            value=session.get("interviewer_notes") or "",
            height=220,
            key=f"workspace_notes_{interview['id']}",
        )

        n1, n2 = st.columns(2)
        if n1.button(
            "Guardar notas",
            use_container_width=True,
            key=f"save_workspace_notes_{interview['id']}",
        ):
            save_interview_session_text(
                user,
                interview["id"],
                transcript,
                notes,
            )
            st.success("Notas guardadas.")

        if n2.button(
            "Analizar transcripción y notas",
            type="primary",
            use_container_width=True,
            key=f"analyze_workspace_notes_{interview['id']}",
        ):
            analysis = analyze_interview_text(
                user["company_id"],
                transcript,
                notes,
            )
            save_interview_session_text(
                user,
                interview["id"],
                transcript,
                notes,
                analysis,
            )
            st.session_state[
                f"workspace_analysis_{interview['id']}"
            ] = analysis
            st.success("Análisis generado.")

        analysis = st.session_state.get(
            f"workspace_analysis_{interview['id']}"
        )
        if not analysis:
            try:
                analysis = json.loads(
                    session.get("analysis_json") or "{}"
                )
            except json.JSONDecodeError:
                analysis = {}

        if analysis:
            st.caption(analysis.get("notice") or "")
            st.dataframe(
                pd.DataFrame(
                    [
                        {
                            "Competencia": item["competency"],
                            "Cobertura de evidencias": item["coverage"],
                            "Puntaje sugerido": item["suggested_score"],
                            "Crítica por defecto": (
                                "Sí" if item["critical_default"] else "No"
                            ),
                            "Indicadores encontrados": ", ".join(
                                item["matched_keywords"]
                            ),
                        }
                        for item in analysis.get("competencies", [])
                    ]
                ),
                use_container_width=True,
                hide_index=True,
            )

    with evaluation_tab:
        render_competency_methodology()
        guides = get_competency_guides(user["company_id"])
        saved_scores = fetch_all(
            """
            SELECT *
            FROM interview_competency_scores
            WHERE interview_id = ?
            """,
            (interview["id"],),
        )
        saved_by_name = {
            item["competency_name"]: item
            for item in saved_scores
        }

        session_now = get_interview_session(interview["id"])
        transcript_now = session_now.get("transcript_text") or ""
        notes_now = session_now.get("interviewer_notes") or ""

        scores = []
        with st.form(f"competency_evaluation_{interview['id']}"):
            for guide in guides:
                saved = saved_by_name.get(guide["name"], {})
                st.markdown(f"#### {guide['name']}")
                st.caption(guide.get("definition") or "")

                c1, c2 = st.columns([2, 1])
                score = c1.slider(
                    "Puntaje",
                    min_value=1,
                    max_value=5,
                    value=int(saved.get("score") or 3),
                    key=f"competency_score_{interview['id']}_{guide['id']}",
                )
                critical = c2.checkbox(
                    "Competencia crítica",
                    value=bool(
                        saved.get("critical")
                        if saved
                        else guide.get("critical_default")
                    ),
                    key=f"competency_critical_{interview['id']}_{guide['id']}",
                )
                st.write(
                    f"**Nivel seleccionado:** "
                    f"{COMPETENCY_SCALE[score]['level']}"
                )
                evidence = st.text_area(
                    "Evidencia observada",
                    value=saved.get("evidence_text") or "",
                    key=f"competency_evidence_{interview['id']}_{guide['id']}",
                )
                competency_notes = st.text_area(
                    "Observaciones",
                    value=saved.get("notes") or "",
                    key=f"competency_notes_{interview['id']}_{guide['id']}",
                )

                scores.append(
                    {
                        "competency": guide["name"],
                        "score": score,
                        "critical": critical,
                        "evidence": evidence,
                        "notes": competency_notes,
                    }
                )
                st.divider()

            submitted = st.form_submit_button(
                "Calcular resultado y guardar evaluación",
                type="primary",
                use_container_width=True,
            )

            if submitted:
                missing_evidence = [
                    item["competency"]
                    for item in scores
                    if not item["evidence"].strip()
                ]
                if missing_evidence:
                    st.error(
                        "Agregá evidencia para: "
                        + ", ".join(missing_evidence)
                    )
                else:
                    report = save_competency_scores(
                        user,
                        interview,
                        scores,
                        transcript_now,
                        notes_now,
                    )
                    st.session_state[
                        f"competency_report_{interview['id']}"
                    ] = report
                    st.success("Evaluación guardada.")
                    st.rerun()

    with report_tab:
        report = st.session_state.get(
            f"competency_report_{interview['id']}"
        )
        if not report:
            try:
                report = json.loads(
                    interview.get("competency_result_json") or "{}"
                )
            except json.JSONDecodeError:
                report = {}

        if not report:
            st.info(
                "Completá la evaluación por competencias para generar el informe."
            )
        else:
            color_map = {
                "VERDE": "🟢",
                "AMARILLO": "🟡",
                "ROJO": "🔴",
            }
            symbol = color_map.get(
                report.get("traffic_light"),
                "⚪",
            )
            r1, r2, r3 = st.columns(3)
            r1.metric("Promedio final", report.get("average"))
            r2.metric(
                "Semáforo",
                f"{symbol} {report.get('traffic_light')}",
            )
            r3.metric(
                "Recomendación",
                report.get("recommendation"),
            )
            st.info(report.get("conclusion") or "")

            if report.get("critical_alerts"):
                for alert in report["critical_alerts"]:
                    st.error(alert)

            st.markdown("#### Fortalezas")
            for item in report.get("strengths") or []:
                st.write(
                    f"• **{item['competency']}**: {item['score']}/5. "
                    f"{item.get('evidence') or ''}"
                )

            st.markdown("#### Competencias que cumplen lo esperado")
            for item in report.get("expected") or []:
                st.write(
                    f"• **{item['competency']}**: {item['score']}/5."
                )

            st.markdown("#### Aspectos de desarrollo")
            for item in report.get("development") or []:
                st.write(
                    f"• **{item['competency']}**: {item['score']}/5. "
                    f"{item.get('evidence') or ''}"
                )

            markdown_report = report_to_markdown(report)
            st.download_button(
                "Descargar informe",
                data=markdown_report.encode("utf-8"),
                file_name=(
                    f"informe_entrevista_{interview['id']}.md"
                ),
                mime="text/markdown",
                use_container_width=True,
            )


# =========================================================
# CENTRO DE ENTREVISTAS
# =========================================================

def render_interviews_redesigned(user: dict) -> None:
    st.subheader("Centro de entrevistas con ALBA")
    st.caption(
        "Prepará, asigná, grabá, evaluá y documentá la entrevista desde un único lugar."
    )

    if not has_permission("view_interviews"):
        st.warning("No tenés permiso para ver entrevistas.")
        return

    if has_permission("manage_interviews"):
        (
            home_tab,
            room_tab,
            guide_tab,
            templates_tab,
            assign_tab,
            tracking_tab,
        ) = st.tabs(
            [
                "Inicio",
                "Sala y evaluación",
                "Guía de competencias",
                "Plantillas",
                "Asignar",
                "Seguimiento",
            ]
        )

        with home_tab:
            st.markdown("### ¿Qué necesitás hacer?")
            c1, c2, c3 = st.columns(3)
            with c1:
                st.markdown(
                    """
                    <div class="alba-card">
                        <h3>1. Preparar</h3>
                        <p>Definí las competencias, preguntas y cuáles son críticas.</p>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
            with c2:
                st.markdown(
                    """
                    <div class="alba-card">
                        <h3>2. Entrevistar</h3>
                        <p>Mostrá a ALBA, grabá la sesión y tomá notas.</p>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
            with c3:
                st.markdown(
                    """
                    <div class="alba-card">
                        <h3>3. Evaluar</h3>
                        <p>Aplicá la escala 1–5 y generá el informe final.</p>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
            render_competency_methodology()

        with room_tab:
            render_interviewer_workspace(user)

        with guide_tab:
            render_competency_guide(user)

        with templates_tab:
            avatar_tab, question_tab = st.tabs(
                ["Avatar ALBA", "Preguntas y plantillas"]
            )
            with avatar_tab:
                render_avatar_settings(user)
                st.markdown("### Videollamada en tiempo real")
                st.write(
                    "Pegá en Streamlit Secrets la URL pública o el código "
                    "oficial de D-ID Agent Embed."
                )
                agent_example = (
                    'DID_AGENT_PUBLIC_URL = "https://..."\n'
                    'DID_AGENT_EMBED_HTML = "<script oficial de D-ID></script>"'
                )
                st.code(agent_example, language="toml")
            with question_tab:
                render_template_management(user)

        with assign_tab:
            render_interview_assignment(user)

        with tracking_tab:
            render_interview_monitoring(user)
    else:
        render_interview_monitoring(user)


def render_audit(user: dict) -> None:
    st.subheader("Auditoría")

    if not has_permission("view_audit"):
        st.warning("No tenés permiso para consultar auditoría.")
        return

    records = fetch_all(
        """
        SELECT
            audit_log.id,
            users.email AS user_email,
            audit_log.event_type,
            audit_log.entity_type,
            audit_log.entity_id,
            audit_log.details,
            audit_log.created_at
        FROM audit_log
        LEFT JOIN users
            ON users.id = audit_log.user_id
        WHERE audit_log.company_id = ?
        ORDER BY audit_log.created_at DESC
        LIMIT 500
        """,
        (user["company_id"],),
    )

    if records:
        st.dataframe(
            pd.DataFrame(records),
            use_container_width=True,
            hide_index=True,
        )
    else:
        st.info("Todavía no hay eventos de auditoría.")


# =========================================================
# PORTAL DE EMPRESA
# =========================================================

def render_company_portal(user: dict) -> None:
    company = get_company(user["company_id"])

    if not company:
        st.error("No se encontró la empresa vinculada.")
        return

    render_company_header(company)

    menu = ["🏠 Inicio", "💼 Búsquedas", "⚡ Asignación rápida"]

    if has_permission("view_candidates"):
        menu.append("👥 Candidatos")
    if has_permission("view_interviews"):
        menu.append("🎥 Entrevistas")
    if has_permission("manage_users"):
        menu.append("🔐 Usuarios")
    if has_permission("manage_company"):
        menu.append("🏢 Empresa")
    if has_permission("view_audit"):
        menu.append("🧾 Auditoría")

    if st.session_state.get("company_menu") not in menu:
        st.session_state["company_menu"] = "🏠 Inicio"

    selected = st.sidebar.radio(
        "Navegación",
        menu,
        key="company_menu",
    )

    st.sidebar.caption(f"Versión {APP_VERSION}")
    st.sidebar.markdown("---")
    st.sidebar.caption(
        "Flujo sugerido: búsqueda → candidatos → entrevista → decisión."
    )

    if selected == "🏠 Inicio":
        render_intuitive_dashboard(user, company)
    elif selected == "💼 Búsquedas":
        render_jobs(user)
    elif selected == "⚡ Asignación rápida":
        render_quick_candidate_assignment(user)
    elif selected == "👥 Candidatos":
        render_applications(user)
    elif selected == "🎥 Entrevistas":
        render_interviews(user)
    elif selected == "🔐 Usuarios":
        render_user_management(user)
    elif selected == "🏢 Empresa":
        render_company_settings(user, company)
    elif selected == "🧾 Auditoría":
        render_audit(user)

def render_candidate_portal(user: dict) -> None:
    candidate = fetch_one(
        "SELECT * FROM candidates WHERE user_id = ?",
        (user["id"],),
    )

    if not candidate:
        st.error("No se encontró el perfil del candidato.")
        return

    st.title("Portal del candidato")
    st.caption(user["email"])

    candidate_menu_options = [
        "Búsquedas abiertas",
        "Mis postulaciones",
        "Mis entrevistas",
        "Mi perfil y CV",
    ]

    if (
        st.session_state.get(
            "candidate_portal_menu"
        )
        not in candidate_menu_options
    ):
        st.session_state[
            "candidate_portal_menu"
        ] = "Búsquedas abiertas"

    menu = st.sidebar.radio(
        "Menú candidato",
        candidate_menu_options,
        key="candidate_portal_menu",
    )

    if menu == "Búsquedas abiertas":
        jobs = fetch_all(
            """
            SELECT
                jobs.id,
                jobs.company_id,
                jobs.title,
                jobs.area,
                jobs.seniority,
                jobs.description,
                jobs.location,
                jobs.work_mode,
                jobs.contract_type,
                jobs.responsibilities,
                jobs.must_have,
                jobs.desirable,
                jobs.competencies,
                companies.name AS company_name,
                companies.logo_blob
            FROM jobs
            JOIN companies
                ON companies.id = jobs.company_id
            WHERE jobs.status = 'ABIERTA'
            ORDER BY jobs.created_at DESC
            """
        )

        if not jobs:
            st.info("No hay búsquedas abiertas.")
            return

        for job in jobs:
            with st.container(border=True):
                c1, c2 = st.columns([1, 5])

                with c1:
                    if job.get("logo_blob"):
                        st.image(
                            job["logo_blob"],
                            width=80,
                        )

                with c2:
                    st.subheader(job["title"])
                    st.write(
                        f"**Empresa:** {job['company_name']}"
                    )
                    st.write(
                        f"**Área:** "
                        f"{job.get('area') or 'No informada'}"
                    )
                    st.write(
                        f"**Seniority:** "
                        f"{job.get('seniority') or 'No informado'}"
                    )
                    st.write(
                        f"**Modalidad:** "
                        f"{job.get('work_mode') or 'No informada'}"
                    )
                    st.write(
                        f"**Ubicación:** "
                        f"{job.get('location') or 'No informada'}"
                    )
                    st.write(
                        f"**Contratación:** "
                        f"{job.get('contract_type') or 'No informada'}"
                    )

                    if job.get("description"):
                        st.markdown("**Descripción**")
                        st.write(job["description"])
                    if job.get("responsibilities"):
                        st.markdown("**Responsabilidades**")
                        st.write(job["responsibilities"])
                    if job.get("must_have"):
                        st.markdown(
                            "**Requisitos excluyentes**"
                        )
                        st.write(job["must_have"])
                    if job.get("desirable"):
                        st.markdown(
                            "**Requisitos deseables**"
                        )
                        st.write(job["desirable"])
                    if job.get("competencies"):
                        st.markdown("**Competencias**")
                        st.write(job["competencies"])

                    already_applied = fetch_one(
                        """
                        SELECT id
                        FROM applications
                        WHERE candidate_id = ?
                          AND job_id = ?
                        """,
                        (
                            candidate["id"],
                            job["id"],
                        ),
                    )

                    if already_applied:
                        st.success("Ya te postulaste.")
                    elif st.button(
                        "Postularme",
                        key=f"apply_{job['id']}",
                    ):
                        application_id = execute(
                            """
                            INSERT INTO applications(
                                candidate_id,
                                job_id,
                                status,
                                created_at,
                                updated_at
                            )
                            VALUES (?, ?, 'RECIBIDA', ?, ?)
                            """,
                            (
                                candidate["id"],
                                job["id"],
                                now_iso(),
                                now_iso(),
                            ),
                        )

                        ensure_company_candidate(
                            job["company_id"],
                            candidate["id"],
                            "POSTULACIÓN_PORTAL",
                            user["id"],
                        )

                        score_application(application_id)

                        log_event(
                            None,
                            user["id"],
                            "CREATE_APPLICATION",
                            "application",
                            application_id,
                            {"job_id": job["id"]},
                        )

                        st.success("Postulación enviada.")
                        st.rerun()

    elif menu == "Mis postulaciones":
        applications = fetch_all(
            """
            SELECT
                jobs.title,
                companies.name AS company,
                applications.status,
                applications.created_at
            FROM applications
            JOIN jobs
                ON jobs.id = applications.job_id
            JOIN companies
                ON companies.id = jobs.company_id
            WHERE applications.candidate_id = ?
            ORDER BY applications.created_at DESC
            """,
            (candidate["id"],),
        )

        if applications:
            st.dataframe(
                pd.DataFrame(applications),
                use_container_width=True,
                hide_index=True,
            )
        else:
            st.info("Todavía no tenés postulaciones.")

    elif menu == "Mis entrevistas":
        render_candidate_interviews(user, candidate)

    else:
        render_candidate_profile(user, candidate)


# =========================================================
# APLICACIÓN
# =========================================================

st.set_page_config(
    page_title=APP_NAME,
    page_icon="🤖",
    layout="wide",
)


def render_fatal_error(
    title: str,
    exc: Exception,
) -> None:
    st.error(title)
    st.exception(exc)

    with st.expander(
        "Detalle técnico completo",
        expanded=True,
    ):
        st.code(traceback.format_exc())

    st.info(
        "Copiá o sacá una captura de este detalle. "
        "La pantalla ya no debería ocultar la causa del error."
    )


def run_application() -> None:
    inject_ui_styles()

    try:
        init_db()
    except Exception as exc:
        render_fatal_error(
            "No se pudo inicializar o actualizar la base de datos.",
            exc,
        )
        st.stop()

    invitation_token = st.query_params.get(
        "invite",
        "",
    )
    if isinstance(invitation_token, list):
        invitation_token = (
            invitation_token[0]
            if invitation_token
            else ""
        )

    if invitation_token:
        render_interview_invitation_access(
            str(invitation_token)
        )
        return

    user = current_user()

    if not user:
        st.title(APP_NAME)
        st.caption(
            "Plataforma de selección con portales para empresas "
            f"y candidatos · Versión {APP_VERSION}"
        )

        login_tab, register_tab, privacy_tab = st.tabs(
            ["Iniciar sesión", "Registrarse", "Privacidad"]
        )

        with login_tab:
            render_login()

        with register_tab:
            render_registration()

        with privacy_tab:
            st.subheader("Política de datos")
            st.write(
                "El correo electrónico funciona como usuario. "
                "Las contraseñas se guardan con hash y salt."
            )
            st.write(
                "DNI, teléfono y ubicación pueden registrarse para gestión, "
                "pero no deben utilizarse para puntuar o rankear."
            )
            st.write(
                "El análisis de coincidencia usa únicamente el contenido "
                "profesional del CV y los criterios de la búsqueda. "
                "La decisión final debe quedar en manos de una persona."
            )
            st.write(
                "Los filtros de edad y ubicación son administrativos y no "
                "forman parte del puntaje ni de la recomendación automática."
            )
            st.write(
                "El banco inteligente propone candidatos con base en la "
                "correspondencia profesional del CV. La asignación a una "
                "búsqueda siempre requiere una acción humana."
            )
            st.write(
                "La asistencia de entrevistas analiza solo texto o "
                "transcripciones. No evalúa voz, acento, apariencia, "
                "gestos ni emociones. Los audios quedan para revisión humana."
            )
            st.write(
                "El avatar ALBA se utiliza únicamente para presentar las "
                "preguntas y no observa ni puntúa al candidato."
            )
            st.write(
                "Las invitaciones se envían al correo detectado en el CV "
                "después de la revisión de RR. HH. Los enlaces son únicos, "
                "tienen vencimiento y no almacenan el token en texto plano."
            )
            st.write(
                "Las grabaciones requieren consentimiento y se utilizan "
                "exclusivamente para la revisión del proceso. La evaluación "
                "automática no analiza apariencia, voz, acento ni emociones."
            )
            st.write(
                "La empresa no debe cargar credenciales, secretos comerciales, "
                "precios, clientes, fórmulas ni procesos confidenciales."
            )

        return

    refreshed_user = refresh_user(user["id"])

    if not refreshed_user:
        st.session_state.pop("auth_user", None)
        st.error("El usuario fue desactivado.")
        st.stop()

    st.session_state["auth_user"] = refreshed_user
    user = refreshed_user

    st.sidebar.success(
        f"{user['full_name']}\n\n{user['email']}"
    )
    st.sidebar.caption(f"Versión {APP_VERSION}")

    audit_warning = st.session_state.pop(
        "audit_warning",
        None,
    )
    if audit_warning:
        st.warning(audit_warning)

    if st.sidebar.button("Cerrar sesión"):
        logout()

    if user["account_type"] == "COMPANY":
        render_company_portal(user)
    elif user["account_type"] == "CANDIDATE":
        render_candidate_portal(user)
    else:
        st.error("Tipo de cuenta no válido.")


try:
    run_application()

except Exception as exc:
    render_fatal_error(
        "Se produjo un error dentro de ALBA.",
        exc,
    )
